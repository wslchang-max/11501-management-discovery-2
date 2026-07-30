import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十二週 課程教學指引：個人與家庭資產負債表、現金流管理與生涯理財規劃")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十二週完整教學指引。本單元將大一新鮮人所學之各項個人理財知識整合為一套系統化的個人財務管理系統。課程深入剖析個人資產負債表 (Balance Sheet, BS) 核心公式（資產 = 負債 + 淨值）、流動/投資/使用性 3 類資產與良性/惡性負債拆解；掌握個人現金流量表 (Cash Flow, CF)、發薪日「先儲蓄後消費 (Pay Yourself First)」30/50/20 黃金分配法則；建立 3~6 個月緊急預備金水庫、區分需要與想要、雪崩法清理 15% 高利卡債；遠期量化人生 5 大財務目標、運用 20 歲開始每月 3,000 元時間複利槓桿，並介紹 F.I.R.E. 提早退休 4% 法則與自動化多帳戶分流管理，輔以主計總處家庭收支調查數據，打造終身受用之財務自主藍圖。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "個人資產負債表與淨值算術\n(Slide 01 - 10)", "個人財務儀表板、資產=負債+淨值公式、3類資產與2類負債拆解、新鮮人淨值計算與家戶資產統計", "🎯 1小時活動：個人資產負債表與淨值算術計算器"),
        ("第二小時\n(00:50-01:40)", "現金流管理、30/50/20 與 F.I.R.E.\n(Slide 11 - 20)", "先儲蓄後消費算術、30/50/20 黃金分配、3-6個月緊急預備金水庫、需要vs想要、雪崩還債法與 4% 法則", "🎯 2小時活動：30/50/20 現金流與緊急預備金試算器"),
        ("第三小時\n(01:40-02:30)", "財務健全度與生涯理財實戰\n(Slide 21 - 30)", "4大財務健康比例、多帳戶自動分流法、家庭財務分工、防詐、大學4年滾存30萬SOP與4大金律", "🎮 3小時小遊戲：生涯理財達人大挑戰 (4大關卡)\n📝 課堂實務作業：個人資產負債表編製與生涯理財計畫")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十二週課程導論：個人與家庭資產負債表、現金流管理與生涯理財規劃", 
         "本頁為第十二週課程導論。講師可以引用巴菲特名言開場：「如果你不知道如何編製個人的財務報表，你就像駕駛一架沒有儀表板的飛機！」說明個人理財不是靠直覺，而是需要建立一套清晰的「個人財務儀表板」。本單元將結合前半學期股票、ETF、外匯、保險與稅務知識，帶領大家學會編製個人資產負債表 (BS) 與現金流量表 (CF)，掌握先儲蓄後消費 30/50/20 黃金法則與 F.I.R.E. 提早退休水庫。",
         "行政院主計總處家庭收支調查 / CFP 個人財務規劃標準。",
         "問學生：『同學們現在知道自己身上到底有多少資產、欠了多少債務，以及你的個人實質淨值是多少嗎？』"),
        (2, 1, "第十二週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密個人資產負債表 (BS) 與淨值算術、3 類資產與 2 類負債拆解，以及現金流量表 (CF) 公式；第二小時聚焦於發薪日先儲蓄後消費 30/50/20 黃金法則、3-6 個月緊急預備金、雪崩法還債與 F.I.R.E. 4% 法則；第三小時傳授 4 大財務健全比例、自動化多帳戶分流、防詐、大學 4 年滾存 30 萬畢業第一桶金 SOP 與 4 大金律。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生拿出手機與筆記本，準備建立專屬的個人資產負債表草稿。"),
        (3, 1, "為什麼企業與個人都需要資產負債表？個人財務體檢兩大支柱",
         "講授評估個人財務健康的兩張核心報表：1. **個人資產負債表 (Balance Sheet, BS)**：紀錄特定時間點的「財務存量」 (Stock)，顯示資產、負債與實質淨值；2. **個人現金流量表 (Cash Flow, CF)**：紀錄特定期間的「財務流量」 (Flow)，顯示每個月收入來源、消費支出與淨盈餘。兩者相輔相成，共同構成個人財務健康診斷儀表板。",
         "理財規劃師 (CFP) 個人財務報表標準。",
         "提問：『月收入很高的人，資產負債表就一定健康嗎？（答案：不一定！如果消費過度或借高利卡債，淨值可能是負的）』"),
        (4, 1, "個人資產負債表核心公式：資產 = 負債 + 淨值",
         "講授個人財務會計黃金恆等式：`個人總資產 = 總負債 + 個人淨值`；`個人淨值 = 個人總資產 - 總負債`。核心觀念：開豪車、住豪宅不代表很有錢！如果總資產 1,500 萬中有 1,400 萬是銀行房貸與車貸，個人真實淨值只有 100 萬元！淨值才是個人真正擁有的實質淨財富。",
         "薩繆爾森《會計學與個人理財》基本恆等式。",
         "強調：『理財的核心目標不是衝大總資產，而是不斷拉高你的「個人淨值」！』"),
        (5, 1, "流動資產 vs. 投資資產 vs. 使用性資產拆解",
         "剖析個人資產的三大分類：1. **流動性資產 (Liquid)**：現金、活存、數位帳戶存款，極高流動性應對日常與緊急預備金；2. **投資性資產 (Investment)**：0050/VOO ETF、台積電股票、債券，能產生股息或資本增值；3. **使用性資產 (Use)**：自住房屋、通勤機車、筆電與手機，會隨時間折舊，無法生利息。",
         "理財規劃師 (CFP) 資產分類標準。",
         "請學生分類：『你手上最新的 iPhone 16 Pro 屬於哪一種資產？（答案：使用性資產，買了立刻折舊，不會生利息）』"),
        (6, 1, "短期負債 vs. 長期負債（信用卡卡債、學貸與房貸）",
         "講授如何區分良性負債 (Good Debt) 與惡性負債 (Bad Debt)：1. **短期高利惡性負債**：信用卡循環利息 (15%)、分期未繳、地下借貸，利息極高且無相應資產，會迅速侵蝕淨值；2. **中長期低利良性負債**：學生貸款 (低利 1.65%)、自住房貸 (2.18%)，利率低且能提升人力資本或換取增值資產。",
         "中華民國銀行公會負債管理原則。",
         "警告：『大一學生絕對不要陷入 15% 信用卡循環利息，那是財務毀滅的黑洞！』"),
        (7, 1, "個人淨值 (Net Worth) 算術：大一新鮮人淨值計算實例",
         "展示大一學生小明 18 歲個人資產負債表編製實例：資產端（高利活存 3 萬 + 0050 ETF 5 萬 + 中古機車 3 萬 = 總資產 11 萬元）；負債與淨值端（機車分期 2 萬 + 學貸 4 萬 = 總負債 6 萬元）；**小明個人實質淨值 (Net Worth) = 11萬 - 6萬 = 5 萬元！**",
         "115管理探索二教案財務編製組。",
         "鼓勵學生：『哪怕剛滿 18 歲淨值只有幾萬元，只要淨值是正的，就是非常棒的起點！』"),
        (8, 1, "實證數據：金管會與主計總處台灣家戶平均資產負債與淨值統計",
         "引述行政院主計總處 2026 最新國情統計通報數據：全台家庭平均每戶資產總額約 1,630 萬元新台幣（房地產占 52%、金融資產占 43%）；平均每戶負債總額約 240 萬元（以房貸為主）；**全台家庭平均每戶淨值高達 1,390 萬元新台幣**，中位數淨值約 900 萬元，呈現台灣社會的總體財富積累水準。",
         "行政院主計總處家庭收支調查與國情統計通報。",
         "提問：『看完台灣家戶平均淨值數據，這對我們大一新鮮人設定未來 10 年理財目標有什麼啟示？』"),
        (9, 1, "個人現金流量表 (Cash Flow) 核心公式：收入 - 支出 = 淨現金流",
         "講授現金流量表公式：`月淨現金流 = 每月總收入 (打工+獎學金+零用錢) - 每月總支出 (食宿+交通+娛樂)`。分析兩大狀態：1. **正現金流 (Positive CF)**：收入 > 支出！產出月盈餘轉入 BS 買 0050 ETF 增加淨值；2. **負現金流 (Negative CF)**：收入 < 支出！入不敷出消耗儲蓄或借卡債，淨值縮水。",
         "理財規劃師 (CFP) 現金流控管標準。",
         "強調：『正現金流是打造個人財富水庫的唯一活水來源！』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：BS 公式（資產 = 負債 + 淨值）；3 類資產與 2 類負債拆解；正現金流產出淨盈餘。預告第 1 小時 Modal 實務活動——「個人資產負債表與淨值算術計算器」，請學生輸入個人存款、股票與負債金額，精算個人淨值。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行資產負債表試算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "翻轉理財順序：發薪日「先儲蓄、後消費 (Pay Yourself First)」算術",
         "剖析月光族存不到錢的根本原因並翻轉理財公式：錯誤公式：`收入 - 消費 = 儲蓄 (剩多少存多少)`，往往月底花光歸零；致富公式：`收入 - 預先儲蓄投資 = 可消費金額 (Pay Yourself First)`！發薪日第一天自動將 20%-30% 轉入投資帳戶，強迫自己用剩餘的錢過日子。",
         "巴菲特理財哲學 \"Do not save what is left after spending\"。",
         "提問：『為什麼發薪當天先自動轉帳儲蓄，比到了月底再看剩多少錢更容易存到錢？（因為約束人性，防止過度消費）』"),
        (12, 2, "大學生打工資產分配黃金法則：30 / 50 / 20 比例法則",
         "講授 30/50/20 個人預算黃金分配模型：以打工月薪 15,000 元為例：1. **50% 必要生活費 (Needs)**：7,500 元（房租、三餐、交通公車票）；2. **30% 預先儲蓄與投資 (Savings)**：4,500 元（定期定額買 0050 ETF + 預備金）；3. **20% 樂享想要 (Wants)**：3,000 元（聚餐、旅遊、遊戲娛樂）。",
         "理財規劃師 (CFP) 30/50/20 個人預算分配模型。",
         "算一算：『如果你的打工月薪是 20,000 元，按照 30/50/20 法則，你每個月應該強制存下多少錢？（答案：6,000 元！）』"),
        (13, 2, "緊急預備金 (Emergency Fund) 水庫：留足 3 ~ 6 個月生活費",
         "講授緊急預備金水庫打造法則：專門應對突發失業、機車故障修車或生病事故。公式：`緊急預備金水庫 = 每月必要生活費 (如 1.2 萬) * 3 ~ 6 個月 = 3.6 萬 ~ 7.2 萬元`。存放管道：必須放在數位帳戶 2.5% 高利活存，兼顧極高流動性與利息，絕不能拿去買股票。",
         "理財規劃師 (CFP) 緊急預備金標準。",
         "強調：『沒有存夠 3 個月緊急預備金之前，絕對不要把手上的現金全部投入股市！』"),
        (14, 2, "辨識「需要 (Needs)」vs.「想要 (Wants)」：記帳 App 與消費坑",
         "教導辨識「需要」與「想要」的消費心法：1. **需要 (Needs)**：生存與課業必需品（餐費、通勤票、教科書）；2. **想要 (Wants)**：滿足慾望的非必需品（150 元星巴克、最新 iPhone、限量潮鞋）。介紹**延遲享樂法則 (Delayed Gratification)**：想買非必需品時，強制冷靜 48 小時，90% 的衝動慾望會自動消退。",
         "行為經濟學《過度消費心理學》。",
         "分享：『記住，省下一杯 150 元手搖飲，就能在 0050 ETF 中多買一股零股！』"),
        (15, 2, "債務管理優先順序：雪球法 (Snowball) vs. 雪崩法 (Avalanche)",
         "剖析兩種科學化還債策略：1. **數學雪崩法 (Avalanche)**：優先還清「利率最高」的債務（如 15% 信用卡卡債 ➔ 機車分期 ➔ 1.65% 學貸），數學上最省總利息負擔；2. **心理雪球法 (Snowball)**：優先還清「金額最小」的債務，快速消除債務筆數，獲得強烈成就感。",
         "Dave Ramsey《Total Money Makeover》還債架構。",
         "建議：『理性理財首選雪崩法，先把 15% 的信用卡高利卡債清掉，節省最多利息費用！』"),
        (16, 2, "學生貸款 (Student Loans) 緩繳條款、利息補貼與理性償還",
         "解密政府學生貸款理財策略：在學期間利息由政府全額補貼，畢業滿 1 年後開始還本付息。學貸利率極低（僅約 1.65%），遠低於通膨率與股票長期報酬率。理性策略：有餘錢時優先定期定額買 0050 ETF (年化 8%)，按月正常還學貸即可，不必急著提前全數清償！",
         "教育部高級中等以上學校學生就學貸款作業要點。",
         "提問：『學貸利率 1.65%，投資 0050 年化 8%，為什麼提前全額還清學貸反而是次優選擇？（答案：機會成本！投資利差高達 6.35%）』"),
        (17, 2, "人生 5 大財務目標：畢業積蓄、購車、結婚、買房與退休",
         "引導學生按時間軸量化人生 5 大財務目標：1. **短期目標 (1-4 年)**：大學畢業第一桶金 (30 萬)、緊急預備金水庫；2. **中期目標 (5-10 年)**：結婚創業預算、買房自備頭期款 (200 萬)；3. **長期目標 (15-40 年)**：子女教育基金、退休金累積 (F.I.R.E. 1,500 萬元水庫)。",
         "理財規劃師 (CFP) 生涯目標量化模型。",
         "請學生在講義上記下自己大學四年畢業時希望達成的第一桶金目標金額。"),
        (18, 2, "時間複利 (Compound Interest) 槓桿：20歲每月 3,000 元的百萬威力",
         "演示時間複利的驚人威力：從 20 歲大一開始，每月定期定額 **3,000 元** 投入年化報酬率 8% 之 ETF，30 年後 (50 歲) 總投入本金僅 108 萬元，但總資產累積至 **4,470,000 元 (447 萬元！)**。結論：在複利公式中，**「時間」比「本金」更具殺傷力！**",
         "理財規劃師 (CFP) 複利時間槓桿試算。",
         "強調：『大一學生最大的財富資產不是銀行存款，而是你擁有整整 40 年的時間複利槓桿！』"),
        (19, 2, "財務獨立提早退休 (F.I.R.E.)：4% 法則 (4% Rule) 資產水庫",
         "解密 F.I.R.E. 運動與 Trinity Study 4% 法則：`F.I.R.E. 退休目標資產 = 每年預計生活費 (如 60 萬) * 25 = 1,500 萬元`。當你累積到 1,500 萬投資在全市場 ETF 時，每年提領 4% (60 萬元) 作為生活費，剩餘資產持續抗通膨滾存，實現提早退休不工作也衣食無憂的財務自由！",
         "Trinity Study 4% Rule 財務自由模型。",
         "算一算：『如果你的目標是退休後每月有 4 萬元生活費 (一年 48 萬)，你的 4% 法則資產水庫是多少？（答案：48萬 × 25 = 1,200 萬元）』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：發薪日先儲蓄（30/50/20 分配）；留足 3-6 個月緊急預備金；雪崩法還卡債；學貸按月還即可；20 歲每月 3,000 元時間複利滾存 447 萬；4% 法則計算 FIRE 水庫。預告第 2 小時 Modal 實務活動——「30/50/20 現金流與緊急預備金試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行現金流與預備金試算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "個人財務健全度 4 大比例指標（儲蓄率、負債比、流動比、投資比）",
         "講授個人財務健檢四大比例指標：1. **月儲蓄率**：儲蓄投資 / 月收入 ➔ **>= 20%~30%**；2. **高利負債比**：高利卡債 / 月收入 ➔ **嚴格保持 0%**；3. **流動性比率**：流動資產 / 月支出 ➔ **保持 3~6 個月**；4. **投資資產比**：投資資產 / 個人淨值 ➔ **>= 50%**（避免全留現金被通膨侵蝕）。",
         "理財規劃師 (CFP) 個人財務健檢標準表。",
         "請學生比對自己的財務數據：『你的高利負債比是不是保持在健康的 0%？』"),
        (22, 3, "打造自動化個人財務水庫：多帳戶分流管理法",
         "教學打造零意志力自動化個人財務水庫：1. **帳戶 A (發薪主帳戶)**：打工薪資入帳，發薪日自動設定轉帳；2. **帳戶 B (投資累積帳戶 30%)**：自動扣款買 0050 ETF，只進不出；3. **帳戶 C (預備金帳戶)**：存放 3-6 個月 2.5% 高利活存；4. **帳戶 D (生活消費帳戶 70%)**：綁定簽帳卡處理日常開銷。",
         "個人自動化理財系統設計架構。",
         "鼓勵學生：『用系統約束人性，發薪日讓銀行自動幫你轉帳理財！』"),
        (23, 3, "婚前與家庭財務管理：雙薪家庭共同基金與財務分工",
         "講授雙薪家庭與伴侶財務管理機制：1. **按收入比例分攤法**：若薪資比為 6:4，共同開銷亦按 6:4 比例投入家庭公積金；2. **設立家庭共同公積金帳戶**：處理房貸、水電、餐費與子女教育；3. **保留個人獨立自由帳戶**：公積金之外，各自保留獨立帳戶自由支配，兼顧共同目標與個人自主。",
         "家庭理財與婚姻財務諮商標準指南。",
         "提問：『為什麼伴侶之間公開透明討論財務，並保留個人獨立自由帳戶能讓婚姻更長久？』"),
        (24, 3, "個人財務資安：防範投資詐騙社群飆股群組與資金盤 (Ponzi)",
         "警告詐騙集團侵蝕個人淨值的陷阱：1. **詐騙 3 特徵**：宣稱「保證獲利穩賺不賠」、要求加入 LINE/Telegram 飆股群組、要求匯款至個人不知名帳戶；2. **理性投資鐵律**：凡宣稱保證年化報酬率 > 10% 且無風險者，100% 是詐騙資金盤 (Ponzi Scheme)！堅決只在金管會核准之券商交易。",
         "內政部警政署 165 反詐騙專線警訊。",
         "警告：『天下沒有白吃的午餐，保證獲利 10% 的飆股群組 100% 是詐騙！』"),
        (25, 3, "實證數據調取：主計總處家庭收支調查 (dgbas.gov.tw) 查閱",
         "手把手教導學生查詢官方 Portal：1. **查詢可支配收入**：登入 `dgbas.gov.tw` 查詢全台前 20% 高收入戶 vs 後 20% 低收入戶收支差距；2. **查詢消費結構**：查閱台灣家戶在食品、居住、醫療與教育上的平均消費比例；3. **個人財務定位**：比對個人收支與全台同齡層數據，設定客觀理財目標。",
         "行政院主計總處家庭收支調查官方 Portal。",
         "演示進入主計總處網站查閱全台家庭收支統計圖表的步驟。"),
        (26, 3, "大一新鮮人個人生涯理財 4 大金律",
         "總結生涯理財 4 大金律：1. **定期盤點資產負債表**（每季編製個人 BS，專注提升「個人淨值」）；2. **堅持先儲蓄後消費**（發薪日自動將 30% 買 0050 ETF，留足 3-6 個月預備金）；3. **善用時間複利槓桿**（20 歲開始每月 3,000 元定期定額）；4. **遠離 15% 高利卡債**（每月全額繳清信用卡）。",
         "金管會與理財規劃師 (CFP) 保護原則。",
         "請學生齊聲朗讀四金律，深化理性理財信念。"),
        (27, 3, "理財試算實例：大學四年滾存 30 萬元畢業第一桶金 SOP",
         "展示大一到大四每月儲蓄 5,000 元滾存 30 萬畢業第一桶金 SOP：打工月薪 15,000 元，按 30% 比例先儲蓄 5,000 元，定期定額買入 0050 / VOO 全市場市值型 ETF (年化 8%)。4 年累積成果：`5,000 元 * 48 個月 = 24 萬本金 + 複利收益 6 萬元 = 30 萬元畢業第一桶金！`",
         "115管理探索二教案資產累積組。",
         "強調：『大學四年只要月存 5,000 元買 ETF，畢業時你就能握有 30 萬元的傲人第一桶金！』"),
        (28, 3, "生涯財務自主與終身財富自由藍圖",
         "總結理財的終極價值：財富自由不是追求奢侈炫耀，而是為自己與家人換取**時間自由、安全感與人生選擇權**。透過健全的淨值與正現金流，你不必被迫做違心工作，能勇敢追求職涯與生活夢想，掌控個人命運！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『掌控你的財務儀表板，創造屬於你自己的富足自由人生！』"),
        (29, 3, "第十二週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十二週 30 頁純教學卡片進行整體串聯：Hour 1 (BS vs CF 兩大支柱 ➔ 資產=負債+淨值 ➔ 3類資產 ➔ 2類負債 ➔ 新鮮人淨值計算 ➔ 全台家戶資產統計 ➔ 淨現金流) ➔ Hour 2 (先儲蓄後消費 ➔ 30/50/20 黃金法則 ➔ 3-6個月緊急預備金 ➔ 需要vs想要 ➔ 雪崩法還債 ➔ 學貸理性償還 ➔ 人生5大目標 ➔ 20歲時間複利 ➔ FIRE 4%法則) ➔ Hour 3 (4大財務健康比例 ➔ 自動化多帳戶分流 ➔ 雙薪家庭財務分工 ➔ 165防詐 ➔ 主計總處Portal ➔ 大學4年滾存30萬SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十三週預告 (半導體產業、護國神山台積電與台灣科技供應鏈)",
         "恭喜學生完成第十二週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人資產負債表編製與生涯理財計畫」。預告第十三週課程主題：「半導體產業、護國神山台積電與台灣科技供應鏈」，下週將帶大家解密台灣護國神山台積電 (2330) 與半導體矽島產業鏈！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十三週半導體產業主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：個人資產負債表與淨值算術計算器",
         "1. 活動目標：幫助學生盤點個人存款、股票、機車與負債，計算個人真實淨值。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入流動資產、投資資產、使用資產與長短期負債，點擊精算。\n"
         "3. 診斷反思：系統算出個人淨值，深化「淨值才是真正拥有的財富」之核心觀念。"),

        ("🎯 第 2 小時實務活動：30/50/20 現金流與緊急預備金試算器",
         "1. 活動目標：讓學生親自精算個人打工月收入之 30/50/20 黃金預算分配與 3~6 個月緊急預備金水庫目標。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入打工月收入，點擊精算。\n"
         "3. 決策學習：系統自動產出 Needs (50%)、Savings (30%) 與 Wants (20%) 具體金額，並計算預備金目標，建立自動化儲蓄心態。"),

        ("🎮 第 3 小時小遊戲：生涯理財達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：個人淨值公式；關卡 2：理財成功黃金公式 Pay Yourself First；關卡 3：30/50/20 預算分配法則；關卡 4：F.I.R.E. 4% 法則目標資產）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 生涯理財達人徽章 (Financial Life Master)」，未滿分獲頒「🥉 理財初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人資產負債表編製與生涯理財計畫",
         "1. 作業題目：請學生盤點個人當前資產與負債，線上填寫個人資產負債表與大學畢業滾存 30 萬元目標計畫。\n"
         "2. 分析要項：(1) 編製當前個人資產負債表；(2) 依據 30/50/20 法則規劃每月 30% 儲蓄投資買 0050 ETF；(3) 設定大學四年滾存 30 萬畢業第一桶金 SOP；(4) 擬定多帳戶自動分流與卡債清理 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十二週_課程教學指引_個人與家庭資產負債表現金流管理與生涯理財規劃.docx'
    doc.save(doc_path)
    print("Created 第十二週_課程教學指引_個人與家庭資產負債表現金流管理與生涯理財規劃.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
