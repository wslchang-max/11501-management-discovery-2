import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十八週 課程教學指引：學期總結、財經探索藍圖與大一新鮮人終身宣言")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十八週大結局完整教學指引。本單元進行全學期 18 週課程大閉環總結（從大一理財健檢、三大報表、0050/00878 定期定額、不動產租買、所得稅報稅，跨越至半導體台積電 3 大護城河、生成式 AI 算力、ESG 碳費與 Web3 冷錢包）；協助大一新鮮人擬定大學 4 年階梯式財經成長藍圖（大一打底 ➔ 大二實務 ➔ 大三實習 ➔ 大四獨立 30 萬積蓄）；精算 20 歲起月存 5,000 元於 31 歲擁有人生第一個 100 萬第一桶金的複利算術；總整合 4 大核心理財公式、3 大前沿科技產業鏈與資安防禦金鐘罩；舉行莊嚴的大一新鮮人「終身財經宣言」宣讀與數位簽署；解析通往 F.I.R.E. 財務自由的 5 大里程碑，並頒發 3 學分通識核心數位結業證書與學習歷程檔案，為大一新鮮人開啟璀璨富足的人生新篇章。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "18週大回顧、4年藍圖與第一桶金\n(Slide 01 - 10)", "18週課程大回顧、管理核心靈魂、大學4年藍圖 (打底/實務/實習/獨立)、100萬複利算術、FOMO克服與期末頒獎", "🎯 1小時活動：大學 4 年財經藍圖與百萬第一桶金試算器"),
        ("第二小時\n(00:50-01:40)", "核心公式總整合、終身宣言與結業證書\n(Slide 11 - 20)", "18週理財公式總整合、產業管理總整合、資安防詐總整合、終身宣言簽署、5大 FIRE 里程碑與結業證書頒發", "🎯 2小時活動：18 週全景知識能力診斷與結業榮譽試算器"),
        ("第三小時\n(01:40-02:30)", "終身學習指引、4大心態與圓滿落幕\n(Slide 21 - 30)", "終身學習資源 (WSJ/FT)、勝出4大心態、導師通道、成績單下載、4大誓言、30年複利SOP與圓滿落幕", "🎮 3小時小遊戲：18 週全景財經通識總大挑戰 (4大關卡)\n📝 課堂實務作業：個人大學 4 年財經探索藍圖與結業感言")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十八週課程導論：學期總結、財經探索藍圖與大一新鮮人終身宣言", 
         "本頁為第十八週大結局課程導論。講師可以引用華倫·巴菲特名言開場：「最好的投資就是投資你自己！知識會像複利一樣隨時間滾雪球發揮驚人威力！」針對大一新鮮人經歷一整學期 18 週學習後的圓滿收官切入。說明本單元將進行全學期 18 週大總結、頒發期末專案獎項、打造大學 4 年財經探索藍圖、試算第一個 100 萬第一桶金，並共同簽署「大一新鮮人終身財經宣言」。",
         "115學年度「管理探索二」課程委員會終章白皮書。",
         "問學生：『同學們還記得 18 週前第一天來到這門課時，自己的財務目標是什麼嗎？現在完成了多少？』"),
        (2, 1, "第十八週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密 18 週大回顧、管理核心靈魂、大學 4 年藍圖 (打底/實務/實習/獨立)、100 萬複利算術、克服 FOMO 與期末頒獎；第二小時聚焦於 18 週理財/產業/資安公式總整合、終身宣言簽署、5 大 FIRE 里程碑與結業證書頒發；第三小時傳授終身學習資源 (WSJ/FT)、勝出 4 大心態、導師通道、成績單下載、4 大誓言與圓滿落幕。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上記下 18 週的核心收穫，準備開啟大學四季的財經藍圖。"),
        (3, 1, "18 週課程大回顧：從大一記帳到 Web3 與 ESG 的財經壯遊",
         "梳理 18 週完整的知識演進脈絡：1. **W1-W3 (理財基礎與槓桿)**：大一財務健檢、資產負債表、好債壞債與時間複利；2. **W4-W8 (商業經營與報表)**：損益表、三大報表、營運資金與訂價策略；3. **W9-W12 (投資與生涯)**：0050 ETF、不動產租買、所得稅報稅與生涯現金流；4. **W13-W16 (產業前沿與資安)**：半導體台積電、生成式 AI、ESG 碳費與 Web3 冷錢包；5. **W17-W18 (專案發表與總結)**：簡報 Pitching 與終身宣言！",
         "115學年度「管理探索二」全學期地圖。",
         "帶領學生快速點過 18 週每一個重要里程碑，展現豐碩的學習脈絡。"),
        (4, 1, "財經與管理的核心靈魂：將知識轉化為翻轉人生的終身實踐力",
         "講授「管理探索二」的核心靈魂——**實踐 (Execution)**：知識若不上台演練、不落實執行，就只是紙上談兵！1. **終身紀律執行**：將 Pay Yourself First 提撥 20% 與 0050/00878 定期定額內化為每個月的自動化習慣；2. **獨立思辨與資安防禦**：不盲從追高殺跌，以數據為依據，嚴防詐騙。",
         "彼得·杜拉克《管理的實踐》。",
         "強調：『知識不是力量，執行的知識才是真正的力量！』"),
        (5, 1, "構建大學四年財經探索藍圖：打底 ➔ 實務 ➔ 實習 ➔ 獨立",
         "為大一新鮮人制定階梯式大學 4 年財經成長藍圖：1. **大一 (打底期)**：建置 3-6 個月預備金，開啟 0050 定期定額，累積 3-5 萬；2. **大二 (實務期)**：考取初階外匯或 ISO 14064 碳盤查證照，精進 Excel/Python；3. **大三 (實習期)**：進入金控或科技業實習；4. **大四 (獨立期)**：達成 30 萬第一桶金，具備完全財務獨立！",
         "115學年度「管理探索二」職涯探索小組。",
         "指導學生在講義上記下自己大學 4 年每一個階段要完成的具體指標。"),
        (6, 1, "人生第一桶金 (100 萬) 時間複利算術：從 20 歲月存 $5,000",
         "解密小資族在 31 歲前攢下 100 萬元的複利神奇公式：`月存 5,000 元 * 12 個月 + 7% 複利 (0050/00878) * 11 年 = 1,048,520 元台幣！` 說明時間的驚人威力：20 歲開始月存 5,000 元，比 30 歲才開始月存 10,000 元輕鬆翻倍！",
         "理財規劃師 (CFP) 複利試算模型。",
         "算一算：『20 歲每天省下 166 元（相當於兩杯咖啡），31 歲就能擁有一百萬第一桶金！』"),
        (7, 1, "理財與心理學：克服 FOMO 追高殺跌陷阱與建構抗脆弱心態",
         "剖析理財與行為金融心理學：理財最大的敵人不是市場，而是鏡子裡的自己！1. **克服 FOMO (害怕錯過)**：看到追高殺跌盲目全押，遇到大跌低點賤賣；2. **建構抗脆弱 (Antifragile) 心態**：堅守定期定額扣款紀律，把市場大跌視為打折大拍賣優惠進場點，波瀾不驚！",
         "塔勒布《抗脆弱》與 Behavioral Finance 心理學。",
         "提醒：『當市場大跌時，理智的人看到的是打折拍賣，恐慌的人看到的是末日！』"),
        (8, 1, "實證數據：全班 18 週學習成效統計與前後知識對比",
         "展示全班 18 週學習成長問卷調查數據：1. **理財知識及格率**：從 W1 的 **32%** 飆升至 W18 的 **98%！** 2. **個人預備金與定期定額開戶率**：全班超過 **85%** 同學已成功建立緊急預備金並開啟定期定額！3. **AI 與資安防禦素養**：100% 通過 Deepfake 與 165 反詐測試。",
         "115學年度「管理探索二」全班學習前/後問卷統計。",
         "與全班同學共同歡呼，慶祝巨大的知識成長與實踐突破！"),
        (9, 1, "頒發期末專案卓越獎項：最佳財務模型獎與最佳團隊獎",
         "隆重頒發 115 學年度期末專案三大桂冠獎項：1. **🏆 最佳財務模型獎**：頒給試算表公式最嚴謹、數據權威之組別；2. **🎨 最佳視覺簡報獎**：頒給莫蘭迪大數字與圖表設計最卓越之組別；3. **👑 最佳發表團隊大獎**：由專家與同儕加權總分第一名獲此殊榮！",
         "115學年度「管理探索二」期末獎項評選委員會。",
         "邀請獲獎組別上台領獎，全班給予最熱烈的掌聲！"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：18 週全景大回顧；管理的實踐靈魂；大學 4 年藍圖（打底/實務/實習/獨立）；20 歲月存 5,000 元於 31 歲擁有一百萬；抗脆弱心態與期末頒獎。預告第 1 小時 Modal 實務活動——「大學 4 年財經藍圖與百萬第一桶金試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行百萬第一桶金達成年齡與 4 年藍圖試算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "18 週核心理財公式總整合：Pay Yourself First 與 4% 規則",
         "總整合 18 週 4 大理財黃金算術公式：1. **先儲蓄後消費**：`收入 - 儲蓄 (20%) = 可支配支出`；2. **50/30/20 預算法則**：`50% 必要 + 30% 想要 + 20% 投資`；3. **4% FIRE 退休金池**：`年度生活費 * 25 (60萬/年 -> 1,500萬)`；4. **72 法則倍增時間**：`72 / 年化報酬率 r% (7.2% -> 10 年翻倍)`。",
         "115管理探索二理財公式總表。",
         "帶領學生複習並熟記四大黃金公式，作為終身理財的指針。"),
        (12, 2, "18 週核心產業管理總整合：台積電、AI 3層鏈與 ESG 3柱",
         "總整合全球三大科技產業管理架構：1. **半導體台積電 (2330)**：純晶圓代工模式 + 技術領先/製造卓越/客戶信任 3 大護城河；2. **生成式 AI 3 層產業鏈**：底層算力 (NVIDIA/廣達) ➔ 中層大模型 (OpenAI) ➔ 上層 Copilot 應用；3. **綠色金融與 ESG 3 柱**：E (RE100/碳費) + S (DEI) + G (治理) ➔ 00878 ESG 配置。",
         "115管理探索二產業管理架構。",
         "總結：『掌握這三大科技產業鏈架構，你在未來看待全球產業趨勢時將具備高瞻遠矚的視野！』"),
        (13, 2, "18 週資安與防詐防線總整合：165、Deepfake 與冷錢包鐵律",
         "總整合個人財產資安三大金鐘罩：1. **165 反詐反查機制**：遇到急用借錢或投資保證獲利，100% 先掛斷撥打 165 或原號碼反查；2. **Deepfake 家族暗語**：與家人設定秘密通關暗語，破解 AI 語音/換臉複製詐騙；3. **冷錢包私鑰主權**：「Not Your Keys, Not Your Coins!」長期加密資產離線保管私鑰。",
         "內政部警政署 165 與 NICS 區塊鏈資安守則。",
         "齊聲複習：『165 反查、Deepfake 暗語、冷錢包私鑰，守護我們的血汗錢！』"),
        (14, 2, "大一新鮮人「終身財經宣言 (Freshman Manifesto)」宣讀與簽署",
         "莊嚴宣讀並數位簽署《大一新鮮人終身財經宣言》：「我承諾：1. 堅持 Pay Yourself First，每月提撥 20% 投資未來；2. 堅持量入為出，絕不染指高利信貸與盲目開槓桿；3. 堅持定期定額配置 0050/00878 等全市場 ETF；4. 堅守資安防線，用資本支持 ESG 永續好企業，成為有責任的現代財經公民！」",
         "115學年度「管理探索二」全體師生共同宣誓。",
         "全班同學全體起立，共同莊嚴宣讀終身財經宣言並點擊數位簽署。"),
        (15, 2, "通往財務自由的 5 大里程碑：預備金 ➔ 零債 ➔ 組合 ➔ 被動 ➔ FIRE",
         "解析登頂財務自由大樓的 5 階階梯：1. **里程碑 1 (預備金)**：存妥 3-6 個月高利活存 ($5萬)；2. **里程碑 2 (零壞債)**：清償高利卡債與信貸；3. **里程碑 3 (資產組合)**：建立 0050+00878 自動化定期定額；4. **里程碑 4 (被動涵蓋必要)**：被動收入 > 每月基本支出 50%；5. **里程碑 5 (F.I.R.E. 達成)**：資產池達 25 倍年支出！",
         "F.I.R.E. 財務自由運動階段定義。",
         "引導學生檢視自己目前處於哪一個里程碑，並訂下邁向下一階梯的目標。"),
        (16, 2, "告別大學小白：從盲目跟風者進化為獨立思辨財經公民",
         "對比大學四年看待財經與投資的思維轉變：1. **過去「理財小白」**：聽 Dcard/Line 群組明牌追高殺跌，看新聞炒作大跌就哭喊斷頭；2. **現在「獨立財經公民」**：登入 MOPS 查閱真實毛利率，看 10-K 年報分析 CapEx 資本支出，堅守 50/30/20 與定期定額紀律！",
         "115學年度「管理探索二」教學目標。",
         "恭喜學生告別理財小白，進化為具備獨立思辨能力的現代財經公民！"),
        (17, 2, "感謝與感動：致 18 週堅持不懈、勇於探索的自己與夥伴",
         "回首 18 週前的陌生與今天的蛻變：每一位同學都完成了 18 週 30 頁純教學卡片、3 大小時 Modal 試算、畫布演練與期末簡報發表！給自己與同班夥伴最熱烈的掌聲，這份堅持與勇氣將成為大學最珍貴的資產。",
         "115學年度「管理探索二」師生致謝紀錄。",
         "請全班同學轉頭對左右兩邊的夥伴說一聲：「謝謝你 18 週的陪伴與鼓勵！」"),
        (18, 2, "大一新鮮人終身受用的 10 大財經黃金金律總複習",
         "總複習 18 週淬鍊出的 10 大終身金律：1. Pay Yourself First 20%；2. 3-6個月緊急預備金；3. 0050/00878 定期定額；4. 租房 48 小時冷靜期與 $R/M < 60\%$；5. 房貸負擔率 $< 33\%$ 30 年均攤；6. 台積電毛利率 53% 與 CapEx；7. GenAI Prompt 4 要素；8. ESG 認明 GRI/SASB；9. Web3 冷錢包 5% 比率；10. 165 與 Deepfake 暗語！",
         "115學年度「管理探索二」10 大黃金金律。",
         "請學生齊聲朗讀 10 大黃金金律，將其銘刻於心。"),
        (19, 2, "學習歷程數位憑證頒發：頒發「管理探索二」結業證書",
         "頒發 115 學年度「管理探索二」3 學分通識核心數位結業證書：恭喜全班同學正式通過審核！每位同學均可於網頁系統下載具備專屬 Hash 雜湊驗證的數位結業證書，直接納入個人大一學習歷程 Portfolio！",
         "115學年度「管理探索二」教務處認證憑證。",
         "演示進入系統下載專屬數位結業證書 PDF 的步驟。"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：18 週理財/產業/資安公式總整合；終身財經宣言簽署；5 大 FIRE 里程碑；告別理財小白；10 大黃金金律；結業證書頒發。預告第 2 小時 Modal 實務活動——「18 週全景知識能力診斷與結業榮譽試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，進行 18 週全景能力診斷與榮譽分數精算。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "畢業後的終身學習指引：優質財經書籍、WSJ/FT 與數據源",
         "提供畢業後的終身學習權衛資源：1. **必讀經典**：《富爸爸窮爸爸》、《漫步華爾街》、《致富心態》、《金字塔原理》；2. **國際媒體**：華爾街日報 (WSJ)、金融時報 (FT)、Bloomberg；3. **官方數據 Portal**：MOPS、TWSE 證交所、NVIDIA IR、TCX 碳交所。",
         "115學年度「管理探索二」推薦閱讀書單。",
         "鼓勵學生養成每天閱讀一篇 WSJ/FT 財經新聞的終身習慣。"),
        (22, 3, "大一新鮮人終身勝出 4 大心態：好奇心、紀律、誠信與利他",
         "講授打造比財富更寶貴的人性硬核品格：1. **永無止境的好奇心 (Curiosity)**：對 AI, Web3, 綠色金融保持開放探索；2. **鋼鐵般的執行紀律 (Discipline)**：堅守 Pay Yourself First 與定期定額；3. **商業誠信與倫理 (Integrity)**：堅守 G 公司治理；4. **利他共享與社會責任 (Altruism)**：用資本支持 ESG 永續企業。",
         "115學年度「管理探索二」核心品格。",
         "強調：『品格決定高度，紀律決定財富！』"),
        (23, 3, "校友回饋與導師陪伴：建立大學 4 年理財與專案諮詢通道",
         "建置課程結束後的陪伴與諮詢通道：1. **線上諮詢**：加入「115 管理探索二校友交流群」，更新理財與實習資訊；2. **導師 Office Hours**：大學四年期間預約導師 1-on-1 諮詢租房合約、實習履歷或理財疑惑；3. **校友分享**：定期邀請學長姐分享金控與科技業經驗。",
         "115學年度「管理探索二」校友會服務。",
         "公佈導師與校友交流群之聯繫方式。"),
        (24, 3, "課程滿意度調查與未來教案優化建議",
         "邀請學生填寫全學期課程滿意度問卷：請點擊系統連結，匿名評估 18 週 30 頁純教學卡片、Canva 莫蘭迪視覺、畫布工具、3 大小時 Modal 試算與小遊戲體驗，提供寶貴建議供下學期教案優化。",
         "115學年度「管理探索二」課程問卷。",
         "給予學生 3 分鐘時間進行線上匿名滿意度問卷填寫。"),
        (25, 3, "實證數據調取：下載個人 18 週學習成績單與數位證書",
         "手把手教導學生線上下載成績單與證書：1. 匯出包含作業、同儕互評與期末發表之 18 週詳細權重成績單；2. 取得「115 學年度管理探索二」3 學分通識結業證書 PDF；3. 直接上傳至教育部大一學習歷程中央資料庫。",
         "115學年度「管理探索二」學習歷程系統。",
         "演示下載 PDF 檔並儲存於個人隨身碟之步驟。"),
        (26, 3, "大一新鮮人終身財經宣言 4 大誓言",
         "總複習終身財經宣言 4 大神聖誓言：1. **誓言先儲蓄後消費** (Pay Yourself First 20%)；2. **誓言定期定額配置全市場 ETF** (0050 / 00878)；3. **誓言堅守資安防線防範詐騙** (165, Deepfake 暗語, 冷錢包)；4. **誓言用資本支持 ESG 永續好企業** (世代正義與低碳生活)。",
         "115學年度「管理探索二」4 大神聖誓言。",
         "請學生再次重溫四大誓言，銘刻於心。"),
        (27, 3, "終身實踐試算實例：30 年資產複利累積與財富自由 SOP",
         "展示 20 歲大一新鮮人邁向 50 歲 1,500 萬 FIRE 的完整 SOP：1. **階段 1 (20-30歲)**：月存 5,000 元買 0050 ➔ 31 歲達成 100 萬第一桶金！2. **階段 2 (30-40歲)**：月存 15,000 元買 0050+00878 ➔ 41 歲達 500 萬！3. **階段 3 (40-50歲)**：複利爆發 ➔ 50 歲達 1,500 萬 F.I.R.E. 池，年領 60 萬股息被動收入！",
         "115管理探索二教案終身理財組。",
         "強調：『這不是不可及的夢想，而是只要你堅持執行 SOP 就必然發生的數學事實！』"),
        (28, 3, "揚帆起航：帶著探索勇氣，開創屬於你的富足與精彩人生！",
         "總結全學期終極寄語：「管理探索二」課程在這裡畫下圓滿句點，但你們精彩的財經與管理人生才剛剛開始。帶上 18 週獲得的知識、紀律與資安金鐘罩，勇敢邁開步伐，開創屬於你們的富足自由人生！",
         "115學年度「管理探索二」課程閉幕致詞。",
         "給予學生最溫暖的祝福與鼓勵。"),
        (29, 3, "第十八週全景知識體系圖與全學期 18 週總閉環",
         "以全景邏輯結構圖呈現 18 週宏觀總閉環：W1-W6 (理財健檢 ➔ 時間複利 ➔ 好債壞債 ➔ 損益表 ➔ 營運資金 ➔ 訂價策略) ➔ W7-W12 (三大報表 ➔ 0050組合 ➔ 不動產租買 ➔ 所得稅報稅 ➔ 生涯現金流) ➔ W13-W18 (台積電 ➔ AI 3層鏈 ➔ ESG 碳費 ➔ Web3 冷錢包 ➔ 期末專案發表 ➔ 4年藍圖與終身宣言)。",
         "115學年度「管理探索二」全學期總閉環圖。",
         "全班一同回顧整學期地圖，確認達成完全的融會貫通。"),
        (30, 3, "課程圓滿落幕：珍重再見，開啟大一新鮮人璀璨的人生新篇章！",
         "恭喜全班同學圓滿修畢 115 學年度「管理探索二」3 學分通識核心課程！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人大學 4 年財經探索藍圖與結業感言」。祝大家：大學四年學業順利、財務獨立、資安無虞、幸福富足！珍重再見！",
         "115學年度「管理探索二」課程委員會。",
         "全班熱烈掌聲，課程圓滿落幕！")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：大學 4 年財經藍圖與百萬第一桶金試算器",
         "1. 活動目標：幫助學生計算從 20 歲起月存 $5,000 元於 0050/00878 (7% 年化報酬)，達成第一個 100 萬第一桶金之具體年齡與時程。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入每月預計儲蓄金額與投資回報率，點擊精算。\n"
         "3. 診斷反思：系統算出達成 100 萬之年限與年齡，驗證「時間複利讓小資族在 31 歲前輕鬆擁有一百萬第一桶金」之結論。"),

        ("🎯 第 2 小時實務活動：18 週全景知識能力診斷與結業榮譽試算器",
         "1. 活動目標：讓學生自我評估在理財、產業、資安與簡報表達四大維度的學習成效，精算 18 週全景能力分數與結業榮譽勳章。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入四大維度得分，點擊精算。\n"
         "3. 決策學習：系統產出平均能力分數與「財經探索金牌領航員」勳章，體現 18 週獨立財務思辨與管理素養。"),

        ("🎮 第 3 小時小遊戲：18 週全景財經通識總大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大全學期主題關卡（關卡 1：Pay Yourself First 公式；關卡 2：台積電純晶圓代工模式；關卡 3：4% FIRE 規則 25 倍資產池；關卡 4：Web3 冷錢包 Not Your Keys, Not Your Coins 鐵律）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 全學期財經通識大滿貫領航員徽章 (Grand Master)」，未滿分獲頒「🥇 財經通識優秀學員」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人大學 4 年財經探索藍圖與結業感言",
         "1. 作業題目：請學生制定大學 4 年階梯式財經成長藍圖，簽署終身財經 4 大誓言並撰寫 18 週結業感言。\n"
         "2. 分析要項：(1) 制定大學 4 年 (打底 ➔ 實務 ➔ 實習 ➔ 獨立 30萬) 之階段目標與 SOP；(2) 簽署 Pay Yourself First 與 0050 定期定額誓言；(3) 規劃 30 年 1,500 萬 FIRE 退休金池行動方案；(4) 撰寫 18 週課程結業感言。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱並頒發數位結業證書。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十八週_課程教學指引_學期總結財經探索藍圖與大一新鮮人終身宣言.docx'
    doc.save(doc_path)
    print("Created 第十八週_課程教學指引_學期總結財經探索藍圖與大一新鮮人終身宣言.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
