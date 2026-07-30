import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79) # #1F4E79
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6) # #2E75B6
    EMERALD_GREEN = RGBColor(0x10, 0xB9, 0x81)
    PURPLE_ACCENT = RGBColor(0x8B, 0x5C, 0xF6)
    DARK_TEXT = RGBColor(0x26, 0x26, 0x26)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第五週 課程教學指引：ETF 狂熱、主動型 ETF 與被動投資")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第五週完整教學指引。本單元專門針對零財金背景大一新鮮人設計，旨在化解社會大眾對於「ETF 狂熱 (0050, 00940)」、「高股息定存化」與「定期定額萬靈丹」的常見迷思。透過生活化比喻（如水果籃便當、微笑曲線、豆腐牌效應），結合臺灣證券交易所 (TWSE) 與投信投顧公會 (SITCA) 權威數據，幫助學生建立獨立審視 ETF 費用率、折溢價風險與資產配置之能力。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "ETF 本質與被動投資\n(Slide 01 - 10)", "John Bogle 指數革命、Vanguard 歷史、全台 ETF 狂熱 (5.5兆)、定期定額微笑曲線與費用率打擊", "🎯 1小時活動：0050 / 00940 定期定額與複利計算盤"),
        ("第二小時\n(00:50-01:40)", "高股息陷阱與主動型 ETF 解禁\n(Slide 11 - 20)", "高殖利率 ≠ 高總報酬、收益平準金新規、市價 vs 淨值折溢價、主動型 ETF 政策解禁與豆腐牌", "🎯 2小時活動：ETF 市價淨值折溢價與套利檢測器"),
        ("第三小時\n(01:40-02:30)", "實證比對與資產配置\n(Slide 21 - 30)", "選牌 4 大金律、0050 vs 0056 10年報酬對決、70/30 核心衛星配置與美股 VOO/QQQ/VT", "🎮 3小時小遊戲：ETF 投資專家大挑戰\n📝 課堂實務作業：自選 ETF 分析報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第五週課程導論：ETF 狂熱、主動型 ETF 與被動投資", 
         "本頁為全週課程之開宗明義導論。講師應引用指數化投資之父約翰·柏格 (John Bogle) 名言：「不要在乾草堆裡找針，把整座乾草堆買下來！」作為開場。針對大一新鮮人，以「買綜合水果便當」作為比喻——單買一顆進口哈密瓜要價 500 元（類似高單價個股），但買一盒包含哈密瓜、蘋果與葡萄的綜合水果盒只需 100 元，即可享受多重水果風味並分散壞果風險。說明 ETF 的底層邏輯正是將符合特定規則的一籃子優質企業打包挂牌交易，讓散戶以極低資本參與國家與產業長期經濟成長。",
         "臺灣證券交易所 (TWSE) ETF 專區 / SITCA 投信投顧公會統計數據。",
         "問學生：『如果同學只有 3,000 元零用錢，你會選擇賭在一家小公司上，還是買進全台灣前 50 大企業的綜合籃子？理由是什麼？』"),
        (2, 1, "第五週 3 小時學習地圖與核心技能樹",
         "本頁為學生建構 3 小時的學習導航地圖。第一小時著重於 ETF 本質、Bogle 指數革命與定期定額微笑曲線；第二小時深入剖析高股息 ETF 迷思、收益平準金新規、折溢價套利與金管會主動型 ETF 政策解禁；第三小時則帶領學生進行選牌 4 大指標檢驗、0050 vs. 0056 10 年實證對決，以及 70/30 核心衛星資產配置。讓學生明確瞭解本週將習得從基礎概念到自主配置的全套終身受用技能。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生對照講義，劃出本週最感興趣的三大關鍵字（如：高股息、折溢價、主動型 ETF）。"),
        (3, 1, "什麼是 ETF？(Exchange Traded Funds: 交易所交易基金的比喻)",
         "詳細拆解 ETF (Exchange Traded Fund) 的定義與運作架構。說明 ETF 兼具「開放式基金的風險分散」與「普通股票的即時交易流動性」。以台灣最代表性的 0050 為例，投信依照「台灣 50 指數」規則，將台積電、聯發科、鴻海等前 50 大企業按市值比例買齊並打包成一股，投資人只要在證券 App 下單買進 1 股 0050，就等於同時擁有這 50 家龍頭企業的微型股權。強調其盤中即時撮合、價格透明度 100% 的獨特優勢。",
         "臺灣證券交易所 (TWSE) 證券統計月報與 ETF 交易機制規範。",
         "提問：『傳統基金每天只有一個下班後的淨值成交，而 ETF 在盤中每一秒都有市價，這對投資人的方便性有何差別？』"),
        (4, 1, "被動投資哲學：約翰·柏格 (John Bogle) 與 Vanguard 指數革命",
         "講述 1975 年約翰·柏格創立 Vanguard 並發行全球首檔標普 500 指數基金的傳奇歷史。當時華爾街主動型基金經理人嘲笑指數基金是「追求平庸」，但半個世紀的實證數據卻給了華爾街沉重一擊：超過 85% 的主動型基金在 10 年期的長期累積報酬上打不贏大盤指數！說明主動基金因為收取高昂管理費 (1.5%-2.0%)、頻繁換股產生摩擦稅費，導致「扣除費用後勝率急劇下降」；而被動投資則靠著低內扣費用與追蹤大盤，成為投資人的長期獲利勝地。",
         "S&P Dow Jones SPIVA Scorecard 報告與 Vanguard 歷史績效對比。",
         "引導思考：『為什麼年薪千萬的華爾街頂尖經理人，長期下來反而打不贏一套固定規則的自動化電腦指數？』"),
        (5, 1, "個股 vs. 共同基金 vs. ETF 的成本與交易機制比較",
         "透過清晰的矩陣對照表，帶領學生全方位比較「個股」、「傳統共同基金」與「被動型 ETF」三大工具。重點突出 ETF 在「交易場所（交易所即時成交）」、「管理費用率（0.1%~0.4% 極低內扣）」、「持股透明度（每日 100% 公開）」與「風險分散」四個維度的綜合優勢。幫助大一新鮮人建立清晰的工具選擇思維，明白為什麼現代理財學界一致推薦新鮮人以 ETF 作為第一步。",
         "金管會證期局 (FSC) 與投信投顧公會 (SITCA) 工具比較資料。",
         "請學生算一算：『如果共同基金每年收 2% 管理費，而 ETF 只收 0.2%，連續投資 30 年下來，兩者的費用差額會達到多少本金？』"),
        (6, 1, "實證數據：台灣 ETF 狂熱 (0050, 00940) 規模突破 5 兆元",
         "引述 2026 最新官方數據，揭示台灣 ETF 市場強勁暴增的現狀。全台 ETF 總資產規模已突破 5.5 兆元新台幣，受益人數超過 1,100 萬人次！將市場三大主流劃分為：1. 市值型 (如 0050 規模破 4,000 億)；2. 高股息型 (如 0056/00878/00940 總規模破 1.8 兆)；3. 債券型 (美債 ETF 規模破 2.5 兆)。分析全民瘋狂搶購高股息 ETF 的社會現象，並預告高配息背後隱藏的投資盲點。",
         "中華民國證券投資信託暨顧問商業同業公會 (SITCA) 2026 官方統計。",
         "提問：『同學們或家人身邊有買 0050 或 00940 嗎？大家買 ETF 最看重的究竟是每月的配息金額，還是未來的資產翻倍？』"),
        (7, 1, "追蹤指數機制：成分股如何選出？(市值加權 vs 股息加權)",
         "解密 ETF 的底層演算法——指數編製規則。對比兩大主流加權邏輯：1. **市值加權法 (Market-Cap Weighted)**：如 0050 追蹤台灣 50 指數，企業市值越大佔比越高，能自動讓台積電等強勢贏家權重放大，實現自動汰弱留強；2. **股息率加權法 (Dividend Weighted)**：如高股息 ETF，依歷史或預估股息率排名篩選，容易偏向成熟成熟產業，卻可能剔除不發股利的高成長科技飆股。",
         "臺灣指數公司 (TIP) 與 FTSE 富時指數編算規則書。",
         "引導討論：『如果台積電因為要投入千億研發而不發高額股息，高股息 ETF 就會把它踢掉，這對長遠的投資報酬是好事還是壞事？』"),
        (8, 1, "定期定額 (DCA) 的數學優勢：微笑曲線 (Smile Curve)",
         "詳細拆解定期定額 (Dollar-Cost Averaging, DCA) 的數學機制。說明定期定額是在固定時間投入固定金額買進，其最大勝率在於「價格高時買到較少單位數，價格大跌時自動買進大量便宜單位數」，進而拉低平均持股成本。透過「微笑曲線」圖解，展示市場經歷高點拉回再反彈時，定期定額投資人如何在谷底累積大量單位數，待行情回升時率先轉虧為盈暴賺。",
         "證券投資人保護中心 (SFIPO) 定期定額投資理財手冊。",
         "提問：『當你看到股市大跌 20% 時，一般散戶會恐慌停損，但定期定額投資人為什麼應該感到高興？』"),
        (9, 1, "總管理費用率 (Expense Ratio) 對長期複利財富的打擊",
         "揭露內扣費用率的隱形殺手本質。ETF 的總管理費用率 (Expense Ratio) 包括經理費、保管費、指數授權費與換股摩擦手續費，這些費用不會寄帳單給投資人，而是天天直接從 ETF 淨值中扣除。計算實例：以 1,000 萬元本金投資 30 年，年費用率 1.5% 的高費用 ETF 會被吃掉超過 280 萬元費用；而 0.15% 的超低費用 ETF 僅扣 35 萬元，兩者差距高達 245 萬元！",
         "投信投顧公會 (SITCA) 各檔 ETF 總內扣費用率公開揭露資訊。",
         "強調：『內扣費用是市場上唯一 100% 確定會發生的成本！選擇低費用率標的是投資人掌控回報的第一步。』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "對第一小時的教學內容進行系統整理：ETF 是籃子股票、Bogle 指數化投資戰勝主動經理人、定期定額利用微笑曲線降本，以及嚴格控管內扣費用率。預告即將進行的第 1 小時 Modal 實務體驗活動——「0050 / 00940 定期定額與時間複利試算盤」，請學生親自輸入每月 3,000~5,000 元預算，計算 20 年後的複利資產累積總額。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行定期定額複利計算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "高股息 ETF 狂熱真相：高殖利率 (Yield) 不等於高總報酬",
         "深入剖析大一新鮮人最常陷入的「高股息陷阱」。強調：**投資總報酬率 = 現金股利 + 資本利得（價差）**。宣稱 8%~10% 高殖利率的 ETF，若為了發股息而選入獲利衰退、股價連年下跌的公司，配出的股息不過是「左手換右手」的自本金發還，除息後更面臨貼息窘境。實證數據顯示，過去 10 年 0050（市值型）總報酬率顯著碾壓高股息 ETF。",
         "臺灣證券交易所 (TWSE) 歷史含息總報酬指數 (Total Return Index)。",
         "提問：『如果你拿到了 1,000 元股息，但股票本金卻虧損了 2,000 元，請問你這一趟投資到底是賺錢還是賠錢？』"),
        (12, 2, "收益平準金 (Equalization Fund) 機制與金管會新制規範",
         "解密「收益平準金」制度的建立背景與監管要求。說明當一檔熱門高股息 ETF 在配息前突然湧入大量新資金時，若無平準金機制，原本舊股東應得的股息會被新進資金嚴重稀釋。平準金是將新進資金的一小部分劃入特別帳戶，專款專用於平穩配息。同時引述金管會 2024 年發布的三大新制作法：配息率不得高於追蹤指數息率、明確揭露配息來源成分比重。",
         "金融監督管理委員會 (FSC) 證期局收益平準金資訊揭露規範。",
         "問學生：『為什麼金管會要嚴格規定投信公司不能拿收益平準金隨意衝高配息率？』"),
        (13, 2, "ETF 的兩大價格：市價 (Market Price) vs. 淨值 (NAV)",
         "清晰界定 ETF 的兩大基本價格概念：1. **市價 (Market Price)**：次級市場（交易所）買賣雙方即時撮合成交的價格，由供需熱度決定；2. **淨值 (Net Asset Value, NAV)**：初級市場投信結算 ETF 裡面所有股票資產價值後算出的每股真實價值。強調黃金鐵律：市價必須緊密貼合淨值，若出現偏差即產生折溢價風險。",
         "臺灣證券交易所 (TWSE) 即時折溢價查詢平台。",
         "舉例說明：『淨值就像便當盒裡食材的成本價（100元），市價是排隊民眾瘋搶喊出的價格（120元），你願意花 120 元買 100 元的便當嗎？』"),
        (14, 2, "折價 (Discount) 與溢價 (Premium) 的形成與套利機制",
         "詳細解析折溢價算式：`折溢價率 = (市價 - 淨值) / 淨值 * 100%`。當市價 > 淨值時稱為**溢價 (Premium)**，代表買貴了；當市價 < 淨值時稱為**折價 (Discount)**，代表買便宜了。說明初級市場授權參與者 (Authorized Participants, AP/PD) 的**申購與贖回套利機制**：當溢價過高時，PD 會在初級市場申購新單位並在次級市場賣出，迅速將市價拉回淨值，因此散戶追高溢價必然受損。",
         "臺灣證券交易所 (TWSE) 初級市場申購贖回與套利機制手冊。",
         "警示學生：『當新聞報導某檔新 ETF 溢價超過 5% 時，為什麼絕對不能在 App 上按買進？』"),
        (15, 2, "指數替換效應：成分股豆腐牌與吃豆腐現象 (Front-Running)",
         "解密被動型 ETF 在指數調整時面臨的「吃豆腐 (Front-Running)」現象。由於被動 ETF 必須完全公開透明地依照指數發布日換股，市場上的主力游資與量化基金會提前預測並買進即將被納入的新成分股（抬高股價），等到千億級 ETF 在生效日被迫高價接盤時再賣出獲利了結。這種「吃 ETF 豆腐」的交易會增加 ETF 建倉成本、產生追蹤偏離。",
         "證券自營商與量化交易法人指數調整套利研究報告。",
         "提問：『如果你知道明天下課全校學生都會去買某家排骨便當，你今天會先做什麼事？這就是吃豆腐的道理。』"),
        (16, 2, "前沿政策解禁：主動型 ETF (Active ETFs) 運作機制",
         "介紹金管會 2025-2026 年開放主動型 ETF (Active ETFs) 審查掛牌的重大政策。對比傳統被動 ETF，主動型 ETF 不再死板追蹤特定指數，而是由專業基金經理人團隊發揮選股能力與換股彈性。其最大亮點在於結合了「傳統主動基金追求超額報酬 (Alpha) 的潛力」與「ETF 在交易所即時交易、透明度高與低成本」的雙重優勢。引用美國 Cathie Wood 的 ARKK 作為國際標竿。",
         "金融監督管理委員會 (FSC) 主動型 ETF 開放法規草案與公聽會紀錄。",
         "引導思考：『當主動型 ETF 解禁後，你認為基金經理人的個人選股能力會變得更重要還是更不重要？』"),
        (17, 2, "主動型 ETF vs. 被動型 ETF 的核心差異對照",
         "透過構造精密的雙軌比較表，幫助學生掌握主動型 ETF 與被動型 ETF 在「選股決策（規則 vs 經理人）」、「核心目標（貼近指數 vs 尋求 Alpha 超額報酬）」、「管理費用率（0.1%~0.4% vs 0.5%~0.9%）」與「持股揭露頻率」上的實質差異。指導學生如何依據個人風險偏好選擇合適的 ETF 工具。",
         "臺灣證券交易所 (TWSE) 主動型與被動型 ETF 規格比較資料。",
         "提問：『想要安心獲取大盤平均成長的人該選哪種？想要賭經理人能挑出飆股的人該選哪種？』"),
        (18, 2, "主題式 ETF (Thematic ETFs) 的熱潮與警訊：AI、半導體",
         "剖析熱門主題式 ETF（如 AI、半導體、電動車、低軌衛星）的優缺點。優點：精準押注高成長熱門產業，短線具備強大爆發力；缺點與警訊：持股集中度過高（前 5 大持股佔比常達 60%~70%），容易面臨單一產業景氣循環谷底的劇烈拉回，且當主題熱度退燒後常面臨成交量萎縮下市風險。",
         "晨星 (Morningstar) 全球主題式 ETF 存續期與績效實證研究。",
         "提醒：『主題式 ETF 適合當作短線的衛星配置，絕不適合拿來當作 100% 的退休金底座！』"),
        (19, 2, "債券型 ETF (Bond ETFs) 的資產配置角色與稅務優勢",
         "介紹台灣市場極具特色的「美債 ETF狂熱 (如 00679B, 00720B)」。說明債券 ETF 在投資組合中擔任「鎖定高利率與防守避險」的角色。特別強調台灣對債券 ETF 的三大稅務優惠：1. 停徵證券交易稅（省下 0.1%）；2. 配息屬於「海外所得」，適用基本所得額 750 萬元免稅額度，免繳二代健保補充保費與個人綜合所得稅。",
         "財政部賦稅署債券型 ETF 停徵證券交易稅與所得稅相關規定。",
         "問學生：『對於高收入大戶來說，為什麼買美債 ETF 比直接領台股股息更容易省下鉅額所得稅？』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：高股息不等於高總報酬、注意收益平準金新制、切勿追高 >1% 溢價之 ETF、瞭解主動型 ETF 政策解禁，以及主題型與債券型 ETF 的配置定位。預告第 2 小時 Modal 實務活動——「ETF 市價淨值折溢價與套利檢測器」，引導學生輸入盤中市價與淨值，計算折溢價率並診斷風險。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行折溢價率診斷演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "大一 ETF 選牌 4 大指標：規模、流動性、費用率與追蹤偏離",
         "歸納大一新鮮人挑選優質 ETF 的四大黃金檢驗指標：1. **資產規模 (AUM)**：優先選擇規模 > 100 億元的標的，遠離規模小於 2 億元觸及清算下市門檻的危險標的；2. **每日成交量 (Liquidity)**：日均量 > 3,000 張，確保買賣不卡關；3. **超低費用率**：總內扣費用 < 0.4%/年；4. **低追蹤偏離度**：淨值貼合指數走勢。",
         "臺灣證券交易所 (TWSE) ETF 清算下市規範與選牌指南。",
         "提問：『如果有一檔新發行的 ETF 宣稱主題很酷，但每天成交量只有 50 張，你敢買嗎？為什麼？』"),
        (22, 3, "追蹤偏離度 (Tracking Error) 與追蹤差距 (Difference)",
         "深入探討「追蹤品質 (Tracking Quality)」。說明為什麼 ETF 的實際淨值報酬率不會 100% 完全等於指數報酬率？兩大主要偏差來源：1. **每日內扣費用**（經理費與保管費扣除）；2. **換股交易摩擦稅費**。計算公式：`追蹤差距 = ETF淨值報酬率 - 追蹤指數報酬率`。優秀的投信公司能將追蹤差距控制在極小範圍。",
         "彭博 (Bloomberg) & Morningstar 追蹤品質評估標準。",
         "引導思考：『當兩檔 ETF 追蹤同一個指數時，追蹤差距越小的代表投信公司的操作技術越好還是越差？』"),
        (23, 3, "實證比對：市值型 0050 vs. 高股息 0056 10 年累積報酬對決",
         "拿出 2016 至 2026 年長達 10 年的真實歷史大數據對決：市值型 0050 10 年累積含息總報酬率高達 **+260% ~ +310%**；而高股息 0056 10 年累積總報酬率為 **+140% ~ +170%**。兩者總報酬相差超過 1.3 倍！解析核心原因：0050 重倉台積電（佔比約 50%），完全享受了全球半導體 AI 暴利；而 0056 則因篩選高股息剔除了高成長股。結論：年輕新鮮人應以市值型為主，退休族才以高股息為主。",
         "臺灣證券交易所 (TWSE) 歷史含息總報酬率指數實證數據。",
         "讓學生進行小組討論：『20 歲的大一學生與 65 歲準備退休的爺爺，兩人的 ETF 配置應該一模一樣嗎？』"),
        (24, 3, "實證數據調取：臺灣證券交易所 (TWSE) ETF 專區查閱",
         "手把手教導大一學生如何直接調取權威官方數據，拒絕網路流言：1. **TWSE ETF 專區 (twse.com.tw/ETF)**：查詢盤中即時折溢價、每日 AUM 規模與配息公告；2. **SITCA 投信投顧公會**：查詢每檔 ETF 每月的「完整內扣總費用率」（含經理費、保管費與交易手續費）；3. **公開資訊觀測站 (MOPS)**：查閱配息組成來源。",
         "臺灣證券交易所 (TWSE) 暨投信投顧公會公開數據服務平台。",
         "演示並要求學生記錄下兩家權威官方數據網站的網址。"),
        (25, 3, "核心與衛星策略 (Core-Satellite Strategy) 實務應用",
         "傳授華爾街最經典的「核心與衛星 (Core-Satellite) 資產配置架構」。針對大一新鮮人提出 **70/30 最佳黃金配比**：70% 配置為 **核心資產 (Core)**，選擇廣基市值型 ETF（如 0050 或 VOO），獲取整體經濟成長總報酬；30% 配置為 **衛星資產 (Satellite)**，選擇高股息 (00878)、主題型 (AI) 或美債 ETF，創造穩定現金流或追求超額爆發力。",
         "Vanguard & BlackRock 經典資產配置理論。",
         "請學生隨堂演練：『如果你每月有 10,000 元可投入，請按照 70/30 策略分配你的核心與衛星 ETF 扣款金額。』"),
        (26, 3, "美股 ETF 投資入門：VOO、QQQ、VT 全球資產配置",
         "帶領學生將眼界放眼全球資本市場，介紹美股三大天王級 ETF：1. **VOO (Vanguard S&P 500)**：追蹤標普 500，內扣費用僅 **0.03%/年**，一次擁有美股 500 大巨擘（蘋果、微軟、輝達）；2. **QQQ (Invesco Nasdaq 100)**：聚焦全球科技創新與 AI 引擎；3. **VT (Vanguard Total World)**：內扣 **0.07%/年**，一口氣打包全球 9,000+ 家企業，實現真正的全地球資產配置。",
         "US SEC Filings & Vanguard / Invesco 官方說明書。",
         "問學生：『台灣市場只占全球股市市值的 2%，只買台股算不算是把雞蛋放在同一個籃子裡？』"),
        (27, 3, "避免過度分散：持有多檔重覆成分股 ETF 的投資盲點",
         "點出散戶常見的「假分散」投資盲點。許多投資人以為同時買進 0050、0052、006208、00940 是分散風險，但透過 MOPS 成分股交叉比對發現：台積電與聯發科在這些 ETF 裡的重覆持股比例高達 70%！這並沒有達成真正的風險分散，反而多付了數筆經理費與手續費。強調：分散投資看的是「底層成分股」，而不是「買了幾張不同代號的貼紙」。",
         "公開資訊觀測站 (MOPS) 成分股交叉比對分析。",
         "提醒學生：『買 ETF 前一定要看前 10 大成分股明細，避免重複買進同一批公司。』"),
        (28, 3, "大一新鮮人 ETF 理性投資 4 大金律",
         "總結大一新鮮人終身受用的 ETF 理性投資四大金律：1. **優先選擇廣基市值型**（以 0050/VOO 為核心）；2. **嚴格控管內扣費用率**（選擇費用 < 0.4%）；3. **堅持定期定額紀律**（利用微笑曲線，不隨市場恐慌停扣）；4. **不盲追過度溢價 ETF**（買進前先查即時折溢價，溢價 > 1% 絕不上當）。",
         "金管會證期局 (FSC) 投資人教育關懷原則。",
         "請學生齊聲朗讀四金律，深化理性投資信念。"),
        (29, 3, "第五週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第五週 30 頁純教學卡片進行整體串聯：Hour 1 (ETF 本質 ➔ Bogle 革命 ➔ 微笑曲線 ➔ 費用率) ➔ Hour 2 (殖利率 ≠ 總報酬 ➔ 收益平準金 ➔ 折溢價套利 ➔ 主動型 ETF) ➔ Hour 3 (選牌 4 大指標 ➔ 0050 vs 0056 對決 ➔ 70/30 核心衛星 ➔ 美股三天王 ➔ 4 大金律)。幫助學生形成系統化的財經素養框架。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第六週預告 (數位金融與高利活存)",
         "恭喜學生完成第五週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「自選 ETF 成分股、費用率與 Core-Satellite 70/30 配置報告」。預告第六週課程主題：「數位金融、高利活存與銀行業變革」，下週將帶大家探討純網銀高利活存優惠條件、信用卡回饋計算與商業銀行利差模式。",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第六週數位金融主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：0050 / 00940 定期定額 (DCA) 與時間複利試算盤",
         "1. 活動目標：讓學生親自體驗定期定額微笑曲線與時間複利的真實威力。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入每月扣款金額（例如 3,000 或 5,000 元）、預期年化報酬率（例如 7.0%）與投資年數（例如 20 年），點擊試算。\n"
         "3. 討論引導：引導學生比較「總投入本金」與「期末資產總額」，體驗時間複利創造的淨利潤，並思考為什麼越早開始定期定額越有優勢。"),

        ("🎯 第 2 小時實務活動：ETF 市價淨值折溢價與套利風險檢測器",
         "1. 活動目標：幫助學生掌握市價與淨值的差異，學會在下單前檢測折溢價率以防追高。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入即時盤中市價與官方淨值，系統自動依據公式 `(市價 - 淨值) / 淨值 * 100%` 計算折溢價率。\n"
         "3. 診斷反思：當折溢價 > +1.0% 時警告嚴重溢價風險；當折溢價 < -1.0% 時提示折價區間，教育學生絕不盲追高溢價 ETF。"),

        ("🎮 第 3 小時小遊戲：ETF 狂熱與被動投資大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：John Bogle 指數革命；關卡 2：高股息殖利率 vs 總報酬；關卡 3：折溢價套利防禦；關卡 4：核心衛星資產配置）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 ETF 投資專家徽章 (ETF Master)」，未滿分獲頒「🥉 ETF 初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：自選 ETF 成分股與費用率分析報告",
         "1. 作業題目：請學生自行挑選 1 檔感興趣的 ETF（如 0050, 00878, 00940 或 VOO），線上填寫報告。\n"
         "2. 分析要項：(1) 說明該 ETF 的追蹤指數與前 5 大成分股；(2) 查詢公會 (SITCA) 數據，記錄其每年總內扣費用率；(3) 擬定個人 Core-Satellite (70/30) 配置方案。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第五週_課程教學指引_ETF狂熱主動型ETF與被動投資.docx'
    doc.save(doc_path)
    print("Created 第五週_課程教學指引_ETF狂熱主動型ETF與被動投資.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
