import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_week1_guide():
    doc = docx.Document()

    # Set 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Document Header
    title_p = doc.add_paragraph()
    title_run = title_p.add_run('115學年度上學期「管理探索二」第一週教學指引\n【資訊爆炸時代的財經素養】（30頁純教學+3小時獨立活動+課堂作業）')
    title_run.font.name = 'Microsoft JhengHei'
    title_run.font.size = Pt(17)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Lesson Meta Info
    meta = doc.add_paragraph()
    m_run = meta.add_run('一、 課程架構與網頁互動工具簡介\n')
    m_run.bold = True
    m_run.font.size = Pt(13)
    m_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    meta.add_run('• 課程主題：第一週 課程導論：資訊爆炸時代的財經素養\n')
    meta.add_run('• 預計授課時間：3 小時（共 180 分鐘）\n')
    meta.add_run('• 教學內容份量：完整 30 頁純教學卡片模組（活動與作業不計入 30 頁，確保教學充實度）。\n')
    meta.add_run('• 網頁互動工具功能：\n')
    meta.add_run('   1. 🖥️ 全螢幕簡報模式 (Fullscreen Mode)\n')
    meta.add_run('   2. 🖊️ 手繪畫布書寫筆刷 (Freehand Pen Tool)\n')
    meta.add_run('   3. 🖍️ 螢光筆重點標記功能 (Fluorescent Highlighter Tool)\n')
    meta.add_run('   4. 🧹 畫布即時清除與多色切換 (Clear & Color Selector)\n')
    meta.add_run('   5. 🌐 中英雙語一鍵即時切換 (Traditional Chinese / English Switcher)\n')

    # Section 2: 3-Hour Time Allocation & Activities Table
    h2 = doc.add_heading('二、 3小時教學進度、獨立活動與課堂作業規劃表', level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    schedule_data = [
        ('第一小時\n(00:00-00:50)', '理論與歷史背景\n(Slide 01 - 10)', '🎯 【第1小時獨立活動】\n「資訊飲食與 SNR 計算盤」\n學生輸入每日看財報與標題黨分鐘數，計算資訊純度分數。', '講授 35 分\n活動 15 分'),
        ('第二小時\n(00:50-01:40)', '偏誤與防禦工具\n(Slide 11 - 20)', '🎯 【第2小時獨立活動】\n「新聞標題框架重組大考驗」\n將聳動標題黨新聞改寫為學術級中立 Signal 報導。', '講授 35 分\n活動 15 分'),
        ('第三小時\n(01:40-02:30)', '實證與工具應用\n(Slide 21 - 30)', '🎮 【第3小時重磅小遊戲】\n「財經新聞真假與雜訊偵測大挑戰」 (4大關卡競賽)\n📝 【第3小時課堂作業】\n「4D Filter 財經新聞檢驗報告表單」現場填寫與提交。', '講授 20 分\n遊戲 15 分\n作業 15 分')
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    headers = ['授課時段', '純教學模組內容 (30頁)', '專屬獨立活動與課堂作業模組', '時間分配']
    widths = [Inches(1.2), Inches(1.8), Inches(3.2), Inches(1.0)]

    for i, t in enumerate(headers):
        hdr[i].text = t
        hdr[i].width = widths[i]
        set_cell_background(hdr[i], '1F4E79')
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, data in enumerate(schedule_data):
        cells = table.add_row().cells
        bg = 'F2F4F8' if row_idx % 2 == 1 else 'FFFFFF'
        for i, text in enumerate(data):
            cells[i].text = text
            cells[i].width = widths[i]
            set_cell_background(cells[i], bg)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Microsoft JhengHei'
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # Section 3: Detailed Activity & Homework Guide
    h3 = doc.add_heading('三、 3大小時活動與課堂作業執行指南', level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph('【第 1 小時活動：資訊飲食與 SNR 計算盤】\n'
                      '• 操作方式：點擊網頁上方「🎯 1小時活動: SNR計算」按鈕。\n'
                      '• 核心邏輯：學生輸入每日閱讀權威數據與標題黨短影音的時間，系統自動計算 SNR 分數並給出健康診斷。')

    doc.add_paragraph('【第 2 小時活動：新聞標題框架重組大考驗】\n'
                      '• 操作方式：點擊網頁上方「🎯 2小時活動: 標題重組」按鈕。\n'
                      '• 核心邏輯：給出聳動標題（如「台積電海外建廠驚爆巨虧？！」），學生練習運用 4D Filter 重組改寫為客觀中立報導。')

    doc.add_paragraph('【第 3 小時小遊戲：財經新聞真假與雜訊偵測大挑戰】\n'
                      '• 操作方式：點擊「🎮 3小時小遊戲: 雜訊偵測」按鈕。\n'
                      '• 關卡包含：房貸限額崩盤疑雲、高股息 ETF 12% 殖利率陷阱、AI 假財報事件、半導體 CAPEX 建廠真實數據比對。')

    doc.add_paragraph('【第 3 小時課堂作業：4D Filter 檢驗報告提交】\n'
                      '• 操作方式：點擊「📝 本週課堂作業」按鈕，現場填寫 Data, Disclosure, Divergence, Duration 分析表單並提交。')

    # Section 4: 30 Pure Slide Summary
    h4 = doc.add_heading('四、 30 頁純教學卡片綱要表', level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for i in range(1, 31):
        doc.add_paragraph(f'• Slide {i:02d}：對應 30 頁純教學卡片第 {i} 頁。', style='List Bullet')

    output_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第一週_課程教學指引_資訊爆炸時代的財經素養.docx'
    doc.save(output_path)
    print(f'Saved updated week 1 guide docx to: {output_path}')

if __name__ == '__main__':
    create_week1_guide()
