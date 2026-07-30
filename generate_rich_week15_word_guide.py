import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十五週 課程教學指引：綠色金融、ESG 永續投資與碳定價經濟學")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十五週完整教學指引。本單元帶領大一新鮮人探索全球永續發展潮流中的「綠色金融與碳經濟學」。課程深入剖析 2050 淨零排放 (Net Zero) 目標與 1.5°C 氣候協定；拆解 ESG 三大柱 (環境 E / 社會 S / 公司治理 G) 評級架構；解密碳定價算術公式（碳費 = 總碳排量 × 指定費率 300元/噸）、歐盟 CBAM 碳邊境調整關稅對台灣出口製造業之衝擊；教導企業漂綠 (Greenwashing) 辨識與 GRI / SASB 國際報告標準；剖析 ESG 篩選 ETF (00878 / 00692)、台灣碳權交易所 (TCX) 運作、範疇 1-3 碳盤查 SOP、赤道原則銀行綠色融資，並手把手建立大一專屬 0050+00878 永續投資組合與綠色職涯展望，輔以環境部氣候變遷署與 TCX 實證數據，引領學生發揮綠色金融社會影響力。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "永續浪潮、ESG 與碳定價經濟學\n(Slide 01 - 10)", "2050淨零目標、ESG三大柱、碳費vs碳稅算術、歐盟CBAM碳關稅、漂綠辨識與環境部公告費率", "🎯 1小時活動：企業碳費成本與 CBAM 關稅衝擊試算器"),
        ("第二小時\n(00:50-01:40)", "ESG ETF、碳權交易與綠色金融\n(Slide 11 - 20)", "00878/00692 ESG ETF篩選、台灣碳權交易所 (TCX)、範疇1-3碳盤查SOP、綠色溢價與赤道原則融資", "🎯 2小時活動：ESG 篩選 ETF 與綠色溢價評估試算器"),
        ("第三小時\n(01:40-02:30)", "永續理財組合與綠色職涯\n(Slide 21 - 30)", "大一ESG資產配置、綠色職涯 (永續管理師)、減碳生活、防假綠色基金詐騙、碳交所查閱與4大金律", "🎮 3小時小遊戲：ESG 永續理財達人大挑戰 (4大關卡)\n📝 課堂實務作業：個人 carbon 體檢與 ESG 永續投資計畫")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十五週課程導論：綠色金融、ESG 永續投資與碳定價經濟學", 
         "本頁為第十五週課程導論。講師可以引用英格蘭銀行前總裁卡尼名言開場：「氣候變遷是終極的市場失靈，但永續金融將是本世紀最大的財富轉型契機！」針對大一新鮮人面臨的全球「2050 淨零排放 (Net Zero)」浪潮切入。說明永續金融已成為全球資本市場的主流語言。本單元將帶領大家拆解 ESG 三大柱、碳費費率算術、歐盟 CBAM 碳關稅、ESG ETF 與碳盤查 SOP。",
         "環境部氣候變遷署 / 金管會綠色金融行動方案 / 台灣碳權交易所。",
         "問學生：『同學們平時購買飲料或衣服時，會注意產品包裝上是否有標示「碳足跡」或「環保標章」嗎？』"),
        (2, 1, "第十五週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密 2050 淨零目標、ESG 三大柱 (E/S/G)、碳費 vs 碳稅算術、歐盟 CBAM 碳關稅、企業漂綠辨識與環境部公告費率；第二小時聚焦於 00878 / 00692 ESG ETF 篩選、台灣碳權交易所 (TCX)、範疇 1-3 碳盤查 SOP、綠色溢價與赤道原則；第三小時傳授大一專屬 0050+00878 ESG 資產配置、綠色職涯、減碳生活、防範假綠色基金詐騙與 4 大金律。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上記下綠色金融與 ESG 的關鍵定義，準備建立專屬的永續投資組合。"),
        (3, 1, "永續浪潮來襲：氣候變遷經濟學與 2050 淨零排放 (Net Zero)",
         "講授 2050 淨零排放 (Net Zero) 的科學目標：透過減碳、再生能源與碳捕捉，使人為溫室氣體排放量與移除量達成抵銷相等的「淨零」狀態。《巴黎氣候協定》目標將全球平均升溫控制在 1.5°C 以內。商業效應：高碳排企業將面臨鉅額碳費用負擔，低碳綠色企業獲得全球資本追捧！",
         "UNFCCC 聯合國氣候變遷公約。",
         "提問：『為什麼全球各大企業紛紛承諾在 2050 年前達成淨零排放？如果不轉型會有什麼後果？（答案：面臨碳關稅、碳費與客戶抽單風險）』"),
        (4, 1, "ESG 三大柱拆解：環境 (E)、社會 (S) 與公司治理 (G)",
         "剖析評估企業永續經營非財務績效的 ESG 三大柱：1. **E (Environmental 環境)**：溫室氣體減量、水資源管理、再生能源 RE100（如台積電承諾 2040 全面使用綠電）；2. **S (Social 社會)**：員工多元包容 (DEI)、職場工安、社區關懷；3. **G (Governance 公司治理)**：董事會獨立性、反貪腐、股東權益保障。",
         "MSCI ESG Rating 評級架構。",
         "用生活比喻：『E 是對地球好，S 是對員工與社區好，G 是對股東與營運誠信好！』"),
        (5, 1, "什麼是碳定價 (Carbon Pricing)？碳費 (Fee) vs. 碳稅 (Tax)",
         "講授碳定價經濟學原理：將排碳的外部污染成本內部化。公式：`企業年度碳費負擔 = 溫室氣體總排放量 (公噸) * 指定碳費費率 (元/公噸)`。區分兩大制度：1. **碳費 (Carbon Fee - 台灣採用)**：環境部徵收，專款專用於國家溫室氣體減量基金；2. **碳稅 (Carbon Tax - 歐洲採用)**：財政部徵收納入國庫通用。",
         "中華民國《氣候變遷因應法》第 28 條規定。",
         "算一算：『如果一家鋼鐵廠一年排放 100 萬噸 CO2，碳費每噸 300 元，一年要繳多少碳費？（答案：高達 3 億元台幣！）』"),
        (6, 1, "歐盟碳邊境調整機制 (CBAM) 與台灣出口製造業衝擊",
         "解密歐盟 CBAM 碳關稅震撼彈：進口至歐盟的高碳排產品（鋼鐵、水泥、鋁材、化肥、電力），進口商必須購買 CBAM 憑證彌補碳價差。衝擊：台灣鋼鐵、金屬扣件 (螺絲螺帽) 為出口歐盟重鎮，若台灣未先行在國內繳納碳費，產品將被歐盟補徵高額碳關稅！",
         "歐盟委員會 CBAM 官方法規白皮書。",
         "強調：『CBAM 逼迫全球所有外銷歐盟的製造業，都必須展開真槍實彈的減碳轉型！』"),
        (7, 1, "企業漂綠 (Greenwashing) 辨識指南與 SASB / GRI 標準",
         "教導辨識假綠色的「企業漂綠 (Greenwashing)」行為：漂綠是指花大錢拍廣告宣傳種植 1,000 棵樹，卻隱瞞工廠每年排放 100 萬噸廢水與碳排！防範心法：查驗企業永續報告書是否通過 **GRI (全球報告倡導組織)** 或 **SASB (永續會計準則委員會)** 第三方獨立查證。",
         "金管會防範企業漂綠行為監管指引。",
         "提醒：『看企業 ESG 不能只看公關廣告，一定要看有沒有 GRI/SASB 的第三方會計師認證報告！』"),
        (8, 1, "實證數據：環境部與金管會台灣碳費徵收費率公告",
         "引述環境部 2026 最新碳費費率公告數據：1. **一般碳費基準費率**：每公噸溫室氣體排放徵收 **300 元台幣**；2. **優惠費率 A/B (自主減量計畫)**：企業若提出符合科學減量目標 (SBTi) 之自主減量計畫，費率大幅降低至 **100 元 ~ 50 元/公噸**，強烈鼓勵企業積極設備轉型。",
         "環境部氣候變遷署 2026 碳費費率公告。",
         "問學生：『為什麼環境部給予提出自主減量計畫的企業高達 50-100 元的優惠費率？（鼓勵設備升級實質減碳）』"),
        (9, 1, "綠色債券 (Green Bonds) 與永續發展債券市場運作機制",
         "講授綠色債券 (Green Bonds) 機制：企業或政府發行固定收益債券，募集資金**限制 100% 用於綠能、提升能效、污染防治等綠色專案**（如台積電發行綠債興建低碳晶圓廠）。介紹櫃買中心 (TPEx) 永續發展債券專區，設有嚴格審核與追蹤機制。",
         "證券櫃檯買賣中心 (TPEx) 永續發展債券專區。",
         "提問：『發行綠色債券與一般債券最大差別是什麼？（答案：資金用途受法律嚴格約束，只能用於綠色專案）』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：2050 淨零排放目標；ESG 三大柱 (E/S/G)；碳費算術（排放量 × 300元/噸）；歐盟 CBAM 碳關稅衝擊；防範企業漂綠。預告第 1 小時 Modal 實務活動——「企業碳費成本與 CBAM 關稅衝擊試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行企業碳費與 CBAM 關稅衝擊試算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "ESG 篩選 ETF 解析：高股息 + ESG 篩選 (00878 / 00692) 原理",
         "解析小資族熱門 ESG ETF 運作原理：1. **國泰永續高股息 (00878)**：追蹤 MSCI 台灣 ESG 永續高股息指數，排除 ESG 評級低劣者，結合永續與高股息；2. **富邦公司治理 (00692)**：篩選臺灣公司治理評鑑前 20% 優良企業，涵蓋台積電等優質成分股。",
         "投信投顧公會 ESG ETF 專區數據。",
         "分析：『00878 為什麼受到小資族歡迎？因為它同時滿足了發放高股息與 ESG 永續篩選兩大優勢！』"),
        (12, 2, "碳權交易 (Carbon Trading) 與台灣碳權交易所 (TCX) 機制",
         "解密碳權交易 (Carbon Credit Trading) 運作：1 單位碳權代表 **1 公噸 CO2e 的排放額度**。減碳成效優異者可將多餘碳權賣給超額排碳者。介紹台灣碳權交易所 (TCX, 總部設於高雄)，提供國內自主減量碳權與國外高品質碳權買賣平台。",
         "台灣碳權交易所 (TCX) 交易規則說明。",
         "提問：『碳權可以買賣，這如何激勵企業主動砸錢投入節能減碳？（答案：減碳不僅能省碳費，還能賣碳權賺錢！）』"),
        (13, 2, "企業碳盤查 (Carbon Footprint) SOP：範疇一、二與範疇三",
         "講授 ISO 14064-1 企業溫室氣體碳盤查三大範疇：1. **範疇一 (Scope 1) 直接排放**：工廠煙囪、公務車汽油燃燒；2. **範疇二 (Scope 2) 間接外購**：外購台電電力、外購蒸氣；3. **範疇三 (Scope 3) 價值鏈上下游**：供應鏈原料製造、員工差旅、產品使用與廢棄。",
         "ISO 14064-1 溫室氣體盤查國際標準。",
         "強調：『蘋果要求供應鏈廠商做的碳盤查，包含了整條價值鏈的範疇三排放！』"),
        (14, 2, "綠色溢價 (Green Premium) 與產品碳足跡對消費者的影響",
         "講授綠色溢價 (Green Premium) 模型：買零碳/低碳產品比起傳統高碳產品額外高出的成本。公式：`綠色溢價 = 零碳產品售價 - 高碳傳統產品售價`。說明隨著技術規模化生產，綠色溢價將逐漸歸零，使低碳產品成為市場主流。",
         "比爾·蓋茲《如何避免氣候災難》綠色溢價模型。",
         "請學生思考：『你願意為了買一杯低碳製造的咖啡，多支付 5 元的綠色溢價嗎？』"),
        (15, 2, "綠色金融與銀行授信：赤道原則 (Equator Principles) 融資",
         "講授銀行業的綠色授信防線——赤道原則 (Equator Principles)：全球大型商業銀行審核大型專案融資 (> 1,000 萬美元) 時，必須評估環境與社會衝擊。台灣主要金控（富邦、國泰、玉山）已簽署赤道原則，停止向新採煤電廠放款！",
         "赤道原則協會 (Equator Principles Association) 指南。",
         "總結：『高污染企業未來在銀行將貸不到任何資金，這就是綠色金融的力量！』"),
        (16, 2, "永續投資策略：負向排除 vs. 正向篩選 vs. 影響力投資",
         "解析 3 種主流 ESG 投資選股戰術：1. **負向排除 (Exclusion)**：剔除菸草、軍火、煤炭高污染企業；2. **正向最佳選股 (Best-in-Class)**：同產業中優先選擇 ESG 評分最高者 (00878/00692)；3. **影響力投資 (Impact)**：資金直接投入能產生可衡量綠色效益之專案。",
         "全球永續投資聯盟 (GSIA) 投資策略分類。",
         "提問：『00878 ETF 屬於哪一種 ESG 投資策略？（答案：正向最佳選股 Best-in-Class）』"),
        (17, 2, "ESG 投資風險：漂綠罰款、碳費侵蝕利潤與 RE100 綠電瓶頸",
         "警告 ESG 投資的隱形轉型風險：1. **漂綠監管重罰風險**：美 SEC 與歐盟加強監管，誇大 ESG 將面臨鉅額罰款與撤資；2. **碳費侵蝕利潤**：高碳排製造業若轉型失敗，年度碳費直接侵蝕獲利；3. **RE100 綠電瓶頸**：台灣綠電供不應求，買不到綠電面臨客戶抽單風險。",
         "金管會與綠電交易市場監管警訊。",
         "提醒：『投資 ESG 概念股不能只看口號，要確認企業是否有足夠的綠電與減碳實力。』"),
        (18, 2, "全球綠能巨頭與台灣綠能供應鏈：台積電 RE100 與風光發電",
         "講授台積電 RE100 綠電承諾與台灣綠能供應鏈：台積電承諾 2040 年達成 100% 全面使用再生能源 (RE100)，並包下全台絕大部分離岸風電廠之發電量！帶動離岸風電水下基礎 (世紀鋼)、綠能能營 (森崴能源) 與儲能 (台達電) 供應鏈爆發。",
         "RE100 氣候組織官方報告與台積電永續報告書。",
         "讚歎：『台積電強勁的綠電需求，成為推動全台灣離岸风電與綠能產業發展的最大買家！』"),
        (19, 2, "台灣科技業綠色轉型：低碳製造與綠色供應鏈競爭力",
         "講授「減碳力就是接單力」的科技業鐵律：Apple 與 Microsoft 強制要求供應鏈於 2030 年達成碳中和。台灣電子廠升級綠色製造技術（高能效機器、廢水回收、智慧電網）。產品碳足跡越低的廠商，越能穩拿蘋果與微軟長期的獨家大單！",
         "Apple 供應鏈清潔能源計畫年報。",
         "強調：『未來科技代工爭奪訂單，不只比誰價格便宜，更比誰的產品碳足跡更低！』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：00878 / 00692 ESG ETF；碳權交易 (TCX)；範疇 1-3 碳盤查 SOP；赤道原則銀行授信；台積電 RE100 綠電買家。預告第 2 小時 Modal 實務活動——「ESG 篩選 ETF 與綠色溢價評估試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行 ESG ETF 與綠色溢價效益試算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "打造大一新鮮人永續理財組合：結合 0050/00878 與 ESG 思維",
         "手把手教學大一專屬永續理財資產配置：**核心組合：0050 (70%) + 00878 (30%)！** 70% 資金配置全市場台積電龍頭成長；30% 資金配置 ESG 永續高股息。實現兼顧資產成長與運用資本支持低碳綠色企業的雙贏目標！",
         "理財規劃師 (CFP) ESG 個人資產配置模型。",
         "鼓勵學生：『讓你的每一塊錢理財資金，都為地球的永續發展發揮正向影響力！』"),
        (22, 3, "大一新鮮人綠色職涯展望：永續管理師、碳盤查員與 ESG 分析師",
         "引導大一學生探索淨零轉型帶爆的新興綠金職缺：1. **永續管理師 (CSO)**：撰寫企業永續報告書、規劃 RE100；2. **碳盤查查驗員 (ISO 14064)**：輔導企業進行範疇 1-3 碳盤查；3. **ESG 永續分析師**：於投信、銀行分析企業 ESG 風險。綠色人才需求暴增！",
         "104 人力銀行綠色人才薪資白皮書。",
         "提示：『不論你主修企管、財金還是工管，考取 ISO 14064 碳盤查證照將大幅提升職場身價！』"),
        (23, 3, "綠色生活與消費減碳：減塑、循環經濟與個人碳足跡 App",
         "教導日常生活減碳與理財並行實務：1. **自備環保杯優惠**：每杯折抵 5 元（一年省下 1,800 元 + 減少 360 個塑膠杯）；2. **個人碳足跡追蹤 App**：記錄搭捷運減少的碳排；3. **循環經濟二手交易**：教科書與衣物循環利用，減少浪費。",
         "環境部全民綠生活推廣專區。",
         "算一算：『大一學生每天自備環保杯買飲料，四年下來能省下多少錢？（答案：超過 7,000 元！）』"),
        (24, 3, "綠色投資資安防詐：防範假冒綠色債券與偽 ESG 基金詐騙",
         "警告打著「愛地球」旗號的高收益詐騙陷阱：詐騙集團宣稱「投資亞馬遜雨林碳權，保證年化報酬率 20%」，99% 是詐騙資金盤！防範心法：買綠色債券或 ESG 基金，**認明金管會「ESG 基金專區」核准之正規合法基金！**",
         "內政部警政署 165 反詐騙專線警訊。",
         "警告：『凡是宣稱保證高報酬的私下碳權投資，100% 是詐騙！』"),
        (25, 3, "實證數據調取：台灣碳權交易所 (tcx.com.tw) 查閱",
         "手把手教導學生登入台灣碳權交易所官網：1. 登入 `tcx.com.tw` 查閱最新國際碳權成報價與交易數量；2. 查閱國內企業自主減量專案碳權核發進度；3. 培養將課堂碳定價理論與真實碳交易數據互相驗證的實證能力。",
         "台灣碳權交易所 (tcx.com.tw)。",
         "演示進入 TCX 官網查閱國際碳權交易走勢圖表的步驟。"),
        (26, 3, "大一新鮮人 ESG 永續理財 4 大金律",
         "總結 ESG 永續理財 4 大金律：1. **優先配置 ESG 篩選 ETF**（選擇 00878 / 00692 用資本支持減碳好企業）；2. **嚴防企業漂綠行為**（查驗 GRI/SASB 獨立第三方認證）；3. **理解碳費對企業獲利衝擊**（追蹤高碳排企業轉型能力）；4. **實踐日常低碳省錢生活**（自備環保杯與搭乘大眾運輸）。",
         "金管會與永續投資 (GSIA) 保護原則。",
         "請學生齊聲朗讀四金律，建立永續理財信念。"),
        (27, 3, "永續投資試算實例：定期定額 ESG ETF 兼顧收益與永續",
         "展示大一打工族每月 3,000 元定期定額買 00878 SOP：打工月薪 15,000 元，按 20% 比例提撥 3,000 元定期定額買 00878 (年化股息率 6.5%)。4 年累積成果：`3,000 元 * 48 個月 = 14.4 萬本金 + 複利 3 萬元 = 17.4 萬元永續積蓄！`",
         "115管理探索二教案永續理財組。",
         "強調：『大學四年定期定額買 ESG ETF，畢業時你將握有 17.4 萬元的綠色永續第一桶金！』"),
        (28, 3, "永續發展與世代正義總整合",
         "總結綠色金融的終極價值：理財不只是賺錢，更是為了實現**「世代正義 (Intergenerational Equity)」**。大一新鮮人透過聰明消費與責任投資，用選票與錢包支持低碳企業，引導資本流向真正守護地球的綠色永續產業！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『成為具備 ESG 永續視野的現代財經公民，共同守護我們的地球！』"),
        (29, 3, "第十五週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十五週 30 頁純教學卡片進行整體串聯：Hour 1 (2050淨零目標 ➔ ESG三大柱 ➔ 碳費公式 [碳排量×300元/噸] ➔ 歐盟CBAM碳關稅 ➔ 漂綠辨識 [GRI/SASB] ➔ 環境部費率公告 ➔ 綠色債券) ➔ Hour 2 (00878/00692 ESG ETF ➔ 碳權交易 [TCX] ➔ 範疇1-3碳盤查SOP ➔ 綠色溢價算術 ➔ 赤道原則融資 ➔ 3大永續策略 ➔ 漂綠與RE100風險 ➔ 台積電RE100) ➔ Hour 3 (0050+00878 ESG配置 ➔ 綠色職涯 ➔ 減碳生活 ➔ 防假綠色基金詐騙 ➔ 碳交所Portal查閱 ➔ 定期定額SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十六週預告 (加密貨幣、區塊鏈基礎、Web3 與數位資產)",
         "恭喜學生完成第十五週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人 carbon 體檢與 ESG 永續投資計畫」。預告第十六週課程主題：「加密貨幣、區塊鏈基礎、Web3 與數位資產風險」，下週將帶大家解密區塊鏈、比特幣 (BTC) 減半、以太坊與防範幣圈詐騙！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十六週加密貨幣與區塊鏈主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：企業碳費成本與 CBAM 關稅衝擊試算器",
         "1. 活動目標：幫助學生計算企業年度溫室氣體碳費負擔與外銷歐盟 CBAM 碳關稅衝擊。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入企業碳排公噸數、國內碳費費率與歐盟外銷比例，點擊精算。\n"
         "3. 診斷反思：系統算出國內碳費與 CBAM 關稅金額，驗證「提出自主減量計畫享 $100 元優惠費率能省下千萬碳費」之結論。"),

        ("🎯 第 2 小時實務活動：ESG 篩選 ETF 與綠色溢價評估試算器",
         "1. 活動目標：讓學生親自精算每月定期定額 00878 / 00692 ESG ETF 之預估綠色股息收益與資產累積。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入每月投資金額與預估股息率，點擊精算。\n"
         "3. 決策學習：系統自動產出每年預估領取之綠色股息金額，深化「定期定額 ESG ETF 兼顧收益與支持綠色轉型」的觀念。"),

        ("🎮 第 3 小時小遊戲：ESG 永續理財達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：ESG 三大柱 E 維度；關卡 2：台灣碳費基準費率 300元/噸；關卡 3：歐盟 CBAM 碳關稅；關卡 4：範疇三 Scope 3 碳盤查）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 ESG 永續理財達人徽章 (ESG Master)」，未滿分獲頒「🥉 綠色初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人 carbon 體檢與 ESG 永續投資計畫",
         "1. 作業題目：請學生進行個人碳足跡體檢，線上填寫專屬 0050+00878 ESG 永續投資與減碳計畫。\n"
         "2. 分析要項：(1) 分析 ESG 三大柱與碳費對企業選股之影響；(2) 說明歐盟 CBAM 碳關稅對台灣出口業衝擊與台積電 RE100 綠電承諾；(3) 規劃 0050 (70%) + 00878 (30%) 永續組合；(4) 擬定日常減碳生活與防假綠色基金詐騙 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十五週_課程教學指引_綠色金融ESG永續投資與碳定價經濟學.docx'
    doc.save(doc_path)
    print("Created 第十五週_課程教學指引_綠色金融ESG永續投資與碳定價經濟學.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
