import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第九週 課程教學指引：保險基礎、風險管理與個人保障規劃")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第九週完整教學指引。本單元專門針對大一新鮮人剛滿 18 歲騎機車通勤、打工上學之現實風險設計。課程深入剖析薩繆爾森四象限風險管理矩陣（避免、減輕、承擔與轉嫁）、大數法則與保險 3 大黃金原則（保險利益防賭博、最大誠信據實說明、損害填補防不當得利）；建構新鮮人「意外險 + 實支實付醫療險 + 重大傷病險」高槓桿保障金三角，並解析 2026 金管會實支實付「正本理賠與差額填補」新規；透過「買定期險 + 省下錢投資 0050 ETF」理財算術與「雙十原則」，傳授拒絕高價人情終身保單的 4 大金律與 3,000 元年保費極致配置方案，輔以 FSC 與 TII 權威數據，幫助學生建立獨立的風險護甲。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "風險管理與保險三大基本原則\n(Slide 01 - 10)", "四象限風險矩陣、大數法則算術、保險利益防道德風險、最大誠信據實說明、損害填補原則與全台 2.6張投保率", "🎯 1小時活動：四象限風險矩陣與應對策略檢測器"),
        ("第二小時\n(00:50-01:40)", "新鮮人保障金三角與政策改革\n(Slide 11 - 20)", "機車打工高風險、意外險、實支實付醫療險雜費自費、2026正本理賠改革、重大傷病一次金、勞保強制險與學團險", "🎯 2小時活動：實支實付醫療險理賠與自負額精算器"),
        ("第三小時\n(01:40-02:30)", "保單算術與人情保單避坑實戰\n(Slide 21 - 30)", "定期險 vs 終身險保費點差、買定期+投資ETF雙贏算術、拒絕人情保單 4 金律、雙十原則、3,000元保單與據實告知法規", "🎮 3小時小遊戲：保險保障達人大挑戰 (4大關卡)\n📝 課堂實務作業：個人第一張保單需求評估報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第九週課程導論：保險基礎、風險管理與個人保障規劃", 
         "本頁為第九週課程導論。講師可以引用諾貝爾經濟學獎得主薩繆爾森 (Paul Samuelson) 名言開場：「保險不是為了讓你發財，而是為了防止你暴貧！」針對大一新鮮人剛滿 18 歲騎機車通勤、打工上學的現實生活切入。說明許多學生誤以為保險是長輩才需要考慮的事，或者被親戚拉去買昂貴的儲蓄險。本單元將帶領大家建立正確的風險轉嫁思維，用最小保費買足最大保障。",
         "金融監督管理委員會 (FSC) 保險局 / 財團法人保險事業發展中心 (TII) 官方數據。",
         "問學生：『同學們剛滿 18 歲拿到機車駕照騎車上學時，有沒有想過萬一發生小摔車住院，醫療費用應該由誰負擔？』"),
        (2, 1, "第九週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密四象限風險矩陣、大數法則、保險利益、最大誠信據實說明與損害填補原則；第二小時聚焦於新鮮人機車/打工高風險、意外險、實支實付醫療險、2026 正本理賠改革、重大傷病險與社會保險底護；第三小時則傳授定期險 vs 終身險保費點差、買定期+投資 ETF 雙贏算術、拒絕人情保單 4 金律與 3,000 元高 CP 值保單配置方案。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上圈出自己目前最擔心的風險類型（如：機車摔車、重大疾病、自費醫療）。"),
        (3, 1, "什麼是風險 (Risk)？四象限風險管理矩陣 (Risk Matrix)",
         "詳細拆解風險管理矩陣 (Risk Matrix) 的四個象限。依據「發生頻率」與「損失程度」劃分：1. 高頻小損 ➔ 預防減輕；2. 高頻大損 ➔ 避開規避；3. 低頻小損 ➔ 自行承擔；4. **低頻大損 (低發生率、極大損失) ➔ 轉嫁保險**（如嚴重車禍、癌症）。強調保險專門用來處理災難型風險，絕非用來處理遺失雨傘等小事。",
         "薩繆爾森《風險管理與保險學》理論與實務架構。",
         "提問：『掉了一把雨傘 vs 騎機車撞到超跑住院，哪一個適合用保險來轉嫁？為什麼？』"),
        (4, 1, "保險的本質：大數法則 (Law of Large Numbers) 與互助分攤",
         "講授保險運作的數學基礎——**大數法則 (Law of Large Numbers)**。說明當觀察樣本足夠龐大時，隨機事件發生的概率會趨近於穩定平均值。公平保費公式：`公平保費 = 預期總理賠金額 / 參與投保總人數 + 營業費用`。數學實例：10,000 名學生每人交 500 元建立 500 萬元水庫，當 5 人發生重大意外時每人領 100 萬元理賠，化解個人毀滅風險。",
         "財團法人保險事業發展中心 (TII) 精算原理說明。",
         "提問：『如果只有 10 個人組一個保險互助會，大數法則還管用嗎？為什麼保險公司需要數十萬保戶？』"),
        (5, 1, "保險 3 大黃金原則 (一)：保險利益原則 (Insurable Interest)",
         "詳細剖析保險利益原則 (Insurable Interest)。法律規定要保人對被保險人的生命或財產必須存在法律上認可的利害關係（如本人、配偶、直系血親）。核心目的：1. **避免賭博行為**：防止不相干的人利用他人生命投保賭博發財；2. **杜絕道德風險 (Moral Hazard)**：防止有人為了詐領保險金而誘發謀殺或縱火犯罪。",
         "中華民國《保險法》第 16 條與第 17 條規定。",
         "問學生：『如果你可以幫隔壁完全不認識的同學買 1000 萬人壽保險，社會上會發生什麼可怕的道德風險？』"),
        (6, 1, "保險 3 大黃金原則 (二)：最大誠信原則與據實說明義務",
         "剖析最大誠信原則 (Utmost Good Faith) 與《保險法》第 64 條**據實說明義務**。投保填寫健康告知書時，必須據實告知近 2~5 年內的住院、手術、癌症或慢性病史。法律嚴厲警告：若故意隱瞞病史，保險公司在 2 年內發現可**直接解除合約，且不給付理賠金、亦不退還已繳保費！**",
         "中華民國《保險法》第 64 條據實說明條款與裁決案例。",
         "提醒學生：『聽信業務員口頭說「這個小病不用填」而隱瞞病史，最後受害的一定是你自己！』"),
        (7, 1, "保險 3 大黃金原則 (三)：損害填補原則 (Indemnity Principle)",
         "剖析損害填補原則 (Indemnity Principle)。規定財產保險與實支實付型保單之理賠上限以「實際發生的損失金額」為限，禁止不當得利！公式：`理賠上限 = min(實際醫療自費金額, 保單限額)`。實例：車禍住院實際自費花 5 萬元，即使買了 3 張實支實付保單，總理賠金最高就是填補這 5 萬元損失，不能靠住院多賺錢。",
         "金融監督管理委員會 (FSC) 損害填補原則規範。",
         "強調：『保險是為了補償損失恢復原狀，絕對不是讓你拿來發財或賺錢的工具！』"),
        (8, 1, "實證數據：台灣人平均每人持有 2.6 張壽險保單",
         "引述保險事業發展中心 (TII) 2026 最新官方統計：全台人身保險投保率高達 260% 以上，平均每人持有 2.6 張保單，保險滲透率高達 12.5%（全球頂級）。剖析新鮮人常見痛點：大部分保費都被花在昂貴的「儲蓄險」上，導致實質意外與醫療保障極低，保費負擔過重壓垮大學生生活費。",
         "財團法人保險事業發展中心 (TII) 人身保險業統計年報。",
         "問學生：『為什麼台灣人買了這麼多保單，遇到大病或意外時卻常常發現保額根本不夠用？』"),
        (9, 1, "保險合約 4 大關鍵角色：要保人、被保險人、受益人與保險公司",
         "拆解保單契約書上的四大主體：1. **要保人**：出錢繳保費的人，擁有解約與指定受益人權利；2. **被保險人**：保障標的人體，事故發生在其身上時啟動理賠（如大一學生小明）；3. **受益人**：領理賠金的人，身故保險金可指定父母；4. **保險公司**：承擔風險並依約給付理賠金。",
         "中華民國《保險法》保險契約基本定義。",
         "提問：『大一學生小明自己打工繳保費，保單上的要保人與被保險人分別是誰？』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：風險矩陣處理「低頻高損」災難；保險 3 大原則（保險利益防賭博、最大誠信據實說明、損害填補防不當得利）；看懂要保人、被保險人與受益人。預告第 1 小時 Modal 實務活動——「四象限風險矩陣與應對策略檢測器」，請學生輸入生活風險情境並判定最佳策略。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行風險矩陣判定體驗。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "大一新鮮人高風險場景：機車通勤與打工傷害剖析",
         "剖析 18 歲大一新鮮人最常面臨的兩大現實風險：1. **機車交通事故**：警政署統計 18-24 歲機車傷亡率居各年齡層之冠，摔車擦傷與骨折頻傳；2. **打工與運動傷害**：餐廳燙傷、搬重物拉傷、籃球韌帶斷裂導致醫療費與停工損失。結論：大一新鮮人買保險第一優先順序絕對是「意外險 + 意外醫療實支實付」。",
         "內政部警政署交通統計與勞動部職業災害統計。",
         "問學生：『同學們騎機車上學時，有沒有遇過差點撞到的驚險瞬間？如果摔車住院需要自費 3 萬元，你負擔得起嗎？』"),
        (12, 2, "新鮮人保障金三角 (一)：意外險 (Accident Insurance)",
         "講授傷害保險（意外險）的核心概念。意外定義必須同時符合**「外來的、突發的、非疾病的」**三大要素。介紹兩大保障項目：1. **意外身故/失能保額**：按失能等級 1-11 級給付 100%~5%（保額 200 萬年保費約 1,200 元）；2. **意外醫療實支實付**：門診自費藥膏凭收據實報實銷（保額 5 萬年保費約 600 元）。",
         "金管會保險局傷害保險單示範條款。",
         "提問：『打工時突然盲腸炎發作住院，意外險會理賠嗎？（答案：不會！因為盲腸炎是疾病，不符合「非疾病」要素）』"),
        (13, 2, "新鮮人保障金三角 (二)：實支實付醫療險 (Reimbursement Medical)",
         "剖析現代醫療趨勢「住院天數短、自費材料貴」。傳統日額型醫療險（住院一天給 1,000 元）已落伍，**實支實付醫療險**包含三大核心：1. **病房費限額**：補貼升等單人病房差額；2. **手術費限額**：補貼自費微創手術；3. **醫療雜費限額 (核心!)**：補貼最昂貴的自費標靶藥物、水晶體與心臟支架 (15-20萬)。",
         "中華民國人壽保險商業同業公會實支實付條款說明。",
         "提問：『如果做一個自費微創手術要花 10 萬元，日額型保單給付 3,000 元 vs 實支實付給付 10 萬元，哪一個才能解決問題？』"),
        (14, 2, "2026 最新政策：金管會實支實付醫療險「正本理賠」改革真相",
         "講授 2026 金管會實支實付醫療險**正本理賠與差額填補改革**。過去舊制（副本雙理賠）：花 5 萬拿副本向 2 家申請理賠 10 萬（靠保險賺錢）；2026 新制：必須使用醫院原始正本收據！第一家賠不足部分由第二家開立差額證明填補，總理賠金回歸「損害填補原則」，絕不超過實際花費。",
         "金融監督管理委員會 (FSC) 實支實付醫療險改革公告。",
         "提醒學生：『新制改革後，實支實付回歸補貼真實醫療花費的本質，不要再聽信副本雙理賠賺錢的話術！』"),
        (15, 2, "新鮮人保障金三角 (三)：重大傷病險 (Major Illness)",
         "介紹重大傷病險的強大功能。理賠範圍直接連結衛福部健保「重大傷病卡」（涵蓋癌症、罕見疾病等 300 多項重大疾病）。好處：1. **一次性高額給付**：只要取得重大傷病卡，保險公司一次給付 100 萬元現金；2. **資金高度自由**：可用於昂貴標靶藥物、看護費、營養品或休養期間生活費。",
         "衛生福利部中央健康保險署重大傷病範圍與發卡統計。",
         "提問：『為什麼一次拿到 100 萬現金理賠，比慢慢申請住院門診給付對重病患者更有幫助？』"),
        (16, 2, "勞保 (Labor Insurance) 與機車強制險 (Compulsory Auto) 底層防護",
         "介紹新鮮人已經擁有的法定社會保險：1. **機車強制險**：政府強制投保，保障車禍「對方受害者」體傷最高 20 萬、身故失能最高 200 萬元；2. **勞工保險**：打工店家依法必須申報，提供職災傷病與醫療補助。強調商業保險是用來補足勞保與強制險保障額度不足的中高層缺口。",
         "勞動部勞工保險局 & 財產保險商業同業公會規範。",
         "提醒：『打工第一天一定要確認店家有沒有幫你申報勞保！這是你的法定權益。』"),
        (17, 2, "學生團體保險 (Student Group Insurance) 的隱形權益與申請",
         "解密每學期學雜費單裡的 500 元學生團體保險（學團險）。保障範圍涵蓋疾病或意外住院門診、失能與身故。一般住院或意外醫療金額超過門檻（如 500 元）即可憑醫院收據與診斷書向學校衛保組申請理賠補助！提醒大一新鮮人不要忽視這項高 CP 值的隱形權益。",
         "教育部學生團體保險大專校院作業要點。",
         "指導學生：『在學校體育課扭傷或打工受傷看門診，記得開立診斷書去衛保組申請學團險理賠！』"),
        (18, 2, "理賠爭議防範：除外責任 (Exclusions) 與等待期 (Waiting Period)",
         "剖析理賠爭議兩大主因：1. **除外責任 (Exclusions)**：故意行為（自殘）、犯罪行為、非法酒駕、整形美容或無必要之療養住院，保險公司不給付；2. **等待期 (Waiting Period)**：醫療險生效後通常有 30 天等待期（癌症險常為 90 天），等待期內發病的疾病不理賠。",
         "金融消費評議中心 (FOI) 保險理賠爭議案例彙編。",
         "提醒：『買了醫療險第二天發燒去住院，為什麼保險公司不理賠？（因為還在 30 天等待期內！）』"),
        (19, 2, "健全保單設計三大原則：先保障後儲蓄、先近後遠、先大後小",
         "歸納健全保單設計三大原則：1. **先保障、後儲蓄**：先用低保費買足高額風險保障（意外+醫療），有餘力再考慮儲蓄；2. **先近後遠**：優先解決大一當下 1-5 年內的現實風險，而非 50 年後的退休；3. **先大後小**：優先防範家庭無法承受的大災難（失能、癌症），而非小額感冒。",
         "理財規劃師 (CFP) 個人保險規劃標準原則。",
         "請學生自我檢視：『你的保單有沒有犯了「花大錢買儲蓄險、意外醫療保障卻是零」的錯誤？』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：保障金三角（意外險+實支實付+重大傷病）；2026 實支實付正本理賠改革；勞保、強制險與學團險底護；看清除外責任與等待期；遵循「先保障後儲蓄」原則。預告第 2 小時 Modal 實務活動——「實支實付醫療險理賠與自負額精算器」，請學生輸入醫療自費金額算理賠補貼。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行醫療理賠精算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "定期險 (Term) vs. 終身險 (Whole Life) 保費點差大解密",
         "詳細比對定期險 (Term Insurance) 與終身險 (Whole Life Insurance) 的保費算術點差。相同的 100 萬元身故保障：定期險保障指定期間，保費極便宜（20歲新鮮人每年僅需約 1,000~1,500 元）；終身險繳費 20 年保障一輩子（內含解約金），保費極昂貴（每年高達 25,000~35,000 元），會直接壓垮大一學生生活費。",
         "財團法人保險事業發展中心 (TII) 費率比對數據。",
         "提問：『大一學生打工一年賺 10 萬，如果買一張年繳 3 萬元的終身險，會對生活造成什麼災難？』"),
        (22, 3, "「先定期、後終身」理財算術：用省下的保費做 0050 ETF 定期定額",
         "剖析聰明理財者的雙贏解法——「Buy Term and Invest the Difference (買定期險 + 投資剩餘價差)」。傳統買終身險：每年硬繳 30,000 元買終身險，20 年共繳 60 萬得到低額保障且無流動積蓄；買定期+投資 ETF：每年花 3,000 元買高額定期險，省下的 27,000 元定期定額買 0050/VOO ETF，20 年後獲得高額保障 + 破百萬元股票資產！",
         "理財規劃師 (CFP) \"Buy Term and Invest the Difference\" 試算。",
         "算一算：『每年省下 27,000 元投入年化報酬率 8% 的 ETF，20 年後複利滾存會變成多少錢？（答案：超過 120 萬元！）』"),
        (23, 3, "拒絕人情保單 4 大金律：不買聽不懂的保單、不為人情買高額終身險",
         "傳授親戚阿姨找你買保險時的理性拒絕防線：1. **聽不懂的絕對不買**（搞不懂理賠條件不簽字）；2. **拒絕高額人情終身險**（不為同情人情壓垮自己）；3. **善用 10 天審閱期契撤權**（收到保單 10 天內享有法定契約撤銷權，無條件退還全額保費）；4. **索取條款明細**（堅持看過官方條款）。",
         "金融消費評議中心 (FOI) 消費者自我保護原則。",
         "強調：『收到保單 10 天內隨時可以無條件契撤全額退費，這是法律賦予你的最高防禦護身符！』"),
        (24, 3, "保費預算黃金比例：雙十原則 (保額 10 倍收入，保費不超過 10%)",
         "介紹理財規劃權威的「雙十原則 (10-10 Rule)」：1. **壽險與身故保額 = 年收入 10 倍**（保障失能或身故後家庭 10 年緩衝）；2. **每年保費支出 <= 年收入 10%**（保費絕不可超過年收入一成）。大一學生打工年收 10 萬，年保費應嚴格控制在 **3,000~5,000 元** 以內！",
         "理財規劃師 (CFP) 雙十原則 (10-10 Rule)。",
         "請學生心算：『如果你一個月打工收入 10,000 元（年收入 12 萬），依據雙十原則，你一年的保費上限是多少？』"),
        (25, 3, "實證數據調取：金管會保險局 (FSC) 與保險事業發展中心 (TII) 查閱",
         "手把手教導學生如何查詢權威官方保險數據：1. **保險事業發展中心 (tii.org.tw)**：查詢全台各大保險公司商品條款與試算費率；2. **金管會保險局 (fsc.gov.tw)**：查詢保險公司違規裁罰紀錄；3. **金融消費評議中心 (foi.org.tw)**：查詢各大保險公司的「理賠申訴率」，挑選申訴率低的優良公司！",
         "金管會保險局 暨 財團法人保險事業發展中心 Portal。",
         "演示評議中心網站查詢各大保險公司申訴率的方式。"),
        (26, 3, "大一新鮮人第一張保單 3,000 元年保費極致高槓桿配置方案",
         "展示一年只需 3,000 元的大一高 CP 值保單組合：1. **定期意外險 + 意外醫療實支 (年保費約 1,200 元)**：身故失能 200 萬、意外醫療實支 5 萬；2. **1年期實支實付醫療險 (年保費約 1,800 元)**：住院雜費 15 萬、病房費每日 1,500 元。總保費約 3,000 元/年，一天不到 9 元即可獲得高額防護！",
         "115管理探索二教案保單精算組。",
         "提醒：『一天只需少喝一杯微糖奶茶，就能幫自己築起堅固的醫療意外護甲！』"),
        (27, 3, "保險核保與告知義務：隱瞞病史導致拒賠與解除合約的法律後果",
         "剖析《保險法》第 64 條據實告知義務的法律後果。業務員口頭說「小病不用填」是極危險的話術。法律規定：若要保人未據實填寫健康告知書，保險公司於契約成立 2 年內得解除契約；解除契約時，保險公司**免負給付理賠金責任，且已繳保費一律不退還！**",
         "中華民國《保險法》第 64 條與法院裁判字號條款。",
         "提醒學生：『健康告知書問什麼就答什麼，誠實是獲得理賠保障的唯一道路！』"),
        (28, 3, "大一新鮮人個人保障理性使用 4 大金律",
         "總結新鮮人終身受用的個人保障四大金律：1. **優先購買定期險**（用超低保費買足高額風險保障）；2. **保障先於儲蓄投資**（做足意外與醫療護甲，省下的錢買 0050 ETF）；3. **據實告知健康史**（誠實填寫，不給保險公司解約拒賠藉口）；4. **堅決拒絕人情保單**（聽不懂不簽，善用 10 天契撤權）。",
         "金管會保險局與消費者文教基金會保護原則。",
         "請學生齊聲朗讀四金律，深化理性保險信念。"),
        (29, 3, "第九週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第九週 30 頁純教學卡片進行整體串聯：Hour 1 (風險矩陣 ➔ 大數法則 ➔ 保險利益 ➔ 最大誠信據實說明 ➔ 損害填補 ➔ 2.6張投保率) ➔ Hour 2 (機車打工風險 ➔ 意外險 ➔ 實支實付醫療險 ➔ 2026正本理賠改革 ➔ 重大傷病 ➔ 勞保強制險學團險 ➔ 除外等待期 ➔ 3大原則) ➔ Hour 3 (定期vs終身點差 ➔ 買定期+投資ETF ➔ 拒絕人情保單 ➔ 雙十原則 ➔ 3,000元方案 ➔ 據實告知法規 ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十週預告 (不動產、房地產與租買決策財務學)",
         "恭喜學生完成第九週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人第一張保單需求評估與人情保單體檢報告」。預告第十週課程主題：「不動產、房地產與租買決策財務學」，下週將帶大家探索房貸本息攤還算術、租房 vs 買房財務抉擇與青安貸款！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十週不動產主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：四象限風險矩陣與應對策略檢測器",
         "1. 活動目標：讓學生親自輸入生活中的風險情境（如機車摔車、遺失雨傘、癌症重大手術），判定風險象限並選擇最佳應對策略。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，選擇發生頻率（高/低）與損失程度（大/小），點擊檢測。\n"
         "3. 診斷反思：系統自動分析並強調「保險專門用來處理低頻率、高損失之災難型風險」，矯正學生買錯保險的迷思。"),

        ("🎯 第 2 小時實務活動：實支實付醫療險理賠與自負額精算器",
         "1. 活動目標：幫助學生掌握實支實付醫療險在升等病房費、自費微創手術與高價自費藥材雜費上的補貼算術。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入醫院自費花費與保單約定上限，點擊精算。\n"
         "3. 決策學習：系統自動計算實支實付給付額與個人最終自負額，深化對 2026 正本理賠與損害填補原則的理解。"),

        ("🎮 第 3 小時小遊戲：保險保障達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：保險適用風險判定；關卡 2：2026正本理賠改革法理；關卡 3：定期險 vs 終身險點差；關卡 4：據實告知義務法律後果）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 保險保障達人徽章 (Insurance Master)」，未滿分獲頒「🥉 保險初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人第一張保單需求評估與人情保單體檢報告",
         "1. 作業題目：請學生評估個人大一生活高風險場景，線上填寫第一張保單規劃報告。\n"
         "2. 分析要項：(1) 列出個人大一生活最常接觸之 2 大風險場景；(2) 依據雙十原則設定 3,000-5,000 元定期險保單項目；(3) 說明「買定期險+投資ETF」比買「終身險」優越之算術；(4) 擬定個人據實告知與拒絕人情保單 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第九週_課程教學指引_保險基礎風險管理與個人保障規劃.docx'
    doc.save(doc_path)
    print("Created 第九週_課程教學指引_保險基礎風險管理與個人保障規劃.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
