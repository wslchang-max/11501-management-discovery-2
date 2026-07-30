import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第八週 課程教學指引：期中學習成果檢核與實戰個案總複習")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第八週期中學習成果檢核完整教學指引。本單元係學期前半段（W1 至 W7）之總評與個案融會貫通樞紐。課程精心設計 5 大跨單元實戰個案（含：通膨 2.5% 個人現金流防禦、台積電 2330 財報與匯率敏感度評析、大學生 10 萬資產 30/50/20 黃金配置、海外旅遊雙幣卡 2.5% 回饋精算，以及美台降息宏觀連動傳導鏈），並針對學生期中會考最常陷入之 4 大觀念陷阱（名目vs實質利率扣除、高股息填息與平準金真相、信用卡 15% 循環利息起算點、美金定存 5% 匯損風險）進行深度對決剖析。輔以 DGBAS、TWSE、FSC 與 CBC 四大官方權威資料庫數據鏈，全面驗證與鞏固大一新鮮人之終身財經素養。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "前半學期觀念精華與實戰個案 A/B\n(Slide 01 - 10)", "W1素養與偏差、W2通膨購買力、W3利率與債券反向、W4本益比ROE、個案 A (通膨防禦) 與個案 B (台積電)", "🎯 1小時活動：W1-W7 觀念對決與迷思快問快答擂台"),
        ("第二小時\n(00:50-01:40)", "後半學期精華與跨單元個案 C/D/E\n(Slide 11 - 20)", "W5 ETF微笑曲線、W6高利活存上限與信用卡15%、W7外匯點差與DDC、個案 C (10萬配置)、個案 D (海外精算) 與個案 E (宏觀傳導)", "🎯 2小時活動：大學生 10 萬元綜合資產配置模擬器"),
        ("第三小時\n(01:40-02:30)", "期中會考模擬、常錯陷阱與總評\n(Slide 21 - 30)", "期中會考 4 大題型心法、4 大常錯陷阱解密、個人習慣成長自測、W9-W18 藍圖前瞻與 7 大黃金守則總整合", "🎮 3小時小遊戲：期中會考大挑戰 (4大領域測驗)\n📝 課堂實務作業：期中學習反思與個人綜理財報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第八週課程導論：期中學習成果檢核與實戰個案總複習", 
         "本頁為第八週期中總複習導論。講師可以引用股神巴菲特 (Warren Buffett) 名言開場：「風險來自於你不知道自己在做什麼！」針對大一新鮮人經過前 7 週（財經素養、通膨、利率、股市、ETF、數位金融、外匯）的探索切入。說明期中檢核不是為了死記公式，而是檢驗自己能否將這些知識內化為終身受用的理性決策邏輯。本單元將帶領大家進行跨單元綜合個案演練。",
         "115學年度「管理探索二」課程委員會期中評量規劃網要。",
         "問學生：『經過前 7 週學習，同學們覺得哪一個單元的觀念對你過去的理財認知衝擊最大？』"),
        (2, 1, "第八週 3 小時學習地圖與期中會考應試策略",
         "本頁為學生構建本週 3 小時的期中複習藍圖。第一小時回顧前半學期 (W1-W4) 觀念精華、通膨防禦個案 A 與台積電評析個案 B；第二小時聚焦於後半學期 (W5-W7) 精華、10 萬資產配置個案 C、海外旅遊個案 D 與宏觀傳導個案 E；第三小時則傳授期中會考 4 大題型心法、4 大常錯陷阱剖析、個人習慣成長自測與 7 大黃金守則總整合。",
         "115學年度「管理探索二」期中審查規範。",
         "引導學生在講義上記下自己的期中會考目標得分，準備進入 3 小時深度複習。"),
        (3, 1, "W1 核心重申：資訊爆炸時代的財經素養與認知偏差",
         "重申第一週核心精華：在資訊爆炸時代建立批判性思考防線。剖析兩大常見心理偏差：1. **FOMO (錯失恐懼)**：看到同學或飆股爆賺就急著梭哈，解法是堅持「不懂的商品絕對不買」；2. **倖存者偏差**：只看到少數爆賺網網紅，忽視 90% 慘賠默默離場的散戶，解法是查閱權威機構統計數據。",
         "第一週課程講義與金管會金融知識普及宣導報告。",
         "提問：『如果你在社群平台上看到朋友曬出一天賺 30% 的對帳單，你的第一反應應該是什麼？』"),
        (4, 1, "W2 核心重申：物價通膨 CPI 與實質購買力公式算術",
         "重申第二週核心精華：通膨是無形的財產稅。消費者物價指數 (CPI) 反映生活成本上漲。強調核心算術公式：`實質利率 = 名目利率 - CPI通膨率`。算術實例：把錢存銀行領 1.7% 名目定存利息，但當年度 CPI 通膨率高達 2.5% 時，實質利率為 -0.8%，代表存款一年後實質購買力無形縮水了 0.8%！",
         "第二週課程講義與主計總處 (DGBAS) CPI 統計月報。",
         "計算題：『如果便當從 100 元漲到 110 元 (通膨 10%)，你的打工時薪沒變，你的實質時薪是增加了還是減少了？』"),
        (5, 1, "W3 核心重申：利率政策、升降息與債券價格反向關係",
         "重申第三週核心精華：央行升降息槓桿與債券價格蹺蹺板。中央銀行透過調整重貼現率控制市場資金水頭。強調**「利率與債券價格的反向關係 (Pse-Saw Rule)」**：央行升息 ➔ 新發行債券給予更高利息 ➔ 舊債券價格下跌；央行降息 ➔ 市場利率下降 ➔ 舊的高利債券變得搶手 ➔ 舊債券價格暴漲！",
         "第三週課程講義與中央銀行 (CBC) 貨幣政策報告。",
         "提醒：『買債券 ETF 想賺價差，一定要等到央行【開始降息】的時候，債券價格才會大漲！』"),
        (6, 1, "W4 核心重申：股市入門、P/E 本益比與杜邦分析 ROE",
         "重申第四週核心精華：買股票就是買企業股權。剖析兩大估值指標：本益比 (P/E = 股價/EPS，代表回本年數) 與股價淨值比 (P/B = 股價/每股淨值)。重申杜邦三因子拆解 ROE：`ROE = 淨利率 × 資產轉速 × 權益乘數`。幫助學生一眼看穿企業獲利是靠高毛利（如台積電）、高轉速（如超商）還是高槓桿（如銀行）。",
         "第四週課程講義與 TWSE 公開資訊觀測站 (MOPS)。",
         "提問：『當一家公司的 P/E 高達 50 倍時，代表市場對它的未來成長性抱持什麼期待？』"),
        (7, 1, "實戰個案 A：通膨 2.5% + 升息環境下的個人現金流防禦",
         "剖析大一學生小明每月 15,000 元生活費的綜合個案。情境：當年度 CPI 通膨率 2.5%、便當漲價 10%、銀行定存 1.7%。若小明把所有錢放在 0.8% 活存，實質購買力將縮水 1.7%。防禦對策：1. 留存 3-6 個月生活費於「數位帳戶 2.5% 高利活存」做緊急預備金；2. 剩餘資金投入低費用率之全市場市值 ETF (如 0050/VOO) 抵抗長期通膨。",
         "115管理探索二教案案例組實務模擬。",
         "請學生算一算：『如果你的月生活費是 12,000 元，你的 3 個月緊急預備金應該是多少錢？』"),
        (8, 1, "實戰個案 B：台積電 (2330) 財報、殖利率與企業價值綜合評析",
         "以半導體霸主台積電為例進行三面綜合評析：1. **獲利品質 (DuPont)**：毛利率 > 53%、ROE > 28%，屬於高技術護城河的高淨利率優質企業；2. **本益比 (P/E)**：歷史 P/E 區間於 15 至 25 倍常態震盪，評估安全邊際；3. **匯率敏感度**：美金營收計價，台幣貶值 1 元，台積電營業利益率提升約 0.4 個百分點！",
         "台積電 (TSMC) 官方財務報告與 MOPS 公開資訊。",
         "引導思考：『為什麼新台幣貶值對台積電的財報獲利表現有顯著的正面助攻效果？』"),
        (9, 1, "前半學期常見 5 大財經觀念迷思解密",
         "考前重點解密 5 大常見觀念誤區：1. 10元股票不一定比1000元股票划算（看市值與P/E）；2. 升息時買債券ETF不會立刻大賺（降息才大漲）；3. 高股息沒填息只是左口袋換右口袋；4. 定存 1.7% 扣除通膨後實質獲利為負；5. 聽新聞明牌買股票往往是散戶接盤高點。",
         "115管理探索二團隊彙整常錯題庫。",
         "齊聲複習：『買股票看價值不看單價，看配息要看有沒有填息！』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：W1-W2 理性防範 FOMO 與通膨實質扣除；W3-W4 降息債券價格暴漲與杜邦 ROE；個案 A/B 通膨防禦與台積電評析。預告第 1 小時 Modal 實務活動——「W1-W7 觀念對決與迷思快問快答擂台」，請學生挑戰 5 大迷思題目，測試迷思抗體得分。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行迷思快問快答挑戰。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "W5 核心重申：ETF 狂熱、定期定額微笑曲線與收益平準金陷阱",
         "重申第五週核心精華：被動投資與指數革命。定期定額微笑曲線 (DCA)：在市場下跌時持續扣款，買到更多便宜單位數，拉低平均成本，待反彈實現笑臉獲利。收益平準金新規：金管會明令禁止用平準金衝高配息率，配息來源需透明揭露。大一新鮮人首選「全市場型市值 ETF (如 0050/VOO)」，搭配 70% 核心 + 30% 衛星配置。",
         "第五週課程講義與投信投顧公會 (SITCA) 統計。",
         "提問：『當股市重挫 20% 時，定期定額扣款的人應該感到恐慌停扣，還是感到高興可以買到更多單位數？』"),
        (12, 2, "W6 核心重申：數位金融、高利活存上限與 15% 循環利息黑洞",
         "重申第六週核心精華：數位金融便利性與新鮮人信用保護。數位帳戶高利活存精算：看清「額度上限 (Cap)」，超過部分回歸 0.8% 基本利率。CDIC 300 萬元保障：純網銀與傳統銀行皆享有最高 300 萬存保。信用卡循環利息黑洞：帳戶只繳最低應繳會啟動 15% 循環利息且往回溯及計息！鐵律：永遠全額繳清。",
         "第六週課程講義與金管會銀行局業務規範。",
         "警告：『信用卡帳單永遠只有【全額繳清】一個選項，絕不要只繳最低應繳金額！』"),
        (13, 2, "W7 核心重申：即期 vs 現金匯率、美台利差與升貶值雙刃劍",
         "重申第七週核心精華：個人外匯與國際貿易。即期匯率 (Spot Rate) 為電子交割，匯率最優惠且免現鈔點差；美台利差 (IRP) 引發資金拋售台幣換美元，推升美元走強；警惕美金高利定存 5% 伴隨的台幣升值匯損風險。海外購物與刷卡結算：堅決拒絕 DDC 新台幣結算，一律選擇「當地貨幣 (USD/JPY/RMB)」。",
         "第七週課程講義與中央銀行 (CBC) 外匯局說明。",
         "提醒：『出國刷卡看到 POS 機問你要用 TWD 還是 JPY 結算，永遠果斷按 JPY！』"),
        (14, 2, "跨單元個案 C：大學生 10 萬元資產配置 (高利活存 vs 0050 vs 美金定存)",
         "剖析大一新鮮人 10 萬元打工積蓄的黃金三合一配置：1. **30% 緊急預備金 (3 萬元)**：存入數位帳戶 2.5% 高利活存，兼顧流動性；2. **50% 核心資產 (5 萬元)**：定期定額投入 0050 或 VOO 全市場市值 ETF，享受長期經濟成長；3. **20% 外幣戰略金 (2 萬元)**：分批逢低買進美金/日圓存入外幣帳戶，預備出國與美股投資。",
         "理財規劃師 (CFP) 組合資產配置黃金比例。",
         "請學生計算：『如果 10 萬塊全部買高股息 ETF，遇到股市暴跌 30% 且急需用錢時，會面臨什麼困境？』"),
        (15, 2, "跨單元個案 D：海外購物與旅遊全套精算 (雙幣卡 + 即期匯率 + 2.5% 回饋)",
         "精算大一學生赴日旅遊 50,000 日圓消費最佳化：1. **提前線上即期買日圓**：日圓低點 (0.210) 在 App 即期買入，比臨櫃現鈔價 (0.216) 省下 300 元台幣；2. **海外刷卡結算**：選日圓結算，拒絕 DDC 避開 5% 惡劣手續費；3. **高回饋雙幣卡**：使用海外回饋 3.0% 雙幣卡，扣除 1.5% 手續費後淨賺 1.5% 回饋！",
         "115管理探索二教案實務精算組。",
         "算一算：『透過「即期價+拒絕DDC+雙幣卡」三管齊下，去日本消費 10 萬日圓能總共省下多少錢？』"),
        (16, 2, "跨單元個案 E：美台升降息對外匯、股市與房貸的傳導鏈分析",
         "剖析宏觀經濟「牽一髮而動全身」的連動傳導鏈。當美國聯準會 (Fed) 啟動降息週期時：1. **外匯 (FX)**：美台利差縮小 ➔ 美元指數 DXY 下降 ➔ 資金湧入亞洲，新台幣面臨升值壓力；2. **股市 (Stocks)**：資金成本下降 ➔ 熱錢流入台股與科技股 ➔ 台股與台積電 P/E 獲得向上調升；3. **房貸 (Mortgages)**：央行停止升息，購屋族負擔平穩。",
         "中央銀行與行政院主計總處宏觀經濟模型。",
         "引導思考：『為什麼美國聯準會坐在華盛頓開會降息，會影響到台灣科學園區的股票價格？』"),
        (17, 2, "財經數據權威資料庫綜合查閱：DGBAS / TWSE / FSC / CBC 數據鏈",
         "介紹大一必知的 4 大台灣官方財經資料庫門戶：1. **主計總處 (dgbas.gov.tw)**：查 CPI 通膨與 GDP 成長率；2. **證交所 (twse.com.tw / MOPS)**：查上市公司財報與 P/E 本益比；3. **金管會銀行局 (bank.gov.tw)**：查數位帳戶統計與存保公告；4. **中央銀行 (cbc.gov.tw)**：查每日匯率收盤與外匯存底。",
         "中華民國四大財經部會官方門戶。",
         "演示從 MOPS 網站搜尋特定上市公司財報的步驟。"),
        (18, 2, "大一新鮮人全方位理財資安與防詐護甲",
         "總結保護資產與清白信用的 4 重防禦裝備：1. **2FA 兩步驟驗證**：開啟 App 2FA，OTP 驗證碼視為最高機密絕不透漏；2. **絕不出租金融帳戶**：出借提款卡觸犯《洗錢防制法》，面臨警示帳戶凍結；3. **簡訊短網址防範**：認明 `.com.tw`，不點擊奇怪 `.xyz` 連結；4. **信用卡全額繳清**：遠離 15% 循環利息，建立良好 JCIC 分數。",
         "內政部警政署 165 反詐騙專線與法務部洗錢防制宣導。",
         "強調：『資安與防詐是你理財道路上的第一道防線，防線破裂一切歸零！』"),
        (19, 2, "財務健全度 4 大指標自測 (緊急預備金、負債比、儲蓄率、流動性)",
         "提供大一新鮮人個人財務健康自我健檢表：1. **緊急預備金**：3-6 個月生活費（存於高利活存）；2. **信用卡負債比**：0%（每月全額繳清）；3. **月儲蓄率**：>= 20% 總收入（發薪日先儲蓄後消費）；4. **流動資產比率**：>= 30% 總資產（保留適當現金，避免跌時被迫賣股）。",
         "理財規劃師 (CFP) 個人財務健檢標準表。",
         "請學生心算自己當前的月儲蓄率與緊急預備金月數。"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：W5-W7 被動投資、高利活存與即期外匯；個案 C/D/E 10萬資產配置、海外旅遊精算與降息傳導鏈；官方 4 大資料庫、資安護甲與 4 大財務健全指標。預告第 2 小時 Modal 實務活動——「大學生 10 萬元綜合資產配置模擬器」，引導學生分配比例並計算預期報酬與回撤。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行 10 萬元資產配置模擬演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "期中會考命題方向與 4 大題型解題心法",
         "解析期中會考評量 4 大維度與高分技巧：1. **觀念觀測題 (30%)**：CPI通膨、名目實質利率、升降息債券價格、P/E；2. **算術應用題 (30%)**：通膨購買力、高利活存上限、海外回饋點差；3. **迷思判讀題 (20%)**：填息陷阱、美金定存匯損、15% 循環利息起算點；4. **綜合個案題 (20%)**：10萬資產配置、雙幣卡、降息傳導鏈。",
         "115學年度「管理探索二」期中會考命題小組。",
         "提醒：『做選擇題時，看到「絕對保證獲利」這種誇大字眼，99% 是錯誤選項！』"),
        (22, 3, "觀念題常錯陷阱 1：名目利率與實質利率的通膨扣除",
         "深度解密會考常錯陷阱 1。題目：「小強把 10 萬元存入銀行獲得 2.0% 年利率，一年後購買力增加 2.0%？」解析：錯！必須扣除 CPI 通膨率。若 CPI 通膨率為 2.8%，實質利率為 `2.0% - 2.8% = -0.8%`，實質購買力反而減少了 0.8%！提醒學生考題一定要看清是否有給定通膨率數字。",
         "115管理探索二考前常錯題庫。",
         "問學生：『如果題目給的名目利率是 3%，通膨率是 3.5%，實質利率是多少？（答案：-0.5%）』"),
        (23, 3, "觀念題常錯陷阱 2：高股息 ETF 填息與收益平準金配息真相",
         "深度解密會考常錯陷阱 2。題目：「某高股息 ETF 宣稱年化配息率 10%，代表投資人穩賺 10% 報酬？」解析：錯！除息當天股價會直接扣除配息金額。若股價沒有「填息」，配息只是自己左口袋換右口袋；且配息可能包含本金返還（收益平準金）。強調配息率不等同於投資報酬率！",
         "115管理探索二考前常錯題庫。",
         "強調：『只有股價填息了，配到手上的股息才是真正的獲利！』"),
        (24, 3, "觀念題常錯陷阱 3：信用卡 0% 分期與循環利息的起算點",
         "深度解密會考常錯陷阱 3。題目：「信用卡只繳最低應繳金額，未繳餘額的 15% 循環利息從月底繳款截止日開始計算？」解析：大錯特錯！銀行計息起算日是「往回溯及自店家刷卡入帳日當天」開始按日計息！而且當月所有新消費都失去免息期！",
         "115管理探索二考前常錯題庫。",
         "警示：『循環利息按日回溯計息非常殘酷，絕不能心存僥倖！』"),
        (25, 3, "觀念題常錯陷阱 4：美元 5% 高利定存與台幣升值匯損陷阱",
         "深度解密會考常錯陷阱 4。題目：「美元定存利率 5% 遠高於台幣定存 1.7%，所以存美元一定比存台幣賺？」解析：錯！忽視了匯率貶值風險。若存入期間新台幣大幅升值 8%（USD/TWD 從 32.5 升至 29.9），美元利息 5% 將完全被 8% 匯損吃掉，總帳反而淨虧損 3%！",
         "115管理探索二考前常錯題庫。",
         "齊聲朗讀：『賺了利息，小心賠了匯差！外幣理財必須兼顧匯率走勢。』"),
        (26, 3, "期中大審視：個人前 7 週財經習慣改變與成長自測量表",
         "引導學生進行期中財經習慣成長自測：從入學前隨意開戶、聽明牌買飆股、看到回饋就刷卡繳最低，轉變為**「精算高利活存上限、定期定額 0050/VOO 微笑曲線、信用卡全額繳清、即期分批換匯與拒絕 DDC 結算」**。見證自己在短短 8 週內的驚人財經素養蛻變！",
         "115管理探索二學生學習成效問卷設計。",
         "請學生為自己在前 8 週的財經素養進步打分（1-10 分）。"),
        (27, 3, "後半學期 (W9-W18) 課程藍圖前瞻：從個人理財邁向產業與全球經濟",
         "前瞻期中過後後半學期的課程藍圖。課程視野將從個人理財邁向產業與全球經濟：W9-W12 聚焦於「保險基礎、房地產租買決策與個人報稅規劃」；W13-W18 踏入「半導體 AI 產業鏈、綠色金融 ESG、加密貨幣 Web3 與期末專案發表」。",
         "115學年度「管理探索二」後半學期課程綱要。",
         "激勵學生：『期中過後，我們將帶大家解密半導體 AI 產業鏈與全球經濟的精彩趨勢！』"),
        (28, 3, "期中全方位財經金律總整合 (7 大黃金守則)",
         "總結陪伴大一新鮮人終身受用的 7 大財經黃金守則：1. 不懂的商品絕對不買；2. 實質利率才是真實獲利 (Real Rate)；3. 指數化被動投資長抱 (DCA)；4. 信用卡帳單全額繳清；5. 換匯看即期拒絕 DDC；6. 緊急預備金留足 3-6 個月；7. 2FA 護甲與零出租帳戶。",
         "115學年度「管理探索二」課程核心價值。",
         "請學生將 7 大黃金守則拍照存留於手機，作為終身理財座右銘。"),
        (29, 3, "第八週全景知識體系圖與期中總結",
         "以全景邏輯結構圖將第八週 30 頁純教學卡片進行整體串聯：Hour 1 (前半學期總複習 ➔ 迷思解密) ➔ Hour 2 (後半學期總複習 ➔ 10萬資產配置 ➔ 雙幣卡精算 ➔ 降息傳導 ➔ 4大官方資料庫 ➔ 財務健檢) ➔ Hour 3 (會考4題型心法 ➔ 4大常錯陷阱 ➔ 成長量表 ➔ 7大黃金守則)。",
         "115學年度「管理探索二」期中全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第九週預告 (保險基礎、風險管理與個人保障規劃)",
         "恭喜學生完成第八週期中總複習與學習成果檢核！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「期中學習反思與個人綜合理財報告」。預告第九週課程主題：「保險基礎、風險管理與個人保障規劃」，下週將帶大家探索薩繆爾森風險管理、醫療險與防癌險配置！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後提交期中報告，並預習第九週保險基礎主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：W1-W7 觀念對決與迷思快問快答擂台",
         "1. 活動目標：測試學生對前半學期 5 大財經迷思的抗體與辨識能力，鞏固正確觀念。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，依次回答 5 道判斷題（包含股票單價錯覺、升息債券價格、高股息填息、信用卡循環利息起算點與美金定存匯損）。\n"
         "3. 診斷反思：針對答錯之題目進行即時解析與觀念修正，確保期中會考不重複踩坑。"),

        ("🎯 第 2 小時實務活動：大學生 10 萬元綜合資產配置模擬器",
         "1. 活動目標：讓學生親自練習 30/50/20 資產配置黃金比例，體驗收益與最大可能回撤風險的權衡。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入高利活存、0050/VOO ETF 與美元定存之配置百分比（加總需為 100%），點擊模擬。\n"
         "3. 決策學習：系統自動計算組合之預期年化報酬率與極端市場最大可能回撤，引導學生依個人風險承受度微調配置。"),

        ("🎮 第 3 小時小遊戲：期中會考大挑戰 (4 大領域綜合測驗)",
         "1. 遊戲機制：包含 4 大領域綜合測驗（關卡 1：通膨與實質報酬算術；關卡 2：央行升降息與債券價格；關卡 3：信用卡循環利息起算點；關卡 4：美元高利定存與匯率風險）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 財金探索期中達人徽章 (Midterm Master)」，未滿分獲頒「🥉 財金探索初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：期中學習反思與個人綜合理財報告",
         "1. 作業題目：請學生回顧前 7 週學習歷程，線上填寫期中個人反思與理財規劃報告。\n"
         "2. 分析要項：(1) 總結衝擊最大的 1 個財經迷思與改變；(2) 健檢個人緊急預備金、月儲蓄率與信用卡繳款習慣；(3) 設定 10 萬元積蓄之 30/50/20 黃金配置；(4) 寫下 2FA 資安與防詐 SOP 承諾。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成審閱與期中評分。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第八週_課程教學指引_期中學習成果檢核與實戰個案總複習.docx'
    doc.save(doc_path)
    print("Created 第八週_課程教學指引_期中學習成果檢核與實戰個案總複習.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
