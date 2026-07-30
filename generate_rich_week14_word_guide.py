import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十四週 課程教學指引：生成式 AI 革命、人工智慧產業鏈與未來職場轉型")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十四週完整教學指引。本單元帶領大一新鮮人探索全球科技最前沿的「生成式 AI (Generative AI) 革命」。課程深入剖析 AI 產業鏈三層結構（算力晶片、LLM 大模型、AI 應用軟體）、NVIDIA GPU 霸權與 CUDA 護城河、台灣 AI 伺服器代工霸權 (廣達/緯創/鴻海/英業達) 與水冷散熱；傳授提示詞工程 (Prompt Engineering) 精準提問 4 要素與學習理財應用；精算企業導入 AI 之 ROI 效益與個人生產力提升；探討 AI 幻覺、著作權爭議、隱私資安與歐盟 AI 法案；傳授 AI 淘金熱「賣鏟人」投資策略，並手把手打造個人 AI 自動化協作工作流、鍛鍊大一新鮮人 3 大硬核能力，輔以 Gartner 與 NVIDIA Investor Relations 實證數據，引領學生駕馭 AI 革命。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "生成式 AI 算力與伺服器供應鏈\n(Slide 01 - 10)", "生成式AI定義、AI三層結構、NVIDIA GPU霸權、AI伺服器代工 (廣達/緯創/鴻海)、水冷散熱與 Gartner 實證數據", "🎯 1小時活動：AI 算力成本與伺服器供應鏈估值計算器"),
        ("第二小時\n(00:50-01:40)", "提示詞工程、企業 ROI 與倫理\n(Slide 11 - 20)", "Prompt Engineering 提問技巧、學習/理財實戰應用、企業 ROI 算術、幻覺與智慧財產權、資安與歐盟 AI 法案", "🎯 2小時活動：企業 AI ROI 效益與個人 AI 賦能時間精算器"),
        ("第三小時\n(01:40-02:30)", "AI 協作工作流與未來競爭力\n(Slide 21 - 30)", "個人 AI 工作流、大一 3 大核心能力、Edge AI (AI PC)、Deepfake 防詐、NVIDIA 財報查閱與 4 大金律", "🎮 3小時小遊戲：生成式 AI 應用達人大挑戰 (4大關卡)\n📝 課堂實務作業：個人 AI 工具協作與未來職場轉型報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十四週課程導論：生成式 AI 革命、人工智慧產業鏈與未來職場轉型", 
         "本頁為第十四週課程導論。講師可以引用 NVIDIA 執行長黃仁勳名言開場：「AI 不會取代你，但懂的使用 AI 的人將會取代你！」針對大一新鮮人面臨的 AI 浪潮切入。說明生成式 AI 已從學術實驗室走向全球企業與日常生活。本單元將帶領大家拆解 AI 產業鏈三層結構（算力晶片、LLM 大模型、AI 應用）、NVIDIA 算力霸權、廣達/緯創/鴻海 AI 伺服器代工、提示詞工程與職場替代/賦能算術。",
         "Gartner / IDC 全球 AI 支出報告 / NVIDIA Investor Relations。",
         "問學生：『同學們平時在做報告或寫作業時，已經開始使用 ChatGPT 或 Claude 幫忙了嗎？』"),
        (2, 1, "第十四週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密生成式 AI 定義、AI 三層結構、NVIDIA GPU 霸權、AI 伺服器代工 (廣達/緯創/鴻海)、水冷散熱與 Gartner 數據；第二小時聚焦於 Prompt Engineering 提問技巧、學習/理財 AI 應用、企業 ROI 算術、幻覺與著作權、歐盟 AI 法案與 AI 投資策略；第三小時傳授個人 AI 工作流、大一 3 大能力、Edge AI PC、Deepfake 防詐與 4 大金律。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生思考自己在未來職場中希望扮演被 AI 替代的角色，還是善用 AI 的超級個人。"),
        (3, 1, "什麼是生成式 AI (Generative AI)？從 ChatGPT 到 Sora",
         "講授生成式 AI (Generative AI) 的技術突破：採用深度學習與 Transformer 網絡架構，從傳統「分析數據」跨越至「自主生成全新內容」。範例：1. **文字與程式碼**：ChatGPT (OpenAI)、Claude (Anthropic)、Gemini (Google) 5 秒生成專業分析與程式碼；2. **多模態與視訊**：Midjourney、DALL-E 3、Sora (OpenAI 視訊生成模型) 實現文字轉高清影片。",
         "OpenAI & Google DeepMind 技術白皮書。",
         "演示文字輸入 prompt 瞬間生成高清圖片與影片的驚人過程。"),
        (4, 1, "AI 產業鏈三層結構：算力晶片、LLM 大模型與上層 AI 應用",
         "拆解 AI 產業鏈獲利結構：1. **底層：算力硬體 (Compute)**：GPU 算力晶片與 AI 伺服器 (NVIDIA, TSMC, 廣達)，算力為新石油，毛利極高；2. **中層：大語言模型 (LLM)**：基礎模型訓練 (OpenAI, Google Gemini, Meta Llama)，研發資本密集；3. **上層：AI 應用軟體 (Apps)**：Copilot (Microsoft 365, Midjourney)，龐大 SaaS 訂閱潛力。",
         "CB Insights 全球 AI 產業鏈結構圖。",
         "用淘金熱比喻：『硬體是賣鏟人，LLM 是開採技術，軟體應用是最終加工售出的金飾！』"),
        (5, 1, "算力就是新石油！NVIDIA GPU 晶片架構與黃仁勳 AI 霸權",
         "剖析 NVIDIA 掌控全球 80% 以上 AI 算力的護城河：1. **CUDA 軟體生態系護城河**：NVIDIA 不僅提供 H100/B200 硬體，更憑藉 15 年累積的 CUDA 平台，綁定全球數百萬 AI 工程師；2. **Blackwell B200 晶片**：擁 2,080 億個電晶體，算力提升 5 倍，推動全球資料中心全面升級為 AI 工廠。",
         "NVIDIA GTC 大會技術報告。",
         "提問：『為什麼其他科技巨頭砸重金設計晶片，短期內依然無法撼動 NVIDIA 的 CUDA 霸權？（因為軟體生態系綁定）』"),
        (6, 1, "AI 伺服器供應鏈：廣達、緯創、鴻海與英業達的角色",
         "講授台灣代工巨頭在全球 AI 伺服器中的獨霸地位：全球 80% 以上的 AI 伺服器均由台灣代工大廠組裝製造！1. **廣達 (2382)**：整機系統組裝 (L10/L11) 與水冷櫃設計，主供 NVIDIA NVL72 與 Google；2. **緯創 (3231)**：NVIDIA GPU 晶基板獨家/主要供應商；3. **鴻海 (2317)**：提供垂直整合零組件與整機櫃製造。",
         "Digitimes Research 全球 AI 伺服器調查。",
         "強調：『無論全球 AI 軟體怎麼大戰，最後的實體伺服器機櫃全都要由台灣廣達與鴻海製造！』"),
        (7, 1, "散熱與關鍵零組件：水冷散熱 (Liquid Cooling) 與 CCL 板",
         "講授 AI 伺服器關鍵硬體升級：AI 伺服器單櫃耗電高達 120kW，傳統風扇散熱無法因應，帶動**水冷散熱 (Liquid Cooling)** 革命！台灣**奇鋐 (3017)、雙鴻 (3324)** 提供水冷板與冷卻分配器 (CDU)。高頻高速傳輸則帶動**台光電 (2383)** 超低損耗銅箔基板 (CCL) 需求。",
         "TrendForce 伺服器散熱與零組件研究。",
         "比喻：『AI 晶片就像超級跑車引擎，水冷散熱就是讓引擎高速運轉而不燒毀的水冷系統！』"),
        (8, 1, "實證數據：Gartner 與 IDC 全球生成式 AI 支出預測",
         "引述 Gartner 與 IDC 2026 最新數據：全球 AI 產業總支出將於 2027 年突破 **5,000 億美元 (約 16 兆台幣)**，年複合成長率 (CAGR) 高達 35%！AI 伺服器出貨金額占比已超過全球伺服器總採購金額的 **50%**，成為雲端巨頭資本支出的第一大主力。",
         "Gartner Worldwide IT Spending Forecast / IDC Worldwide AI Tracker。",
         "問學生：『微軟、微軟、亞馬遜與 Google 每年砸幾百億美元買 AI 伺服器，展現了什麼訊號？』"),
        (9, 1, "大一新鮮人看 AI 職場：哪些工作被替代？哪些被賦能？",
         "剖析 AI 時代未來職涯發展的雙刃劍：1. **高度易被 AI 替代的工作**：重複性數據輸入、基礎翻譯、初級程式除錯、常規客服；2. **被 AI 大幅賦能的工作**：跨領域整合者、進階系統架構師、提示詞工程師、AI 財務分析師與策略決策者。結論：學會駕馭 AI 的人將成為職場超級個體。",
         "世界經濟論壇 (WEF)《未來就業報告》。",
         "激勵學生：『不要害怕被 AI 替代，要努力學習成為使用 AI 的超級個人！』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：AI 3 層結構；NVIDIA CUDA 護城河；台灣廣達/緯創/鴻海包辦 80% AI 伺服器代工；水冷散熱趨勢；Gartner 5,000 億美元支出預測。預告第 1 小時 Modal 實務活動——「AI 算力成本與伺服器供應鏈估值計算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行 AI 算力建置成本與伺服器估值試算。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "提示詞工程 (Prompt Engineering) 基礎：精準提問四要素",
         "講授精準操控 AI 的黃金提示詞公式：`黃金提示詞 = 角色設定 (Role) + 背景任務 (Context) + 限制條件 (Constraint) + 輸出格式 (Output Format)`。實例：「你是一位台股資深分析師 (Role)，請分析台積電最新的毛利率表現 (Context)，列出 3 大優劣勢，不要使用複雜專業術語 (Constraint)，以表格形式輸出 (Output Format)。」",
         "OpenAI Prompt Engineering 最佳實務指南。",
         "帶領學生在講義上演練撰寫一組包含 4 要素的專業 Prompt。"),
        (12, 2, "AI 工具在大一學習與財務分析實戰：數據處理與圖表",
         "展示 AI 在大一學習與理財分析的 3 大實戰應用：1. **論文與財報速讀**：上傳 100 頁英文年報至 Claude / ChatGPT，10 秒摘要核心亮點；2. **財務數據試算**：使用 Advanced Data Analysis 自動跑 Python 代碼畫股票圖表；3. **簡報大綱生成**：輸入主題，AI 快速梳理 30 頁邏輯架構。",
         "115管理探索二教案 AI 輔助教學組。",
         "演示將一份資產負債表數據貼入 ChatGPT 並自動生成分析報告的過程。"),
        (13, 2, "企業引進 AI 的 ROI 算術：人力縮減 vs. 算力伺服器成本",
         "講授企業導入 AI 的投資報酬率算術：`AI ROI = (節省人力與時間效益 - AI 伺服器與 SaaS 訂閱費) / AI 總建置成本 * 100%`。麥肯錫實證數據：成功導入 Copilot AI 的企業，員工每處理一份客戶報告可節省 40% 時間，ROI 於 1.2 年內成功回收成本。",
         "McKinsey 麥肯錫《Generative AI 經濟價值報告》。",
         "提問：『企業每個月花 $30 美元為員工訂閱 Copilot，省下的時間價值超過多少？（答案：遠超過 30 美元！）』"),
        (14, 2, "幻覺 (Hallucination) 與著作權陷阱：AI 產出真實性查證",
         "警告盲目信任 AI 的兩大陷阱：1. **幻覺 (Hallucination)**：大語言模型是「機率文本生成器」，有時會本正經地一本胡說八道（引述不存在的法條或虛構數據）；2. **著作權與侵權爭議**：AI 訓練集可能包含未授權內容，直接複製商業銷售可能面臨著作權訴訟。鐵律：**永遠進行事實查證 (Fact-checking)！**",
         "智慧財產局 AI 著作權爭議說明。",
         "提醒：『將 AI 產出的數據或引用文獻拿去 Google 反查，是防範幻覺的必備習慣。』"),
        (15, 2, "數據隱私與資安風險：企業機密外洩與 Deepfake 詐騙",
         "講授數據資安防禦：1. **企業資安機密外洩**：三星工程師曾將未公開晶片原始碼貼入免費版 ChatGPT，導致機密遭訓練集吸收；2. **Deepfake 偽造技術**：詐騙集團利用 3 秒語音樣本複製親友聲音，甚至複製臉部影像進行視訊詐騙。警告：**絕對不能將機密數據與個人身分證字號餵給免費 AI！**",
         "國家資通安全研究院 (NICS) 資安警訊。",
         "提醒：『使用 AI 工具時，永遠不要勾選「允許官方使用我的數據訓練模型」選項！』"),
        (16, 2, "AI 倫理與法規：歐盟 AI 法案 (EU AI Act) 與全球監管",
         "講授全球首部硬性監管 AI 的法案——歐盟 AI 法案 (EU AI Act) 風險三分級：1. **不可接受風險**：社交評分、即時生物識別（全球完全禁止）；2. **高風險**：醫療設備、自動化求職篩選（嚴格審查）；3. **有限風險**：聊天機器人、Deepfake（**強制標註「本內容由 AI 生成」！**）。",
         "歐盟委員會 (European Commission) EU AI Act 官方條文。",
         "討論：『為什麼生成式 AI 產出的圖片或影片，法律要求必須明確標註「AI 生成」？（防止偽造詐騙與誤導）』"),
        (17, 2, "投資 AI 產業鏈理性策略：硬體賣軍火商 vs 上層軟體公司",
         "剖析淘金熱「賣鏟人」投資策略：在金淘熱潮中，買賣鏟子的軍火商最穩定賺錢！1. **淘金賣鏟人 (硬體供應鏈)**：不論哪家大模型贏，都需要向 **NVIDIA、台積電、廣達** 採購算力與伺服器，獲利能見度最高；2. **上層軟體公司**：競爭激烈且易被模型升級替代，需選擇具備黏性用戶生態系者 (Microsoft)。",
         "理財規劃師 (CFP) 科技產業投資架構。",
         "結論：『投資 AI 優先選擇提供算力與伺服器的硬體龍頭，風險最低且最直接受惠。』"),
        (18, 2, "全球科技巨頭 AI 大戰：Microsoft+OpenAI vs. Google vs. Meta",
         "解析三大科技陣營搶奪次世代作業系統的大戰：1. **Microsoft + OpenAI**：ChatGPT-4o 與 Copilot 綁定 Office 365 企業生態（閉源收費）；2. **Google (Alphabet)**：Gemini 1.5 Pro 與搜尋引擎/Android 整合；3. **Meta (Facebook)**：Llama 3 **完全開源 (Open Source)**，號召全球工程師社群免費使用。",
         "MIT Technology Review 全球 AI 趨勢分析。",
         "提問：『Meta 為什麼要將強大的 Llama 3 大模型完全免費開源？（答案：建立生態系，打擊競爭對手的閉源壁壘）』"),
        (19, 2, "台灣科技業在全球 AI 浪潮中的硬體獨霸地位 (市占率>80%)",
         "展現台灣科技業在全球 AI 浪潮中的硬體獨霸地位：從晶圓代工 (TSMC 100%)、GPU 基板 (緯創 90%)、散熱 (奇鋐 70%) 到伺服器整機 (廣達/鴻海 80%)，台灣硬體代工市占率高達 80% 以上！強大的工程整合能力讓矽谷大廠紛紛來台設立 AI 研發中心。",
         "經濟部技術處 AI 產業鏈數據。",
         "讚歎：『全球每一次 AI 算力的升級，都是台灣科技製造業的巨大利多！』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：Prompt 4 要素；AI 在學習/理財之應用；企業 AI ROI 算術；防範幻覺與歐盟 AI 法案；投資硬體賣鏟人策略；台灣 80% AI 伺服器獨霸地位。預告第 2 小時 Modal 實務活動——「企業 AI ROI 效益與個人 AI 賦能時間精算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行企業 AI ROI 與個人時間效益試算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "打造個人 AI 協作工作流：融入日常學習、理財與簡報",
         "手把手教導打造個人每日 AI 自動化多工具協作 SOP：1. **階段 1：收集與摘要 (Information)**：使用 Perplexity / Claude 速讀 50 篇國際財經新聞；2. **階段 2：邏輯分析與試算 (Analysis)**：使用 ChatGPT Code Interpreter 跑 Python 分析數據；3. **階段 3：視覺產出 (Output)**：使用 AI 生成 Markdown 結構，compile 成美觀 HTML 簡報！",
         "個人 AI 生產力工作流設計。",
         "鼓勵學生：『建立自己的 AI 多工具工作流，你一個人就能完成過去一個團隊的工作量！』"),
        (22, 3, "大一新鮮人面對 AI 時代 3 大關鍵能力：提問、批判與整合",
         "講授機器拿不走的勝出關鍵能力：1. **精準提問力 (Prompting)**：拆解複雜任務為子目標，引導 AI 給出最優解答；2. **批判思考力 (Critical Thinking)**：進行事實查證 (Fact-checking)，辨識幻覺與邏輯破綻；3. **跨領域整合力 (Cross-domain Synthesis)**：將財經、管理、資工與設計跨界結合，解決真實問題。",
         "Harvard Business Review《AI 時代的核心人才能力》。",
         "強調：『答案越來越廉價，提問與批判思考能力反而變得無比昂貴！』"),
        (23, 3, "邊緣 AI (Edge AI) 趨勢：AI PC 與 AI 智慧手機換機潮",
         "講授算力從雲端延伸至個人終端裝置的趨勢：AI PC 與 AI 智慧手機內建神經網絡處理單元 (NPU)，能在離線狀態下於筆電/手機端即時執行生成式 AI 任務。Apple Intelligence 與 Copilot+ PC 引爆全球終端換機潮，帶動華碩 (2357)、宏基與聯發科新一波成長。",
         "IDC 全球 AI PC 出貨量市場報告。",
         "提問：『離線狀態下在手機端直接執行 AI 任務，有什麼好處？（答案：極速響應、保護個人數據隱私、不需上傳雲端）』"),
        (24, 3, "AI 時代個人資安防禦：防範 Deepfake 語音/影像複製詐騙",
         "講授 Deepfake 防詐實務：詐騙集團利用 3 秒語音樣本複製親友聲音（AI 語音複製），或視訊通話看到親友臉孔（AI 換臉），要求緊急匯款。破解防詐 2 招：1. **設定家人專屬「防詐通關暗語」**；2. 掛斷電話後，**親自撥打原舊電話號碼重新反查確認！**",
         "內政部警政署 165 反詐騙專線 Deepfake 防詐指引。",
         "警告：『接到親友緊急借錢電話或視訊，絕對要先掛斷，用原電話反查或問暗語！』"),
        (25, 3, "實證數據調取：NVIDIA 財報 (investor.nvidia.com) 實務查閱",
         "手把手教導學生登入 NVIDIA 官方 IR 網站：1. 登入 `investor.nvidia.com` 查閱季度財報中 Data Center 營收年增率 (YoY > 200%)；2. 追蹤全美四大雲端巨頭 (Microsoft, Amazon, Google, Meta) 的資本支出展望；3. 培養閱讀原版英文財報的數據能力。",
         "NVIDIA Investor Relations (investor.nvidia.com)。",
         "演示進入 NVIDIA IR 網站查閱營收報告與投影片的步驟。"),
        (26, 3, "大一新鮮人 AI 時代個人競爭力 4 大金律",
         "總結 AI 時代 4 大金律：1. **積極擁抱 AI 成為超級個人**（融入日常學習理財，提升 300% 生產力）；2. **鍛鍊批判思考與查證能力**（對 AI 產出進行事實查證，嚴防幻覺）；3. **理性投資 AI 硬體賣鏟人**（優先布局半導體與 AI 伺服器龍頭）；4. **嚴防 Deepfake 複製詐騙**（設立家族防詐通關暗語）。",
         "金管會與 AI 科技倫理保護原則。",
         "請學生齊聲朗讀四金律，建立擁抱 AI 的理性信念。"),
        (27, 3, "學習實力試算實例：運用 AI 工具提升 300% 學習與理財 SOP",
         "展示大一新鮮人小明運用 AI 進行通識報告與理財試算實例：傳統人工方式閱讀書籍搜尋網頁耗時 15 小時；AI 協作 SOP（Perplexity 整理架構 + ChatGPT 跑 Python + Claude 生成內容）僅需 2 小時完成高品質報告！節省下來的 13 小時拿去打工賺錢或運動學習，效益極大化。",
         "115管理探索二教案 AI 生產力研究組。",
         "強調：『善用 AI SOP，大學四年你的學習與理財效率將超越同齡人 3 倍以上！』"),
        (28, 3, "擁抱 AI 革命：成為駕馭 AI 的終身學習者與領航者",
         "總結 AI 革命的終極價值：AI 是人類歷史上最強大的文明加速器。面對人工智慧時代，恐懼與抗拒毫無意義。大一新鮮人應以積極開放的心態，將 AI 視為個人智力的延伸，保持好奇心並堅守同理心與道德良知，成為駕馭 AI 的領航者！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『駕馭 AI 工具，開創屬於你的無限可能性未來！』"),
        (29, 3, "第十四週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十四週 30 頁純教學卡片進行整體串聯：Hour 1 (GenAI突破 ➔ 3層產業結構 ➔ NVIDIA CUDA護城河 ➔ AI伺服器組裝 ➔ 水冷散熱 ➔ Gartner數據 ➔ 職場替代vs賦能) ➔ Hour 2 (Prompt 4要素 ➔ 學習理財AI實戰 ➔ 企業AI ROI ➔ 幻覺與著作權 ➔ 數據資安 ➔ 歐盟AI法案 ➔ 投資賣鏟人 ➔ 巨頭AI大戰 ➔ 台灣市占80%) ➔ Hour 3 (個人AI工作流 ➔ 大一3能力 ➔ Edge AI PC ➔ Deepfake防詐暗語 ➔ NVIDIA IR查閱 ➔ 300%效率SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十五週預告 (綠色金融、ESG 永續投資與碳定價經濟學)",
         "恭喜學生完成第十四週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人 AI 工具協作與未來職場轉型報告」。預告第十五週課程主題：「綠色金融、ESG 永續投資與碳定價經濟學」，下週將帶大家探索 ESG 三大柱、碳費與綠色金融！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十五週綠色金融與 ESG 主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：AI 算力成本與伺服器供應鏈估值計算器",
         "1. 活動目標：幫助學生計算企業建置算力中心之 AI 伺服器採購成本與電力功耗，評估台灣供應鏈價值。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入 AI 伺服器機櫃數量、單機預算與功耗，點擊精算。\n"
         "3. 診斷反思：系統算出百萬/億級台幣建置金額與功耗，驗證「算力為新石油，台灣代工與水冷散熱極度受惠」之結論。"),

        ("🎯 第 2 小時實務活動：企業 AI ROI 效益與個人 AI 賦能時間精算器",
         "1. 活動目標：讓學生親自精算企業引進 Copilot AI 後員工節省時間所創造的年收益與投資報酬率 (ROI)。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入團隊人數、平均月薪與 AI 提升時間比例，點擊精算。\n"
         "3. 決策學習：系統自動產出每月與每年省下的薪資時間價值，深化「導入 AI 能迅速帶來數倍正向 ROI」的觀念。"),

        ("🎮 第 3 小時小遊戲：生成式 AI 應用達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：AI 產業鏈受惠硬體層 NVIDIA；關卡 2：黃金提示詞 4 要素公式；關卡 3：台灣 AI 伺服器代工市占率 80%；關卡 4：Deepfake 防詐暗語）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 AI 應用領航員徽章 (AI Master)」，未滿分獲頒「🥉 AI 探索者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人 AI 工具協作與未來職場轉型報告",
         "1. 作業題目：請學生設計專屬 AI 協作工作流，評估個人主修科系未來職場面對 AI 之替代與賦能。\n"
         "2. 分析要項：(1) 運用黃金提示詞 4 要素撰寫一組 Prompt；(2) 規劃日常學習 3 階段 AI 協作 SOP；(3) 評估個人主修科系未來職場任務替代風險並擬定 3 大關鍵能力計畫；(4) 擬定防範 Deepfake 詐騙之家族暗語 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十四週_課程教學指引_生成式AI革命人工智慧產業鏈與未來職場轉型.docx'
    doc.save(doc_path)
    print("Created 第十四週_課程教學指引_生成式AI革命人工智慧產業鏈與未來職場轉型.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
