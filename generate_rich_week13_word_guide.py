import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十三週 課程教學指引：半導體產業、護國神山台積電與台灣科技供應鏈")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十三週完整教學指引。本單元帶領大一新鮮人進入產業探索的核心——「半導體產業」。課程深入剖析摩爾定律 (Moore's Law)、IC 設計 (聯發科)、晶圓代工 (台積電) 與封裝測試 (日月光) 三大分工體系；解密張忠謀創立純晶圓代工模式之破壞式創新與台積電 (2330 TSMC) 三大護城河（技術領先、製造卓越、客戶信任）；拆解 CoWoS 2.5D/3D 先進封裝與 AI 晶片瓶頸；解析地緣政治與美國/日本熊本/德國海外建廠、毛利率 53% 定價權密碼、3-4 年半導體景氣循環、3 大財報指標 (CapEx/GM/利用率)；並說明 0050 ETF 中台積電占 50% 權重之投資策略、ASML EUV 曝光機、商業間諜防禦與 MOPS 營收查閱，建立前瞻之科技產業視野。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "半導體產業鏈與台積電護城河\n(Slide 01 - 10)", "摩爾定律、IC設計/代工/封測三分工、張忠謀破壞式創新、台積電3大護城河與TWSE實證數據", "🎯 1小時活動：台積電先進製程與毛利率估值試算器"),
        ("第二小時\n(00:50-01:40)", "CoWoS 先進封裝與地緣政治\n(Slide 11 - 20)", "CoWoS技術、美日德全球建廠、毛利率53%密碼、半導體景氣循環、3大財報指標與巨頭競合", "🎯 2小時活動：半導體產業鏈與 CapEx 估值精算器"),
        ("第三小時\n(01:40-02:30)", "0050 成分股與新鮮人產業觀\n(Slide 21 - 30)", "0050台積電權重50%、科技職涯探索、ASML EUV曝光機、商業間諜防禦、MOPS查閱與4大金律", "🎮 3小時小遊戲：半導體產業達人大挑戰 (4大關卡)\n📝 課堂實務作業：台積電財務指標與矽島產業鏈分析報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十三週課程導論：半導體產業、護國神山台積電與台灣科技供應鏈", 
         "本頁為第十三週課程導論。講師可以引用張忠謀名言開場：「沒有台積電的晶圓代工模式，就沒有今天全球科技與 AI 的榮景！」針對大一新鮮人身處台灣「半導體矽島」的優勢切入。說明晶片是現代智慧手機、電動車與 AI 伺服器運算的靈魂。本單元將帶領大家剖析台積電 (2330) 的三大護城河、IC 設計與封測分工、CoWoS 先進封裝與台灣科技供應鏈的全球地位。",
         "台灣證券交易所 (TWSE) / 台積電 (2330) 企業年報與 MOPS。",
         "問學生：『同學們知道你口袋裡的 iPhone 或 Android 手機，裡面的 CPU 晶片是在台灣哪一個地方製造出來的嗎？』"),
        (2, 1, "第十三週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密生活中的半導體、摩爾定律、IC設計/代工/封測三分工、張忠謀破壞式創新與台積電 3 大護城河；第二小時聚焦於 CoWoS 先進封裝、美日德全球建廠、毛利率 53% 定價權密碼、半導體景氣循環與 3 大財報指標 (CapEx/GM/利用率)；第三小時傳授 0050 ETF 台積電權重 50%、科技職涯、ASML EUV 曝光機、商業間諜防禦與 MOPS 營收查閱 SOP。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上記下台積電股票代號 2330，準備學習評估科技龍頭股的財報。"),
        (3, 1, "晶片在哪裡？生活中的半導體與摩爾定律 (Moore's Law)",
         "講授半導體 (Semiconductor) 的材料特性（導電性介於導體與絕緣體之間，如矽 Silicon）。幾乎所有電子產品都依賴晶片運算。介紹**摩爾定律 (Moore's Law)**：積體電路上可容納的電晶體數目，約每 18 ~ 24 個月翻倍，效能提升 100%。這條法則由 Intel 創辦人 Gordon Moore 提出，驅動了過去 50 年全球資訊革命。",
         "TSMC 科技博物館半導體歷史。",
         "提問：『如果汽車產業也遵守摩爾定律，現在一台賓士車應該賣多少錢、跑多快？（答案：只要幾塊錢且比火箭快！）』"),
        (4, 1, "半導體產業鏈三大分工：IC 設計、晶圓代工與封裝測試",
         "講授半導體產業專業垂直分工體系：1. **上游：IC 設計 (Design)**：設計電路藍圖（無晶圓廠 Fabless），如 NVIDIA、高通、台灣聯發科 (2454)；2. **中游：晶圓代工 (Foundry)**：將電路圖刻在矽晶圓上，高資本支出，台灣台積電 (2330) 獨霸；3. **下游：封裝測試 (OSAT)**：切割封裝並測試，台灣日月光 (3711) 全球第一。",
         "工研院產科國際所 (ISTI) 半導體產業鏈地圖。",
         "用生活比喻：『IC 設計像建築師畫藍圖，晶圓代工像營造廠蓋鋼筋水泥，封裝測試像室內裝潢與驗屋！』"),
        (5, 1, "為什麼晶圓代工模式能改變世界？張忠謀與台積電的破壞式創新",
         "剖析 1987 年張忠謀創立台積電的破壞式創新：傳統 IDM 模式（Intel、Samsung）自己設計又自己製造，與客戶存在直接競爭。台積電首創「純晶圓代工 (Pure-play Foundry)」：**「絕對不與客戶競爭！」**贏得全球所有 IC 設計公司 (Apple, NVIDIA, AMD) 的 100% 信任，帶動了全球無晶圓廠 IC 設計公司的創業大爆發。",
         "張忠謀自傳與台積電企業史。",
         "強調：『不與客戶競爭的誠信原則，是台積電稱霸全球晶圓代工的最深厚基石！』"),
        (6, 1, "台積電 (2330) 的三大護城河：技術領先、製造卓越與客戶信任",
         "剖析台積電對手無法超越的 3 大護城河：1. **技術領先 (Technology Leadership)**：全球首家量產 3nm/2nm 先進製程，良率 (Yield) 碾壓競爭對手；2. **製造卓越 (Manufacturing Excellence)**：全天候 24 小時極致自動化生產，成本控制與規模龐大；3. **客戶信任 (Customer Trust)**：嚴守客戶商業機密與 IP，成為 Apple 與 NVIDIA 最放心獨家合作夥伴。",
         "台積電 (2330) 投資人關係 (IR) 簡報。",
         "提問：『為什麼對手即使砸幾百億美元蓋晶圓廠，依然無法搶走台積電的 Apple 訂單？（因為良率與信任無法在一夜之間複製）』"),
        (7, 1, "從晶圓 (Wafer) 到晶片 (Chip)：奈米 (nm) 製程與昂貴 Fab",
         "講授微觀晶片製造工程：晶圓 (Wafer) 由高純度單晶矽棒切成（常用 12 吋/300mm）。奈米製程 (3nm/2nm) 指閘極長度微縮度，越小代表相同面積容納更多電晶體，運算更快且省電。強調晶圓廠 (Fab) 造價極度昂貴：一座最新 2nm 晶圓廠造價高達 **200 億美元 (約 6,500 億台幣)**，構建了巨額資本壁壘。",
         "TSMC 技術論壇 (Technology Symposium)。",
         "算一算：『一座 2nm 晶圓廠要 6,500 億台幣，相當於可以蓋幾座台北 101 大樓？（答案：約 11 座 101 大樓！）』"),
        (8, 1, "實證數據：台灣證券交易所 (TWSE) 台積電市值、毛利率與研發",
         "引述台灣證券交易所 (TWSE) 與 MOPS 2026 最新數據：台積電 (2330) 總市值突破 **25 兆元台幣 (約 8,000 億美元)**，名列全球前 10 大企業！毛利率 (Gross Margin) 長期穩定保持在 **53% 以上**（高科技製造業奇蹟）；每年投入研發費用高達 **500 億元台幣以上**，占營收近 8%，持續築高技術壁壘。",
         "台灣證券交易所 (TWSE) 與 MOPS 2026 最新數據。",
         "問學生：『全台上市櫃公司有近 2,000 家，台積電一家公司的市值占了整個台股總市值的幾成？（答案：高達 35% 以上！）』"),
        (9, 1, "台灣半導體矽島產業鏈：聯發科、日月光與設備材料鏈",
         "講授台灣完整半導體生態系：1. **IC 設計霸主聯發科 (2454)**：全球智慧型手機晶片市占率第一，擴展至 AI 天璣晶片與車用；2. **封測全球第一日月光 (3711)**：全球最大的半導體封裝測試廠，包辦先進 SiP 與 CoWoS 後段封測；加上漢唐、家登、辛耘等設備材料供應鏈，打造矽島傳奇。",
         "經濟部產業發展署半導體產業生態系報告。",
         "強調：『台灣半導體強大不只因為台積電，而是擁有一條世界最密集的產業鏈聚落！』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：三分工體系（設計/代工/封測）；純代工模式不與客戶競爭；台積電 3 護城河；市值破 25 兆台幣與 53% 超高毛利率。預告第 1 小時 Modal 實務活動——「台積電先進製程與毛利率估值試算器」，請學生輸入營收與毛利率精算估值。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行台積電估值試算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "什麼是 CoWoS 先進封裝？超越摩爾定律 (More than Moore)",
         "解密 AI 晶片的算力催化劑——**CoWoS (Chip-on-Wafer-on-Substrate)** 2.5D/3D 先進封裝技術：將 GPU (如 NVIDIA B200) 與高頻寬記憶體 (HBM) 垂直/水平緊密拼接在同晶圓上，大幅縮短傳輸距離、降低功耗並提升資料頻寬。鐵律：**沒有 CoWoS 封裝產能，NVIDIA 的 AI 晶片就無法順利出貨！**",
         "TSMC CoWoS 技術白皮書。",
         "提問：『為什麼光靠 3nm 晶片還不夠，一定要加上 CoWoS 先進封裝，AI 晶片才能發揮最強算力？』"),
        (12, 2, "地緣政治與半導體供應鏈：美國 CHIPS Act、日本熊本與德國廠",
         "解密「矽盾 (Silicon Shield)」與台積電全球建廠布局：1. **美國亞利桑那廠 (Fab 21)**：因應美《晶片法案》(CHIPS Act) 補貼，設置 4nm/3nm 先進製程；2. **日本熊本廠 (JASM)**：攜手 Sony 與 Toyota 滿足車用與感測器需求；3. **德國德勒斯登廠 (ESMC)**：深入歐洲車用樞紐。說明地緣政治下的晶片供應鏈分散趨勢。",
         "台積電全球營運拓展說明。",
         "討論：『為什麼各國政府都砸重金邀請台積電去當地蓋晶圓廠？（因為晶片已成為戰略國家安全物資）』"),
        (13, 2, "晶圓代工定價權 (Pricing Power) 與毛利率 53% 財報密碼",
         "剖析先進製程的高定價權：在 3nm 先進製程領域，台積電擁有超過 **90% 的全球獨家市占率**！帶來極強的**訂價權 (Pricing Power)**！公式：`毛利率 = (營收 - 成本) / 營收 >= 53%`。即使每片先進晶圓售價高達 20,000 美元，Apple 與 NVIDIA 依然搶著預訂產能，確保超高獲利能力。",
         "台積電 (2330) 法人說明會 (Earnings Call)。",
         "強調：『擁有別人做不出來的獨家技術，企業就握有訂價權，可以維持 53% 的超高毛利率！』"),
        (14, 2, "半導體景氣循環 (Semiconductor Cycle)：庫存調整與終端需求",
         "講授半導體產業 3-4 年景氣循環：1. **繁榮擴張期 (Boom)**：終端需求暴增、搶產能、代工漲價，產能利用率 > 100%；2. **庫存調整期 (Bust)**：消費電子疲軟、客戶砍單、產業去庫存，庫存天數 (DOI) 攀升；3. **復甦上升期 (Recovery)**：新應用 (AI/車用) 帶動需求重啟，資本支出重包。",
         "WSTS 世界半導體貿易統計組織報告。",
         "提問：『在半導體景氣去庫存谷底時，聰明的長期投資人會怎麼做？（答案：冷靜分批布局，等待復甦週期）』"),
        (15, 2, "台灣電子零組件出口占比與經濟部/財政部海關數據",
         "引述經濟部與財政部海關最新統計數據：台灣電子零組件出口占全台總出口金額高達 **40% ~ 45%！** 其中積體電路 (IC) 為最核心大宗。半導體產業每年為台灣帶來數千億美元的貿易順差，維繫了新台幣匯率與外匯存底的強健穩定。",
         "財政部關務署進出口貿易統計報告。",
         "數據強調：『半導體產業一個產業就貢獻了台灣近一半的外銷出口金額！』"),
        (16, 2, "大學生看懂半導體財報 3 大指標：CapEx、GM 與利用率",
         "教導評估晶圓代工廠 competitive 的 3 大指標：1. **資本支出 (CapEx)**：每年買曝光機蓋廠金額，增加代表對未來先進製程需求極具信心；2. **毛利率 (Gross Margin)**：反映技術獨佔與成本控制，標竿 53%；3. **產能利用率 (Utilization)**：設備滿載運作程度，> 95% 代表訂單塞爆。",
         "理財規劃師 (CFP) 科技股財報分析標準。",
         "演示如何從台積電季報尋找 CapEx、毛利率與利用率三個關鍵數字。"),
        (17, 2, "科技股投資風險：景氣循環谷底、資本支出過度與斷鏈風險",
         "警告半導體投資的潛在風險：1. **景氣循環谷底風險**：手機/PC 需求下滑導致利用率下降、毛利率壓縮；2. **資本支出過度沉沒**：若建廠過度，巨大的固定折舊費用會侵蝕獲利；3. **地緣政治與斷鏈風險**：戰爭、地震或貿易制裁中斷關鍵材料或設備供應。",
         "公開資訊觀測站 (MOPS) 台積電風險因素揭露。",
         "提醒：『投資半導體股不能只看順風期的高獲利，也要評估逆風期的折舊壓力。』"),
        (18, 2, "全球半導體巨頭競合：TSMC vs. NVIDIA vs. Intel vs. ASML",
         "解析全球科技巨頭生態圈：NVIDIA 是全球 AI 晶片霸主（GPU 設計），100% 依賴台積電 CoWoS 產能，是台積電最大客戶之一；ASML 是荷蘭極紫外光曝光機獨家供應商，是台積電關鍵設備軍火商；Intel 與 Samsung 既是代工對手，同時亦委託台積電代工先進晶片！",
         "TrendForce 集邦科技半導體研究報告。",
         "總結：『在全球半導體舞台上，台積電處於所有科技巨頭環繞的生態圈核心位置！』"),
        (19, 2, "台灣科技供應鏈群聚效應：竹科、中科、南科與高雄 S 廊帶",
         "講授台灣「一小時半導體供應鏈」的聚落奇蹟：從竹科、中科、南科到高雄半導體 S 廊帶。工程師、設備商搭高鐵可在 1 小時內趕赴全台任何一座 Fab 廠調試設備。全世界沒有任何國家能複製台灣這種超高效率的半導體地理群聚聚落！",
         "國科會三大科學園區統計年報。",
         "讚歎：『高鐵串聯的矽島 S 廊帶，是台灣最無法被取代的實體競爭障礙！』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：CoWoS 封裝為 AI 晶片心臟；美日德海外建廠布局；毛利率 53% 定價權；半導體 3-4 年循環；3 大財報指標 (CapEx/GM/利用率)。預告第 2 小時 Modal 實務活動——「半導體產業鏈與 CapEx 估值精算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行 CapEx 與產能效益試算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "從投資者角度看台積電 (2330) 與 0050 ETF 權重成分",
         "剖析 0050 ETF 的台積電權重成分：在元大台灣 50 (0050 ETF) 中，台積電 (2330) 權重占比高達 **50% 以上！** 投資觀點：買入 0050 ETF 等同於一半資金參與台積電成長，另一半資金分散至聯發科、鴻海、日月光等供應鏈，是小資族打包台灣半導體的最優解。",
         "元大投信 0050 成分股最新比重公告。",
         "提問：『如果你想投資台積電但又怕單一股票波動太大，買 0050 ETF 有什麼好處？（答案：兼具半導體龍頭成長與指數分散效益）』"),
        (22, 3, "大一新鮮人科技產業職涯探索：工程師 vs 供應鏈 vs 理財",
         "引導大一學生進行科技產業跨領域職涯探索：1. **研發與製程工程師**：電機、物理、化學、材料背景，於 Fab 廠挑戰先進製程極限；2. **供應鏈與營運管理**：商管、資管背景，負責全球資材採購與客戶專案；3. **科技財務與法務**：財金、法律背景，處理百億美元 CapEx 投資與 IP 專利。",
         "104 人力銀行科技產業人才白皮書。",
         "鼓勵：『文法商管科系的同學，只要掌握半導體產業鏈與外語能力，同樣能進入台積電擔任營運高階主管！』"),
        (23, 3, "關鍵材料與設備供應商：ASML 極紫外光 (EUV) 曝光機",
         "介紹全球最精密機器——ASML EUV 光刻機：一台售價高達 1.8 億美元 (約 58 億台幣)。波長僅 13.5nm，使用雷射擊打錫滴產生極紫外光，在晶圓下記錄幾奈米電路。台積電擁有全球最大規模的 EUV 曝光機陣容，確保先進製程產能遙遙領先。",
         "ASML 艾司摩爾官方技術年報。",
         "展示 ASML EUV 機器龐大的物理體積（需要 3 架波音 747 貨機才能運輸一台！）。"),
        (24, 3, "產業資安與智慧財產權 (IP)：防範商業間諜與技術外洩",
         "講授國家安全與營業秘密保護：外國競爭對手高薪挖角與竊取 2nm 參數的商業間諜威脅。中華民國《國家安全法》修法將核心關鍵技術竊密列為重罪。台積電實施最高規格專利與營業秘密保護系統 (PIP)，嚴禁攜帶未授權裝置進入 Fab 廠區。",
         "中華民國《國家安全法》與《營業秘密法》條文。",
         "警告：『商業機密是科技業的生命線，竊取核心關鍵技術將面臨重刑起訴！』"),
        (25, 3, "實證數據調取：公開資訊觀測站 (mops.twse.com.tw) 查閱營收",
         "手把手教導學生使用 MOPS Portal 查詢台積電數據：1. 登入 `mops.twse.com.tw` 輸入代號 `2330`，查閱每月 10 日前通報之單月合併營收與 YoY 年增率；2. 查閱每季公布之先進製程 (3nm/5nm) 營收占比與毛利率數據，養成數據驗證投資習慣。",
         "公開資訊觀測站 (mops.twse.com.tw)。",
         "演示手機進入 MOPS 網站查詢台積電營收與季報 PDF 的步驟。"),
        (26, 3, "大一新鮮人半導體與科技投資 4 大金律",
         "總結半導體投資 4 大金律：1. **認清護城河與定價權**（優先選擇技術獨占與毛利率 > 50% 之龍頭台積電）；2. **善用 0050 ETF 分散風險**（一次打包台積電與整體供應鏈）；3. **理解景氣循環不追高**（在 3-4 年去庫存谷底分批布局）；4. **密切關注地緣政治**（評估海外建廠成本與供應鏈移轉）。",
         "金管會與半導體產業分析保護原則。",
         "請學生齊聲朗讀四金律，建立理性科技投資觀念。"),
        (27, 3, "投資試算實例：定期定額參與台灣半導體成長股利複利 SOP",
         "展示定期定額零股買入台積電 (2330) 或 0050 ETF SOP：每月只需 1,000 ~ 3,000 元即可在券商 App 設定自動定期定額買零股。台積電採按季配息（每季約 4 元以上），將發放之季度股利自動再買入零股，發揮股利滾存複利威力！",
         "115管理探索二教案科技投資組。",
         "提醒：『大一學生每月存 2,000 元買台積電零股，畢業時你就成了護國神山的股東！』"),
        (28, 3, "矽島台灣與國際財經宏觀展望",
         "總結台灣在全球 AI 與半導體時代的核心地位：隨著生成式 AI 爆發，世界對晶片算力需求無窮無盡。大一新鮮人立足台灣，具備第一線觀察頂尖科技巨頭動向的優勢，理解晶片如何串連全球地緣政治、外匯與資本市場，開擴宏觀國際視野！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『身在矽島台灣，要對台灣的科技產業優勢充滿自信與國際視野！』"),
        (29, 3, "第十三週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十三週 30 頁純教學卡片進行整體串聯：Hour 1 (摩爾定律 ➔ 3大階段分工 ➔ 純代工模式 ➔ 台積電3護城河 ➔ 先進製程與Fab造價 ➔ 2330市值/毛利53% ➔ 矽島生態系) ➔ Hour 2 (CoWoS先進封裝 ➔ 美日德建廠 ➔ 53%毛利定價權 ➔ 3-4年循環 ➔ 電子出口40% ➔ 3財報指標 ➔ 巨頭競合 ➔ S廊帶) ➔ Hour 3 (0050台積電權重50% ➔ 跨領域科技職涯 ➔ ASML EUV ➔ 商業間諜防禦 ➔ MOPS營收查閱 ➔ 定期定額SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十四週預告 (生成式 AI 革命、人工智慧產業鏈與未來職場)",
         "恭喜學生完成第十三週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「台積電財務指標與矽島產業鏈分析報告」。預告第十四週課程主題：「生成式 AI 革命、人工智慧產業鏈與未來職場轉型」，下週將帶大家探索 ChatGPT/Claude AI 算力革命與未來職場轉型！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十四週生成式 AI 主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：台積電先進製程與毛利率估值試算器",
         "1. 活動目標：幫助學生掌握晶圓代工廠估值方法，分析先進製程營收與毛利率對市值的驅動。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入台積電預估營收、毛利率與 P/E 本益比倍數，點擊精算。\n"
         "3. 診斷反思：系統算出估算毛利與市值估值，驗證「53%+ 超高毛利率是台積電市值突破 25 兆台幣核心支柱」之結論。"),

        ("🎯 第 2 小時實務活動：半導體資本支出 (CapEx) 與產能效益試算器",
         "1. 活動目標：讓學生親自精算晶圓代工廠龐大 CapEx 投資在產能利用率高檔時對未來營收增長之效益。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入年度 CapEx 億美元與 Fab 利用率，點擊精算。\n"
         "3. 決策學習：系統自動算出折合台幣 CapEx 與未來新增營收，深化「高產能利用率能確保龐大 CapEx 快速轉化為營收」的觀念。"),

        ("🎮 第 3 小時小遊戲：半導體產業達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：純晶圓代工模式創始者張忠謀；關卡 2：台積電毛利率 53% 標竿；關卡 3：CoWoS 先進封裝地位；關卡 4：0050 ETF 台積電權重 50%）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 半導體產業達人徽章 (Semiconductor Master)」，未滿分獲頒「🥉 科技初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：台積電財務指標與矽島產業鏈分析報告",
         "1. 作業題目：請學生分析台積電 3 大護城河與 0050 權重，線上填寫台積電財務指標與矽島產業鏈分析報告。\n"
         "2. 分析要項：(1) 分析台積電 3 大護城河與 53% 毛利率定價權；(2) 說明 CoWoS 先進封裝對 AI 晶片出貨與海外建廠影響；(3) 評估 0050 ETF 中台積電占 50% 權重之投資效應；(4) 登入 MOPS 查閱台積電每月營收並擬定定期定額零股買入 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十三週_課程教學指引_半導體產業護國神山台積電與台灣科技供應鏈.docx'
    doc.save(doc_path)
    print("Created 第十三週_課程教學指引_半導體產業護國神山台積電與台灣科技供應鏈.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
