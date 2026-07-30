import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十七週 課程教學指引：期末專案發表、實戰財經簡報與總評量")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十七週完整教學指引。本單元為全學期學習成果總驗收「期末專案發表、實戰財經簡報與總評量」。課程深入剖析 Guy Kawasaki 商業簡報 10/20/30 法則（10 頁投影片、20 分鐘發表、30 號字體）；傳授數據視覺傳達力、大數字卡片與對比圖表設計；演練控時、眼神接觸、肢體語言與雙語專有名詞流暢度；建立 Q&A 評審問答防守與 Backup Slide 展現數據口袋戰術；由全班 6 大小組依次登場發表（包含 30 萬理財第一桶金、租房買房決策、台積電護城河與 0050、AI ROI 算術、ESG 00878、Web3 冷錢包 5% 配置）；實施「Keep / Improve / Learn」同儕匿名互評機制與專家評審團總講評，並指導學生將發表成果納入大一學習歷程 Portfolio，輔以教育部商業簡報競賽評分標準，引領學生邁向卓越財經溝通者。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "簡報法則、數據表達與小組發表 1-2\n(Slide 01 - 10)", "期末專案架構、10/20/30法則、數據視覺化、控時與Q&A戰術、第一/二小組專案發表與講評", "🎯 1小時活動：簡報發表控時與評分標準試算器"),
        ("第二小時\n(00:50-01:40)", "小組發表 3-6、同儕互評與專家講評\n(Slide 11 - 20)", "第三組 (半導體)、第四組 (AI ROI)、第五組 (ESG)、第六組 (Web3) 發表、同儕互評機制與專家總講評", "🎯 2小時活動：期末同儕互評與組別總成績精算器"),
        ("第三小時\n(01:40-02:30)", "履歷 Portfolio、表達氣場與 4 大金律\n(Slide 21 - 30)", "將專案納入履歷 Portfolio、3大表達能力、著作權保護、10分鐘 Pitching SOP 與 4 大金律", "🎮 3小時小遊戲：財經簡報表達達人大挑戰 (4大關卡)\n📝 課堂實務作業：個人期末簡報反思與學習歷程報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十七週課程導論：期末專案發表、實戰財經簡報與總評量", 
         "本頁為第十七週課程導論。講師可以引用賈伯斯名言開場：「簡報不是報告數據，而是傳達你改變世界的願景！」針對大一新鮮人經過 16 週學習後的驗收切入。說明簡報發表是商業與管理舞台的核心競爭力。本單元將帶領大家拆解 10/20/30 簡報法則、數據視覺化表達、Q&A 防守戰術，並進行全班 6 大小組發表、專家同儕雙重評分。",
         "教育部大一簡報競賽評分指南 / 115管理探索二課程委員會。",
         "問學生：『同學們在過去 16 週的課程中，哪一個主題（如租房買房、台積電或 AI）最啟發你？你們小組專案選了什麼主題？』"),
        (2, 1, "第十七週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密期末專案架構、10/20/30法則、數據視覺化、演練技巧、Q&A 戰術、評分標準與第一/二小組發表；第二小時聚焦於第三至六小組專案發表、同儕互評機制 (Keep/Improve/Learn)、專家總講評與財經 Storytelling；第三小時傳授履歷 Portfolio 打造、3 大表達能力、專案落地 SOP、智財權保護與 4 大金律。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上記下 10/20/30 法則與評分標準，準備在發表舞台展現最佳表現。"),
        (3, 1, "期末專案簡報架構：從痛點分析、財務試算到價值主張",
         "剖析專業財經簡報三大邏輯區塊：1. **痛點與現狀分析 (Pain Points)**：明確指出大一新鮮人面臨的理財/租房痛點，引用數據佐證；2. **財務模型與試算 (Financial Model)**：展示精確試算表（如 30 萬儲蓄、房貸負擔率、ETF 報酬）；3. **落地行動方案 (Actionable SOP)**：提供步驟化 SOP 與風險防禦。",
         "McKinsey 麥肯錫商業簡報架構指南。",
         "強調：『一份好的財經簡報，必須具備深刻的痛點、嚴謹的財務試算與可執行的 SOP！』"),
        (4, 1, "專業財經簡報 10/20/30 法則：10頁、20分鐘與 30號字",
         "講授矽谷創業大師 Guy Kawasaki 傳授的黃金簡報公式：`黃金簡報法則 = 10 頁精準簡報 + 20 分鐘發表時間 + 30 號內文 Huge 字體`。說明 10 頁投影片精簡聚焦，20 分鐘留出評審 Q&A，30 號 Huge 字體逼迫講者只放關鍵字，絕不朗讀文字稿！",
         "Guy Kawasaki《The Art of the Start》簡報法則。",
         "提問：『為什麼 Guy Kawasaki 要求內文字體不得小於 30 號？（答案：逼迫發表者精簡文字，用口頭演講說服聽眾）』"),
        (5, 1, "視覺化數據表達：用圖表與關鍵數字打動評審與投資人",
         "講授數據視覺化表達技巧：讓數據自己說話！1. **善用大數字卡片**：把關鍵結論放大（如「300% 效率提升」或「17.4 萬畢業積蓄」）；2. **圖表對比勝過千言萬語**：用長條圖呈現「租房 vs 買房 30 年資產差異」，讓評審一目了然。",
         "HBR《Visualizing Data for Impact》。",
         "對比：『密密麻麻的文字表格讓評審昏昏欲睡；大數字卡片與莫蘭迪對比圖表讓評審眼前一亮！』"),
        (6, 1, "簡報演練技巧：控時、眼神接觸、肢體語言與雙語流暢度",
         "演練發表前 4 項基本功：1. **控時精準**：演練至時間結束前 30 秒完美收尾；2. **眼神接觸 (Eye Contact)**：掃視評審全場，保持自信笑臉；3. **肢體語言 (Body Language)**：自然手勢，避免背對觀眾；4. **雙語專有名詞**：流暢切換 ROI, CapEx, ESG, Prompt, Blockchain。",
         "Toastmasters 國際演講協會商業簡報技巧。",
         "指導學生在台下進行 1 分鐘雙語自我介紹與控時演練。"),
        (7, 1, "評審問答 (Q&A) 應對戰術：防守型回答 vs. 延伸價值答辯",
         "剖析評審 Q&A 高分應對戰術：被問到尖銳問題時，切忌盲目硬凹或啞口無言。完美 3 步驟：1. 感謝評審提問；2. 簡潔回答核心數據；3. 展現邏輯：「此部分我們於備用投影片 (Backup slide) 有詳細精算！」展現充足準備。",
         "華爾街投行 Pitching Q&A 應對守則。",
         "示範：『秀出事先準備好的 Backup Slide，讓評審對你們小組的嚴謹態度大加讚賞！』"),
        (8, 1, "實證數據：商業簡報評分標準表與教育部比賽指引",
         "解析期末簡報四大權重評分標準：1. **財務與管理邏輯嚴謹度 (35%)**：公式試算正確、數據權威；2. **痛點解決與價值主張 (25%)**：SOP 具落地可行性；3. **視覺設計與圖表品質 (20%)**：美觀大方、大數字；4. **上台表達與 Q&A 氣場 (20%)**：控時精準、雙語流暢。",
         "115學年度「管理探索二」期末簡報評分標準表。",
         "引導學生對照評分標準表，確保小組專案每一分都拿在點子上。"),
        (9, 1, "第一小組登場：大一新鮮人 30 萬理財第一桶金專案發表",
         "第一小組登場發表專案《大一新鮮人 30 萬理財第一桶金規劃》：1. **痛點**：大一零頭期款與薪資無規劃；2. **財務模型**：每月打工 15,000 元提撥 3,000 元定期定額買 0050+00878，配合 2.5% 高利活存，4 年累積 30 萬元！評審講評：財務模型嚴謹，0050+00878 具避險效果！",
         "第一小組期末專案簡報備忘錄。",
         "全班給予第一小組熱烈掌聲，評審進行即時點評。"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：10/20/30 簡報法則；數據視覺化大數字；Q&A 戰術與 Backup Slide；評分標準 4 維度；第一小組發表展示。預告第 1 小時 Modal 實務活動——「簡報發表控時與評分標準試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行發表控時與加權分數試算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "第二小組登場：租房 vs. 買房財務決策與房貸負擔率發表",
         "第二小組登場發表專案《租房 vs. 買房財務決策》：1. **痛點**：求學租房租金負擔與未來 30 年買房巨額房貸壓力；2. **財務模型**：建立 $R/M < 60\%$ 租金房貸比與 $<33\%$ 房貸負擔率試算模型，評估 30 年淨資產差異。評審講評：引用地政司數據量化決策，具高度實用性！",
         "第二小組期末專案簡報備忘錄。",
         "評審就租金指數與房貸負擔率提出 Q&A 互動。"),
        (12, 2, "第三小組登場：半導體台積電護城河與 0050 投資組合發表",
         "第三小組登場發表專案《半導體台積電護城河與 0050 戰略》：1. **痛點**：想投資台積電但資金不足且懼怕高波動；2. **財務模型**：零股定期定額買台積電 + 配置 0050 ETF (台積電權重 50%)，追蹤毛利率 53% 與 CapEx。評審講評：半導體供應鏈剖析透徹，架構專業！",
         "第三小組期末專案簡報備忘錄。",
         "評審讚賞第三小組對 CapEx 與毛利率之掌握。"),
        (13, 2, "第四小組登場：生成式 AI 工具輔助企業效率與 ROI 發表",
         "第四小組登場發表專案《生成式 AI 提示詞工程與個人效率提升》：1. **痛點**：害怕被 AI 替代，不知如何將 AI 融入學習；2. **實踐 SOP**：展示 Prompt 4 要素 (Role/Context/Constraint/Format) 與 3 階段 AI 協作 SOP，縮短 75% 報告時間！評審講評：ChatGPT 現場演示完美，掌握超級個人精髓！",
         "第四小組期末專案簡報備忘錄。",
         "第四小組現場演示 10 秒生成 Markdown 簡報架構。"),
        (14, 2, "第五小組登場：綠色金融 ESG (00878) 永續資產配置發表",
         "第五小組登場發表專案《綠色金融 ESG 永續資產配置》：1. **痛點**：希望理財同時避免資金流入高污染企業；2. **財務與永續模型**：計算台灣碳費 (300元/噸) 與 CBAM 關稅衝擊，建構 00878 ESG 高股息永續配置。評審講評：結合氣候經濟學與永續理財，具世代正義影響力！",
         "第五小組期末專案簡報備忘錄。",
         "評審方針對 CBAM 碳關稅對台灣鋼鐵業衝擊提出討論。"),
        (15, 2, "第六小組登場：Web3 數位資產 5% 避險配置與冷錢包發表",
         "第六小組登場發表專案《Web3 數位資產 5% 避險配置》：1. **痛點**：想參與 Web3 爆發力但懼怕高波動與 Telegram 詐騙；2. **資安與財務模型**：95% 傳統資產 + 5% 加密貨幣 (BTC/ETH)，提領至冷錢包離線保管私鑰。評審講評：冷錢包資安強調到位，防詐 SOP 切中要害！",
         "第六小組期末專案簡報備忘錄。",
         "第六小組展示冷錢包硬體設備與私鑰保管 SOP。"),
        (16, 2, "專家同儕互評 (Peer Review) 機制：建設性回饋 (Feedback)",
         "講授同儕匿名互評機制：評分表三元素：1. **Keep (優點讚賞)**：指出發表組別表現優異的圖表或財務邏輯；2. **Improve (改善建議)**：溫和指出公式瑕疵或控時細節；3. **Learn (啟發獲益)**：寫下自己獲得的最大靈感。強調同儕學習價值。",
         "同儕互評 (Peer Assessment) 教育學理論。",
         "指導學生進入網頁系統填寫全班 6 大小組同儕互評表。"),
        (17, 2, "教官與專家評審總講評：亮點、邏輯破綻與優化建議",
         "教官與專家評審團總講評：1. **兩大驚艷亮點**：6 組皆成功展現完整的 Excel/Python 財務模型試算，且雙語專有名詞流暢度超越預期；2. **兩大進步提醒**：注意簡報內文文字量，切忌朗讀投影片，Q&A 時善用 Backup Slide 展現數據口袋！",
         "期末專案評審團講評紀錄。",
         "評審團給予全班同學極高的評價與讚賞。"),
        (18, 2, "卓越簡報案例拆解：將枯燥財經數據轉化為動人故事",
         "拆解故事化表達 (Storytelling) 案例：對比「0050 預估年化報酬率 7.2%」與「想像 20 歲的小明，每天省 50 元珍奶錢存 0050，50 歲退休時擁有 447 萬尊嚴退休金！」。展示故事化表達如何瞬間打動評審與聽眾人心。",
         "HBR《Storytelling That Moves People》。",
         "引導學生體會故事化表達在商業 Pitching 中的巨大感染力。"),
        (19, 2, "簡報道具與現場互動：問答設計、QR Code 與多媒體呈現",
         "展示提升發表體驗的現場互動技巧：1. **開場問答**：「現場手上有冷錢包的請舉手！」；2. **試算器 QR Code**：簡報最後一頁提供網頁試算器 QR Code，讓評審現場手機掃描試算；3. **多媒體演示**：流暢切換 10 秒 AI 影片或 HTML 畫布書寫。",
         "TED Talks 簡報視覺設計技巧。",
         "展示在簡報投影片上放 QR Code 供評審即時掃描互動的範例。"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：6 大小組發表亮點；同儕互評三元素 (Keep/Improve/Learn)；專家總講評提醒；財經 Storytelling 故事化表達；現場 QR Code 互動。預告第 2 小時 Modal 實務活動——「期末同儕互評與組別總成績精算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，進行同儕互評與專家分數加權算術演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "打造個人職場簡報亮點：將期末專案納入大一學習歷程 Portfolio",
         "手把手教學將期末專案轉化為履歷表亮點：1. **上傳 GitHub / 個人網站**：將網頁版雙語簡報與 Python 財務模型上傳 GitHub，產出專屬作品連結；2. **履歷亮點撰寫**：「大一完成『全套 18 週財經管理探索實戰專案』，具備 0050/00878 財務模型與 Prompt 工程能力。」",
         "104 人力銀行大一學習歷程與履歷設計。",
         "鼓勵學生：『將這份專案放進履歷，大學四年你將擁有遠超同齡人的作品集 Portfolio！』"),
        (22, 3, "大一新鮮人簡報表達 3 大核心能力：邏輯、視覺與氣場",
         "總結簡報表達 3 大核心能力：1. **金字塔邏輯架構力 (Logic)**：結論先行 (Bottom-line first)，MECE 分類邏輯；2. **數據視覺傳達力 (Visuals)**：莫蘭迪色系、大數字卡片與對比圖表；3. **自信表達氣場 (Aura)**：眼神接觸、控時精準、雙語流暢。",
         "McKinsey Pyramid Principle《金字塔原理》。",
         "強調：『這 3 大能力將跟隨你終身，成為你未來職場升遷與創業簡報的秘密武器！』"),
        (23, 3, "簡報後續追蹤：專案計畫修訂 SOP 與行動方案落地執行",
         "講授簡報後的落地執行 SOP：發表結束不是終點，而是行動計畫落地的起點！1. 依據同儕與專家講評，修訂專案 Excel 財務公式；2. 正式開啟個人 2.5% 高利活存與 0050/00878 定期定額扣款；3. 每半年檢視資產負債表與現金流量表。",
         "PDCA 專案持續改善循環。",
         "提醒：『把簡報裡的財務規劃真正付諸實行，才算真正完成這門課的修練！』"),
        (24, 3, "期末發表資安與智財權保護：避免簡報侵犯著作權",
         "講授發表資安與學術倫理規範：避開簡報侵權陷阱！1. 圖片使用 CC0 免費授權或 AI 生成圖片；2. 引用數據明確標註出處（如：資料來源：證交所/環境部）；3. 未經授權不得直接複製商業網路圖片。",
         "教育部學術倫理與智慧財產權宣導指引。",
         "提醒：『保護智慧財產權，是每一位專業溝通者必須堅守的底線！』"),
        (25, 3, "實證數據調取：查看全班同儕互評數據與期末簡報總分",
         "手把手教導學生線上查看評分結果：1. 檢視財務邏輯 (35%)、痛點解決 (25%)、視覺設計 (20%) 與表達氣場 (20%) 得分分布；2. 查閱同儕給予的「Keep / Improve / Learn」匿名回饋卡片；3. 確立個人成長軌跡。",
         "115管理探索二期末評分系統。",
         "指導學生進入系統讀取同儕給予的建設性留言卡片。"),
        (26, 3, "大一新鮮人實戰財經簡報 4 大金律",
         "總結財經簡報 4 大金律：1. **貫徹 10/20/30 黃金簡報法則**（10 頁、20 分鐘、30 號字）；2. **數據視覺化與大數字卡片**（對比圖表打動評審）；3. **沉著 Q&A 展現數據口袋**（善用 Backup Slide 延伸解答）；4. **專案納入履歷 Portfolio 落地執行**（上傳 GitHub 產出連結）。",
         "麥肯錫與 Guy Kawasaki 簡報保護原則。",
         "請學生齊聲朗讀四金律，建立終身受用的商務簡報習慣。"),
        (27, 3, "簡報試算實例：10 分鐘精準財經 Pitching 流程 SOP",
         "展示 10 分鐘 Pitching 時間分配 SOP：1. 前 2 分鐘 (Hook & Pain)：故事開場，點出痛點；2. 中間 5 分鐘 (Solution & Model)：展示試算模型圖表、0050+00878 配置與 SOP；3. 最後 3 分鐘 (Impact & Summary)：呈現 30 萬資產累積成果，開啟 Q&A！",
         "115管理探索二教案簡報演練組。",
         "強調：『掌握這套 10 分鐘 Pitching SOP，未來面對任何投資人與長官都能游刃有餘！』"),
        (28, 3, "展現學習成果：邁向卓越財經溝通者的總里程碑",
         "總結期末發表的終極價值：期末發表是你邁向商業與管理舞台的起點。經過 17 週的訓練，大家已蛻變為能獨立建構財務模型、進行專業簡報發表的準專業人才！保持表達熱情，將簡報表達與財務分析化為終身競爭力！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『恭喜大家在舞台上展現無與倫比的精彩發表！』"),
        (29, 3, "第十七週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十七週 30 頁純教學卡片進行整體串聯：Hour 1 (專案架構 ➔ 10/20/30法則 ➔ 視覺化大數字 ➔ 演練 4 基本功 ➔ Q&A戰術 ➔ 評分標準 ➔ 第一/二小組發表) ➔ Hour 2 (第三至六小組發表 [台積電/AI/ESG/Web3] ➔ 同儕互評機制 [Keep/Improve/Learn] ➔ 專家總講評 ➔ 財經Storytelling ➔ 現場QR Code互動) ➔ Hour 3 (履歷Portfolio ➔ 3大溝通能力 [邏輯/視覺/氣場] ➔ 後續落地執行 ➔ 智財權保護 ➔ 期末成績統計 ➔ 10分鐘Pitching SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十八週預告 (學期總結、財經探索藍圖與大一終身宣言)",
         "恭喜學生完成第十七週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人期末簡報反思與學習歷程報告」。預告第十八週課程主題：「學期總結、財經探索藍圖與大一新鮮人終身宣言」，下週將進行全學期 18 週總複習、頒發專案獎項與宣言簽署！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第十八週學期總結與終身宣言主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：簡報發表控時與評分標準試算器",
         "1. 活動目標：幫助學生了解商業簡報四大維度（財務邏輯 35%、痛點解決 25%、視覺設計 20%、表達氣場 20%）加權比重，精算專案總成績與等第。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入四大維度得分，點擊精算。\n"
         "3. 診斷反思：系統產出加權總成績，驗證「財務邏輯占 35% 為簡報靈魂」之結論。"),

        ("🎯 第 2 小時實務活動：期末同儕互評與組別總成績精算器",
         "1. 活動目標：讓學生精算專家評審團給分 (60% 權重) 與全班同儕互評平均得分 (40% 權重) 之加權總成績。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入專家評審與同儕平均得分，點擊精算。\n"
         "3. 決策學習：系統產出最終總成績，體現全班高度民主客觀的同儕學習與建設性評閱精神。"),

        ("🎮 第 3 小時小遊戲：財經簡報表達達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：10/20/30 簡報法則中 30 代表 30pt 字體；關卡 2：財務邏輯占 35% 最高權重；關卡 3：Q&A 善用 Backup Slide 防守；關卡 4：同儕互評 Keep/Improve/Learn 3 元素）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 財經簡報表達達人徽章 (Pitch Master)」，未滿分獲頒「🥉 簡報初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人期末簡報反思與學習歷程報告",
         "1. 作業題目：請學生反思小組期末發表表現，撰寫專案納入大一履歷 Portfolio 之修訂計畫。\n"
         "2. 分析要項：(1) 總結 10/20/30 法則與視覺化數據貫徹心得；(2) 整理專家與同儕 Keep/Improve 講評修訂 SOP；(3) 撰寫上傳 GitHub 產出網頁連結納入 Portfolio 計畫；(4) 擬定個人理財 SOP 落地執行細節。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十七週_課程教學指引_期末專案發表實戰財經簡報與總評量.docx'
    doc.save(doc_path)
    print("Created 第十七週_課程教學指引_期末專案發表實戰財經簡報與總評量.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
