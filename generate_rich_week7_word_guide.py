import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第七週 課程教學指引：個人外匯、匯率變動與國際貿易")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第七週完整教學指引。本單元專門針對零財金背景大一新鮮人設計，帶領學生跨出新台幣圈，探索全球貨幣相對價格。課程深入剖析「即期匯率 vs. 現金匯率」與「買入價 vs. 賣出價」之點差算術、旅遊換匯 3 大省錢管道與雙幣信用卡鎖匯機制；解密《經濟學人》大麥克指數 (PPP)、美台利差 (IRP) 與資本流動；客觀分析新台幣升值與貶值對出國留學、iPhone進口與台積電出口獲利的雙刃劍影響，並引述台灣 5,700 億美元外匯存底水庫與央行「柳樹理論」匯率政策。結合中央銀行 (CBC) 與臺灣銀行 (BOT) 權威數據，幫助學生建立獨立的外匯換匯、美金高利定存風險控管與跨國網購資產配置能力。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "外匯基礎與牌告匯率算術\n(Slide 01 - 10)", "即期 vs 現金匯率、銀行買入賣出點差、旅遊換匯 3 大管道、雙幣卡鎖匯與美元指數 DXY 霸權", "🎯 1小時活動：旅遊與網購換匯點差省錢計算器"),
        ("第二小時\n(00:50-01:40)", "匯率理論與升貶值雙刃劍\n(Slide 11 - 20)", "大麥克 PPP、美台利差 IRP、貿易流 vs 資本流、台幣升貶值贏家輸家、5700億外匯存底、柳樹理論與日圓套利", "🎯 2小時活動：大麥克指數 (PPP) 幣值高低估檢測器"),
        ("第三小時\n(01:40-02:30)", "實證比對與新鮮人外匯實戰\n(Slide 21 - 30)", "換匯 4 金律、美金定存高利 5% 匯損風險、台幣 27.5 貶至 32.5 留學 iPhone 對決、10% 美元配置與均值回歸", "🎮 3小時小遊戲：外匯交易達人大挑戰\n📝 課堂實務作業：自選外幣換匯與升貶值影響報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第七週課程導論：個人外匯、匯率變動與國際貿易", 
         "本頁為第七週課程導論。講師可以引用諾貝爾經濟學獎得主薩繆爾森 (Paul Samuelson) 名言開場：「匯率是連繫一國經濟與世界經濟最敏感的神經網路！」針對大一新鮮人出國旅遊換外幣（日圓、美金、韓元）、Amazon/淘寶跨境網購或關注台積電出口新聞切入。說明許多學生不明白為什麼牌告看板會有四個價格，也不明白台幣升值或貶值對自己口袋裡的錢有何實質影響。本單元將帶領大家拆解個人外匯與國際貿易。",
         "中央銀行 (CBC) 外匯局 / 臺灣銀行 (Bank of Taiwan) 牌告匯率發布說明。",
         "問學生：『同學們大學期間有沒有計畫出國旅遊或留學？大家換外幣時最擔心遇到什麼問題？』"),
        (2, 1, "第七週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密外匯基礎、即期 vs 現金匯率點差、銀行買賣視角、旅遊換匯 3 大管道、雙幣卡與美元指數 DXY；第二小時聚焦於購買力平價説 (大麥克指數 PPP)、美台利差 (IRP)、新台幣升貶值贏家與輸家、央行 5,700 億美元外匯存底、柳樹理論與日圓套利；第三小時則傳授換匯 4 金律、美元高利定存匯損風險、台幣 27.5 貶至 32.5 實證對決、10% 美元配置與均值回歸思維。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生劃出本週最感興趣的三大外匯關鍵字（如：即期匯率、大麥克指數、日圓套利）。"),
        (3, 1, "什麼是匯率？(Exchange Rate: 兩種貨幣的相對價格)",
         "詳細拆解匯率 (Exchange Rate) 的定義。匯率是兩種貨幣之間的相對價格。以美金兌新台幣 (USD/TWD = 32.5) 為例，代表買進 1 美元需要支付 32.5 元新台幣。對比新台幣升值 (Appreciation) 與貶值 (Depreciation) 的定義：USD/TWD 從 32.5 降至 30.0 代表台幣變強（升值）；從 32.5 升至 35.0 代表台幣變弱（貶值）。澄清學生常搞混的數字升降與貨幣強弱關係。",
         "中央銀行 (CBC) 外匯局每日匯率發布說明。",
         "提問：『當 USD/TWD 數字從 32 變成 30 時，代表新台幣是變貴了還是變便宜了？（答案：變貴/升值！因為只需 30 元就能換到 1 美元）』"),
        (4, 1, "揭開銀行牌告匯率看板：即期匯率 (Spot) vs. 現金匯率 (Cash)",
         "對比銀行牌告看板上的「即期匯率 (Spot Rate)」與「現金匯率 (Cash Rate)」。即期匯率是電子帳戶數字交割（在 App 上把台幣帳戶扣款轉入外幣帳戶），銀行不需負擔實體現鈔押運、保險與保管成本，因此匯率最優惠；現金匯率是去臨櫃點交實體紙鈔，銀行需加收現鈔處理與資金積壓成本，因此匯率較貴（點差大）。教導大一新鮮人平時先在 App 用即期匯率買好外幣，出國前再領現鈔。",
         "臺灣銀行 (Bank of Taiwan) 歷史牌告匯率說明。",
         "請學生算一算：『如果在 App 用即期匯率換 10 萬日圓比去臨櫃用現金匯率便宜 200 元台幣，這 200 元可以在日本多吃什麼？』"),
        (5, 1, "買入價 (Bank Buy) vs. 賣出價 (Bank Sell) 的銀行視角點差",
         "解決大一新鮮人看匯率看板時最常搞混的難題——永遠從「銀行的角度」看！銀行賣出價 (Bank Sell)：你拿台幣向銀行買外幣（銀行賣給你），價格較貴；銀行買入價 (Bank Buy)：你把外幣賣回給銀行換台幣（銀行向你買），價格較便宜。兩者之間的差額即為「銀行點差利潤 (Spread)」。結論：買了外幣若立刻賣回給銀行，會因為點差的存在立刻產生虧損！",
         "臺灣銀行 (Bank of Taiwan) 牌告匯率買賣價標準定義。",
         "強調：『牌告看板上的【買入】是指銀行向你買，【賣出】是指銀行賣給你，千萬不要搞錯角色！』"),
        (6, 1, "實戰旅遊換匯 3 大管道：臨櫃換鈔 vs. 線上結匯 vs. 外幣 ATM",
         "評比大一出國旅遊換外幣的 3 大實務管道：1. **銀行臨櫃換鈔**：按「現金賣出價」，點差最貴且常需 100 元手續費，適合臨時急用；2. **線上結匯 (指定機場提領)**：在 App/網頁用「即期優惠價」預定並於桃園機場提領，免手續費且最省錢省時；3. **外幣 ATM 提領**：持金融卡至指定外幣 ATM 24 小時提領日圓/美元，直接從台幣帳戶扣款按優惠現鈔價。",
         "兆豐銀行 & 臺灣銀行外匯服務管道比較。",
         "問學生：『如果你明天早上 6 點要搭飛機去日本，最推薦使用哪一種換匯管道提領日圓現鈔？』"),
        (7, 1, "雙幣信用卡 (Dual Currency Card) 的扣款機制與鎖匯優勢",
         "介紹專為旅遊與外匯愛好者設計的「雙幣信用卡」。雙幣卡同時綁定台幣帳戶與指定外幣帳戶（如日圓/美元）。國內消費扣台幣，國外消費直接從外幣帳戶扣除外幣！最大優勢在於「鎖匯」：當你發現日圓處於歷史低點（如 0.20）時，先在 App 上用即期匯率買好日圓存著；半年後去日本刷雙幣卡直接扣除這筆便宜日圓，完全免受屆時匯率暴漲影響。",
         "玉山銀行、國泰世華雙幣信用卡業務公告與鎖匯說明。",
         "提問：『如果你預計半年後要去日本玩，看到現在日圓創歷史新低，你會怎麼利用雙幣卡超前部署？』"),
        (8, 1, "實證數據：全台外幣存款突破 14 兆新台幣",
         "引述中央銀行 2026 最新統計，揭示全台國人外幣存款總額已創下歷史新高，突破 **14.2 兆元新台幣**（佔整體總存款近三成）。解析國人外幣持有結構：第一名美元 (USD, 佔比 > 70%)，主要因全球儲備貨幣地位、高利率美幣定存 (4%~5%) 與美股投資需求；第二名日圓 (JPY, 佔比 ~ 15%)，主要因旅遊需求與便宜日圓逢低囤積；其餘為人民幣、歐元與澳幣。",
         "中央銀行 (CBC) 外匯局 2026 全體金融機構外幣存款統計月報。",
         "提問：『為什麼台灣人手裡握有高達 14 兆元的外幣？大家存美元主要是為了賺利息還是為了出國買東西？』"),
        (9, 1, "美元指數 (DXY) 與全球貨幣的連動王者地位",
         "介紹外匯市場的指標之王——**美元指數 (U.S. Dollar Index, DXY)**。DXY 是衡量美元兌全球 6 大主要貨幣（歐元 57.6%、日圓 13.6%、英鎊 11.9%、加幣 9.1%、瑞典克朗 4.2%、瑞士法郎 3.6%）的綜合加權指數。美元指數走強 (DXY > 105) 代表資金回流美國資產，非美貨幣（台幣、日圓、歐元）普遍面臨貶值壓力；美元指數回落代表資金湧入新興市場。",
         "洲際交易所 (ICE) 美元指數編算規範與歷史走勢。",
         "引導思考：『當新聞報導「美元指數創下 20 年新高」時，你可以預測新台幣兌美元會升值還是貶值？』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：匯率是兩種貨幣相對價格；即期匯率比現金匯率省錢；牌告買賣價需從銀行視角看；線上結匯與雙幣卡是旅遊省錢神力；美元指數 DXY 主導全球資金流向。預告第 1 小時 Modal 實務體驗活動——「旅遊與網購換匯點差省錢計算器」，請學生輸入預計換匯金額，精算即期 vs 現金換匯能省下多少錢。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行即期 vs 現金換匯點差精算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "購買力平價説 (PPP) 與經濟學人「大麥克指數」",
         "詳細講授**購買力平價説 (Purchasing Power Parity, PPP)** 與《經濟學人》創立的「大麥克指數 (Big Mac Index)」。PPP 認為長期而言全世界相同商品經匯率換算後價格應相同（一價定律）。計算算式：`PPP隱含匯率 = 台灣大麥克售價 (75元 TWD) / 美國大麥克售價 (5.6美元 USD) = 13.39 TWD/USD`。若實際市場匯率為 32.5，代表依大麥克購買力計算，新台幣被市場低估了約 58%！",
         "英國《經濟學人》(The Economist) 官方 Big Mac Index 最新報告。",
         "問學生：『如果在台灣買一個大麥克要 75 元，在美國要 5.6 美元，這說明台灣的物價與實質生活成本比美國便宜還是貴？』"),
        (12, 2, "利率平價説 (IRP)：美台利差與資金跨國搬家",
         "剖析**利率平價説 (Interest Rate Parity, IRP)**。說明資金永遠追逐高回報。當美國聯準會升息使美國利率高達 5.25% 而台灣利率僅 2.0% 時（美台利差廣達 3.25%）：外資與台灣投資人會賣掉台幣、換成美元，存入美國高利定存或買美債。市場上大量拋售台幣買進美元，使 USD/TWD 匯率從 30 一路推升至 32.5，形成台幣貶值壓力。",
         "中央銀行 (CBC) 理監事會議美台利差與匯率分析報告。",
         "提問：『如果你手上有 100 萬，放在台灣銀行存定存利息 1.7%，放在美國銀行存定存利息 5.0%，你會想把錢搬去哪裡？』"),
        (13, 2, "影響匯率供需的兩大力量：進出口貿易流 vs. 跨國資本投資流",
         "對比決定一國貨幣升貶值的兩大水龍頭：1. **經常帳貿易流 (Trade Flows)**：台積電等出口商賺取美元外銷財，換回台幣發薪水繳稅，推升台幣升值；2. **金融帳資本流 (Capital Flows)**：外資熱錢買賣台股、本國投信買美債 ETF。強調現代外匯市場中，金融帳資本流的交易數量遠大於實體貿易流，是造成短線匯率劇烈波動的主因。",
         "中央銀行 (CBC) 國際收支帳 (Balance of Payments) 統計數據。",
         "引導思考：『當外資一天在台股賣超 500 億並匯出台灣時，新台幣當天會急升還是急貶？』"),
        (14, 2, "新台幣升值 (Appreciation) 的贏家與輸家分析",
         "深入分析新台幣升值（如 USD/TWD 從 32.5 強升至 28.0）的經濟影響。**升值大贏家**：1. 出國旅遊國人（換外幣超便宜）；2. 進口業者與消費者（進口石油、糧食、iPhone 成本下降，緩解通膨）。**升值大輸家**：出口廠商（如工具機、電子廠），美元營收換回台幣變少，面臨大幅匯兌損失 (FX Loss) 與國際價格競爭力下降。",
         "經濟部國家發展委員會 (NDC) 匯率變動對產業衝擊評估報告。",
         "提問：『當台幣大升值時，去日本玩的人很開心，但新竹科學園區的電子廠老闆為什麼會睡不著覺？』"),
        (15, 2, "新台幣貶值對出口產業 (台積電) 與進口物價的影響",
         "分析新台幣貶值（如 USD/TWD 從 30.0 貶至 33.0）的雙刃劍影響。**出口利多**：台積電晶圓代工以美金計價，當台幣貶值 1 元，台積電營業利益率可提升約 0.4 個百分點，財報獲利亮眼；**輸入型通膨警訊**：台灣能源 97% 依賴進口，台幣貶值導致進口天然氣、原油與小麥價格暴漲，加劇國內物價上漲壓力。",
         "台積電 (TSMC) 法人說明會財務敏感度分析數據。",
         "問學生：『台幣貶值雖然幫台積電多賺了錢，但為什麼一般民眾在夜市買便當或去加油時會覺得錢變小了？』"),
        (16, 2, "央行外匯存底 (FX Reserves) 的功用與台灣 5,700 億美元實力",
         "介紹台灣安然渡過 1997 亞洲金融風暴與 2008 金融海嘯的硬實力底氣——**外匯存底 (Foreign Exchange Reserves)**。外匯存底是中央銀行持有的外幣資產總和（美債、黃金）。台灣外匯存底高達 **5,750 億美元**（全球第六大）。核心功用：1. 清償國際貿易進口帳款；2. 當金融危機外資抽離時，央行能動用外匯存底拋售美元、護盤新台幣。",
         "中央銀行 (CBC) 外匯局 2026 最新外匯存底統計。",
         "提問：『台灣沒有加入聯合國，為什麼能靠著 5,700 億美元的外匯存底在國際金融界屹立不搖？』"),
        (17, 2, "央行匯率政策：柳樹理論 (Willow Tree Policy) 與動態穩定",
         "講述台灣央行最經典的匯率戰略——**「柳樹理論 (Willow Tree Policy)」**。台灣採取「有管理的浮動匯率制度」，匯率應像柳樹枝條一樣，順應國際美元走勢「隨風擺動」，不硬性死盯固定匯率；但當市場出現熱錢暴進暴出、波動過度劇烈時，央行會進場「拉尾盤」熨平波動，維持新台幣動態穩定。",
         "中央銀行 (CBC) 匯率政策與金融穩定報告。",
         "問學生：『為什麼央行希望匯率像柳樹一樣順風擺動，而不是像大樹幹一樣硬碰硬死盯不變？』"),
        (18, 2, "日圓貶值 (JPY) 狂潮：利差交易 (Carry Trade) 運作與旅遊效應",
         "解密日圓兌台幣一度跌破 0.21 的背後機制。日本央行長期維持超低負利率，而美國利率高達 5.25%。全球對沖基金發起**日圓套利交易 (JPY Carry Trade)**：借超便宜日圓 ➔ 換成美元 ➔ 買美債賺 5% 利差。大量拋售日圓導致日圓歷史性暴跌，引發台灣全民赴日旅遊掃貨與買房狂潮。",
         "日本銀行 (BOJ) & 國際清算銀行 (BIS) 跨境套利交易報告。",
         "提問：『當日圓從 0.30 貶到 0.21 時，去日本買一雙 10,000 日圓的球鞋，換算成台幣便宜了多少錢？』"),
        (19, 2, "企業外匯避險 (FX Hedging)：遠期外匯 (Forward Contract) 簡介",
         "介紹出口廠商如何防止匯率吃光利潤。假設鴻海接獲蘋果 1 億美元訂單，約定 6 個月後交貨收款。若鴻海擔心 6 個月後台幣大升值，鴻海會與銀行簽訂**遠期外匯合約 (Forward Contract)**，事先鎖定 6 個月後的履約匯率（如 32.0）。不管半年後市場匯率如何波動，都按 32.0 履約，將匯率風險完全轉移給銀行！",
         "台北外匯市場發展基金會與企業外匯避險實務手冊。",
         "引導思考：『如果鴻海不做了遠期外匯避險，半年後台幣突然升值 10%，鴻海這一單會發生什麼事？』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：PPP 大麥克指數評估貨幣高低估；IRP 美台利差引發資金搬家；台幣升貶值對旅遊進口與出口各有利弊；台灣 5,700 億外匯存底與柳樹理論；遠期外匯鎖匯避險。預告第 2 小時 Modal 實務活動——「大麥克指數 (PPP) 幣值高低估檢測器」，引導學生輸入台美大麥克售價，計算 PPP 隱含匯率。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行大麥克 PPP 幣值高低估檢測演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "大一新鮮人外匯換匯 4 大指標：即期、分批、外幣戶與拒絕 DDC",
         "歸納大一新鮮人聰明換匯四大黃金守則：1. **優先選擇即期匯率**（在 App 用即期價買賣，省下現鈔點差）；2. **分批買進平滑成本**（分 3-4 批逢低買進，平滑波動）；3. **善用外幣數位帳戶**（存美金日圓領高利活存，搭配雙幣卡）；4. **海外結算拒絕 DDC**（國外刷卡一律選當地貨幣，絕不選 TWD 避免 5% 手續費）。",
         "臺灣銀行與兆豐銀行外匯理財實務指南。",
         "提問：『在國外網站刷卡結算時，選 TWD 結算會被扣高達 5% 的 DDC 惡劣手續費，大家記住了嗎？』"),
        (22, 3, "美幣定存 (USD Time Deposit) 高利 5% 的風險：賺利息賠匯差",
         "剖析大一新鮮人最常陷入的「美金高利定存陷阱」。廣告大字「美元定存年利率 5.5%」非常吸引人，但忽略了**匯率風險 (FX Risk)**。算術實例：存入 10,000 美元獲得 5.5% 利息（本利和 10,550 美元），但買進時 USD/TWD = 32.5（本金 32.5 萬台幣），一年後台幣升值至 29.5，換回台幣僅得 31.1 萬元。雖然賺了 5.5% 美元利息，但賠了 8% 匯差，總帳淨虧損 14,000 元台幣！",
         "中央銀行與金管會銀行局外幣定存風險警示手冊。",
         "強調名言：『外幣理財第一鐵律——【賺了利息，也要小心賠了匯差】！』"),
        (23, 3, "實證案例：新台幣從 27.5 貶至 32.5 對留學與 iPhone 的成本計算",
         "透過真實歷史大數據，展示當新台幣兌美元大幅貶值 18%（從 27.5 貶至 32.5）時，對一般民眾購買力的衝擊：1. **美國留學學費 (每年 50,000 美元)**：USD/TWD=27.5 時需 137.5 萬台幣，貶至 32.5 時需 162.5 萬台幣，學費一年暴增 25 萬台幣！2. **進口 iPhone ($1,000 美元)**：售價從 27,500 元暴漲至 32,500 元！",
         "教育部海外留學費用統計與 Apple 台灣官方定價歷史。",
         "請學生算一算：『如果計畫 2 年後去美國留學，現在台幣處於弱勢，你應該怎麼做匯率避險規劃？』"),
        (24, 3, "實證數據調取：中央銀行 (CBC) 外匯局與臺灣銀行每日牌告查閱",
         "手把手教導學生如何查詢權威官方外匯數據：1. **中央銀行外匯局 (cbc.gov.tw)**：查詢每日主要貨幣兌新台幣收盤匯率、外匯存底與交易量；2. **臺灣銀行牌告匯率 (rate.bot.com.tw)**：查詢 19 種外幣的歷史即期與現金買賣匯率走勢圖；3. **經濟部國際貿易署**：查詢台灣進出口統計。",
         "中央銀行 (CBC) 暨臺灣銀行 (BOT) 官方數據門戶。",
         "演示臺灣銀行歷史牌告匯率走勢圖查詢方式。"),
        (25, 3, "大一新鮮人外匯資產配置：10% 美元短期定存 + 旅遊外幣預備金",
         "傳授大一新鮮人理性的外匯資產配比規劃：1. **10% 美元資產**：利用數位帳戶於 USD/TWD < 31.0 時分批買進，存入美金高利活存/定存，作為未來美股或美債 ETF 備用金；2. **旅遊預備金 (小額分批)**：出國前 6 個月逢低分 3 批買進日圓/韓元存入外幣帳戶，辦理雙幣卡方便海外消費。",
         "理財規劃師 (CFP) 外匯資產配置實務建議。",
         "請學生演練：『如果你有 5 萬元存款，請按照 10% 美元 + 旅遊金策略，規劃你的外幣買進計畫。』"),
        (26, 3, "跨境網購 (Amazon, Taobao) 匯率與海外刷卡手續費最佳化",
         "教導學生在 Amazon、Taobao 或 Steam 跨境網購時的精明結算 3 步驟：1. **結算幣別**：Amazon 選美元/日圓；Taobao 選人民幣。絕對不選 TWD！2. **卡片選擇**：挑選海外回饋率 > 2.5% 的信用卡（扣除 1.5% 手續費仍淨賺 1.0%）；3. **關稅注意**：單筆 > 2,000 元或半年超過 6 次將被收取進口關稅與營業稅！",
         "財政部關務署海外網購進口稅費說明。",
         "提醒：『網購填寫報關 EZ WAY 易利委 App 認證時，要留意半年的 6 次免稅額度！』"),
        (27, 3, "避免追高殺低：外匯市場無升不貶、無貶不升的常態輪動",
         "點出外匯市場強烈的**「均值回歸 (Mean Reversion)」特質**。新台幣過去 20 年在 27.5 至 35.0 區間內常態擺動。警示：當 USD/TWD 飆到 33.5、全網新聞熱捧時追高買美元，常買在相對高點，面臨隨後台幣強彈的匯損。理性思維：在台幣強勢 (< 30.0) 時分批儲備美元；在台幣弱勢 (> 33.0) 時分批減碼結算。",
         "中央銀行 20 年新台幣匯率歷史走勢統計。",
         "提醒：『外匯市場沒有永遠上漲的貨幣！切勿在新聞鋪天蓋地報導暴漲時盲目追高。』"),
        (28, 3, "大一新鮮人個人外匯理性使用 4 大金律",
         "總結新鮮人終身受用的個人外匯四大金律：1. **即期分批換匯**（App 線上分 3-4 批買進，平滑點差）；2. **拒絕 DDC 雙幣結算**（海外消費一律選當地貨幣）；3. **警惕美幣高利匯損**（算清匯率風險，不盲追高利）；4. **堅守均值回歸不追高**（台幣強勢時分批儲備，不追高暴漲幣別）。",
         "金管會與中央銀行外匯消費者保護原則。",
         "請學生齊聲朗讀四金律，深化外匯理性投資信念。"),
        (29, 3, "第七週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第七週 30 頁純教學卡片進行整體串聯：Hour 1 (外匯基礎 ➔ 即期vs現金點差 ➔ 銀行買賣視角 ➔ 換匯3管道 ➔ DXY) ➔ Hour 2 (PPP大麥克指數 ➔ IRP美台利差 ➔ 貿易流vs資本流 ➔ 台幣升貶值對決 ➔ 外匯存底 ➔ 柳樹理論 ➔ 遠期避險) ➔ Hour 3 (換匯4金律 ➔ 美金定存匯損 ➔ 台幣27.5貶至32.5實證 ➔ 10%美元配置 ➔ DDC避坑 ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第八週預告 (期中學習成果檢核與實戰複習)",
         "恭喜學生完成第七週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「自選外幣換匯與升貶值影響報告」。預告第八週課程主題：「期中學習成果檢核與實戰個案總複習」，下週將針對前七週（財經素養、通膨、利率、股市估值、ETF、數位金融、外匯）進行綜合實戰個案總複習與期中會考挑戰！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並積極複習前七週全套課程準備期中會考。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：即期 vs. 現金換匯點差省錢計算器",
         "1. 活動目標：讓學生親自計算即期匯率與現金匯率的價差金額，養成線上換匯與機場提領的省錢習慣。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入預計換匯之台幣金額（例如 30,000 元）、App 即期賣出價（例如 0.2150）與臨櫃現金賣出價（例如 0.2200），點擊試算。\n"
         "3. 診斷反思：引導學生計算即期換匯多換得的日圓金額與折合省下的台幣，體會看懂牌告點差實質省錢效益。"),

        ("🎯 第 2 小時實務活動：大麥克指數 (PPP) 幣值高低估檢測器",
         "1. 活動目標：幫助學生理解購買力平價説 (PPP) 的一價定律，學會檢測各國貨幣是否被市場高估或低估。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入台灣大麥克售價（75元）、美國大麥克售價（5.6美元）與當前市場匯率（32.5），系統自動算出 PPP 隱含匯率 (13.39)。\n"
         "3. 決策學習：比對市場匯率與 PPP 隱含匯率，判定新台幣被低估約 58%，討論實體購買力與金融帳熱錢流動對匯率的拉扯。"),

        ("🎮 第 3 小時小遊戲：外匯交易達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：即期匯率 vs 現金匯率；關卡 2：銀行賣出價 vs 買入價視角；關卡 3：新台幣貶值的經濟影響；關卡 4：海外結算 DDC 避坑）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 外匯交易達人徽章 (Forex Master)」，未滿分獲頒「🥉 外匯初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：自選外幣換匯與升貶值影響報告",
         "1. 作業題目：請學生自行挑選 1 種計畫持有或旅遊的外幣（如美元 USD、日圓 JPY），線上填寫報告。\n"
         "2. 分析要項：(1) 記錄臺灣銀行當前即期與現金賣出價並比較點差；(2) 分析新台幣對該外幣大升值 10% 對個人旅遊與出口企業之優缺點；(3) 擬定個人分批換匯與外幣配置規劃。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第七週_課程教學指引_個人外匯匯率變動與國際貿易.docx'
    doc.save(doc_path)
    print("Created 第七週_課程教學指引_個人外匯匯率變動與國際貿易.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
