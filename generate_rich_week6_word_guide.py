import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第六週 課程教學指引：數位金融、高利活存與銀行業變革")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第六週完整教學指引。本單元專門針對大一新鮮人進入大學後第一件金融大事「申辦數位帳戶與第一張信用卡/簽帳卡」設計。課程深入剖析商業銀行存放款淨利差 (NIM) 商業模式、純網銀 (LINE Bank、樂天、將來銀行) 高利活存限額算術、海外刷卡 1.5% 手續費與 DDC 陷阱、分期 0% 與 15% 循環信用利息危險，以及 2FA 兩步驟驗證、防範釣魚簡訊與洗錢防制法人頭帳戶警告。結合金管會銀行局與中央銀行 (CBC) 權威數據，建立新鮮人終身受用的數位金融安全與理財習慣。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "商業銀行本質與數位帳戶\n(Slide 01 - 10)", "存放款淨利差 (NIM)、信用創造乘數、全台 2,300萬數位帳戶、純網銀三雄比對與存保 300萬保障", "🎯 1小時活動：數位帳戶高利活存與利息試算盤"),
        ("第二小時\n(00:50-01:40)", "支付革命與點數經濟陷阱\n(Slide 11 - 20)", "支付演進、信用卡生態系、海外 1.5% 手續費算術、180天點數過期、分期 0% 與 15% 循環利息", "🎯 2小時活動：信用卡現金回饋與海外手續費計算器"),
        ("第三小時\n(01:40-02:30)", "資安防護與新鮮人理財實戰\n(Slide 21 - 30)", "2FA 雙重驗證、釣魚防詐、人頭帳戶警告、大一 1薪轉+2數位戶開戶配置與 JCIC 聯徵分數", "🎮 3小時小遊戲：數位金融達人大挑戰\n📝 課堂實務作業：數位帳戶權益比較報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第六週課程導論：數位金融、高利活存與銀行業變革", 
         "本頁為第六週課程導論。講師可以引用比爾·蓋茲 (Bill Gates) 經典名言開場：「銀行服務是必要的，但實體銀行不是！(Banking is necessary, banks are not!)」針對大一新鮮人剛滿 18 歲踏入大學、面臨人生第一次開戶與申辦信用卡/簽帳卡的興奮感切入。許多學生被廣告上的「高利活存 8%」、「刷卡回饋 10%」大字吸引，但往往不了解銀行背後的商業模式與行銷算術。本單元將帶領大家揭開數位金融與商業銀行的神秘面紗。",
         "金融監督管理委員會 (FSC) 銀行局 / 中央銀行 (CBC) 官方資料。",
         "問學生：『同學們上大學後，有沒有去實體銀行排隊開過戶？排了多久？跟手機開戶比起來體驗如何？』"),
        (2, 1, "第六週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密商業銀行本質、存放款淨利差 (NIM)、信用創造機制、純網銀三雄與 CDIC 300萬存款保障；第二小時聚焦於支付演進、信用卡生態系、海外 1.5% 手續費算術、180天點數過期與 15% 循環利息黑洞；第三小時則傳授 2FA 雙重驗證資安防護、釣魚簡訊辨識、人頭帳戶洗錢警告，以及 1 薪轉 + 2 數位戶的理財配置與 JCIC 聯徵信用分數建立。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上圈出自己目前最想申辦的數位帳戶名稱，帶著問題進入本週課程。"),
        (3, 1, "商業銀行的核心商業模式：存款、放款與淨利差 (NIM)",
         "詳細講解商業銀行 (Commercial Banks) 作為「金融中介 (Financial Intermediary)」的核心盈利邏輯。銀行透過吸收社會大眾的存款並支付低利息（如活存 0.8% 或定存 1.7%），再將資金貸放給購屋族或企業並收取高利息（如房貸 2.2%、車貸 4.0%、信貸 6.0%）。兩者之間的差額即為**淨利差 (Net Interest Margin, NIM)**。說明 NIM 是商業銀行最主要的獲利來源。",
         "中央銀行 (CBC) 全體銀行存放款加權平均利率統計月報。",
         "提問：『如果你把 10 萬元存進銀行領 0.8% 利息，銀行拿這筆錢借給別人收 6% 信貸利息，銀行這一趟賺了多少錢？』"),
        (4, 1, "信用創造機制：銀行如何「憑空製造貨幣」？",
         "解密金融體系中最神奇的「部分準備金制度 (Fractional Reserve Banking)」。以央行法定準備率 10% 為例：小明存入 10,000 元，銀行保留 1,000 元準備金，將 9,000 元貸給小華；小華買手機後店家存入銀行，銀行再保留 900 元準備金，將 8,100 元貸給小強... 透過貨幣乘數公式 `貨幣乘數 = 1 / 法定準備率 = 10 倍`，最初的 10,000 元現金最終在社會上創造出高達 100,000 元的廣義貨幣 (M2)！",
         "中央銀行 (CBC) 貨幣供給量 M1A, M1B, M2 統計數據說明。",
         "引導思考：『如果全台灣所有人同一天跑去銀行把存款全部領出來，銀行給得出錢嗎？這就是擠兌 (Bank Run) 的由來。』"),
        (5, 1, "傳統銀行實體分行 vs. 數位存款帳戶 (Digital Savings)",
         "對比傳統實體分行與數位存款帳戶的營運成本結構。實體銀行負擔昂貴的黃金地段租金、水電保全與大量櫃員薪資，因此僅能給予 0.8% 基本活存利率；而數位帳戶 / 純網銀「無實體店面」，營運成本極低，能將省下來的數億元租金與人事成本，轉化為 2.0%~8.0% 的高利活存回饋給年輕數位原生代顧客，實現雙贏。",
         "金融監督管理委員會 (FSC) 銀行局數位存款帳戶業務規範。",
         "提問：『如果你是銀行總經理，省下一家分行每年 1,000 萬的租金，你會拿來做什麼行銷推廣？』"),
        (6, 1, "實證數據：全台數位存款帳戶突破 2,300 萬戶",
         "引述金管會 2026 最新官方統計，揭示全台數位存款帳戶開戶數已強勢突破 **2,350 萬戶**（平均每人擁有 1 戶以上）。將數位帳戶分為兩大陣營：1. **第一類 (傳統銀行數位戶)**：如台新 Richart、國泰世華 CUBE、永豐 DAWHO，戶數超過 1,800 萬戶；2. **第二類 (純網路銀行)**：如 LINE Bank、樂天、將來銀行，戶數超過 380 萬戶。顯示數位金融已成為大學生日常生活的一部分。",
         "金融監督管理委員會 (FSC) 銀行局 2026 最新數位存款帳戶統計報告。",
         "請班上同學舉手統計：『手上有台新 Richart 或 LINE Bank 的同學請舉手？看看哪一家在大學生中最普及？』"),
        (7, 1, "純網銀 (Neobanks) 革命：LINE Bank、樂天、將來銀行三雄比對",
         "詳細比對台灣經金管會特許發行的三家純網路銀行 (Neobanks)：1. **LINE Bank**：結合 LINE 2,100 萬用戶生態圈，主打「聊天室直接轉帳免記帳號」與 LINE Points 即時回饋；2. **樂天銀行**：結合日本樂天集團與國票金控，主打樂天點數與赴日旅遊跨境優惠；3. **將來銀行**：結合中華電信與兆豐銀行，提供口袋帳戶自由劃分與高額免費跨轉。",
         "各純網銀 2026 官方公開財務報告與權益公告。",
         "提問：『跟朋友吃完飯要分帳時，用 LINE Bank 聊天室直接轉帳跟輸入 16 位數銀行帳號相比，哪一個比較不容易出錯？』"),
        (8, 1, "拆解高利活存真相：牌告利率 vs. 優惠利率限額算術",
         "揭露高利活存行銷大字背後的算術陷阱。大一新鮮人常被「高利活存 8%」大字吸引，卻忽視細則中的**「金額上限 (Cap)」與「指定任務」**。算術實例：某銀行宣稱 8% 活存，但限額僅 **10,000 元**，超過部分回歸 0.8% 基本利率。存入 10 萬元，一個月領到的利息僅為 `(10,000 * 8% / 12) + (90,000 * 0.8% / 12) = 66 + 60 = 126 元`！實質有效利率僅 1.51%。",
         "中央銀行與公平交易委員會 (FTC) 不實廣告與消費警示範例。",
         "提醒：『看高利活存不能只看 % 數，一定要先看「優惠金額上限」是多少！』"),
        (9, 1, "中央存款保險公司 (CDIC) 300 萬元保障機制",
         "解答學生常見疑慮：「純網銀沒有實體分行，萬一倒閉我的錢會不會消失？」說明不論傳統銀行或純網銀，只要經金管會核准並參加存保，均享有**中央存款保險公司 (CDIC)** 的法定保障。每一存款人在同一家銀行的本金加利息，最高保障額度為 **新台幣 300 萬元**。且保費由銀行繳納，存款人完全免費！",
         "中央存款保險公司 (CDIC) 存款保險條例與保障額度規定。",
         "問學生：『如果你有 500 萬元現金，為了得到 100% 的存款保險保障，你應該怎麼存放？（答案：分存在 2 家不同的存保銀行）』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：商業銀行靠存放款利差 (NIM) 盈利；信用創造放大貨幣供給；數位帳戶靠省租金提供高利活存；精算優惠金額上限；認明 CDIC 300萬存款保障。預告第 1 小時 Modal 實務體驗活動——「數位帳戶高利活存與利息精算盤」，引導學生輸入本金、優惠利率與上限，計算實質收益。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行高利活存利息精算演練。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "支付方式大演進：現金 ➔ 塑膠卡片 ➔ 行動支付 (Apple Pay, TWQR)",
         "梳理人類支付方式的三代演進：1. **實體現金 (Cash)**：匿名性高但無回饋且遺失無法找回；2. **塑膠卡片 (Cards)**：簽帳卡/信用卡，提供現金回饋與延後付款；3. **行動支付 (Mobile)**：分為 NFC 感應（Apple Pay、Google Pay）與 QR Code 掃碼（LINE Pay、街口、全支付、TWQR）。說明國家推動 TWQR 電子支付跨機構共用碼的戰略意義。",
         "財金資訊公司 (FISC) 全民通用支付 TWQR 規範說明。",
         "問學生：『大家上大學後，一整天出門消費完全不拿出實體新台幣硬幣與鈔票的比例有多高？』"),
        (12, 2, "信用卡交易四角架構：發卡行、收單行、清算機構與手續費",
         "詳細拆解信用卡刷卡幕後的四角架構：1. **發卡銀行 (Issuer)**：核發卡片給持卡人，承擔信用風險；2. **收單銀行 (Acquirer)**：為店家安裝刷卡機並收取 1.5%~2.5% 手續費；3. **清算組織 (Networks)**：Visa、Mastercard、JCB 轉接訊息；4. **商家**：負擔刷卡手續費。說明商家願意負擔手續費是因為能帶來更高的客單價與免收偽鈔優勢。",
         "Visa & Mastercard 國際組織刷卡手續費拆解規範。",
         "提問：『為什麼有些夜市小吃攤或早餐店不願意支援信用卡刷卡？（答案：不想負擔 2% 的刷卡手續費與稅務留痕）』"),
        (13, 2, "信用卡現金回饋算術：國內 2% / 國外 3% 與海外交易手續費",
         "詳細剖析海外刷卡手續費算術。當你在國外網站（如 Steam、Amazon、Netflix）或出國旅遊刷卡時，發卡行會收取 **1.5% 海外交易國際手續費**（1.0% 清算組織費 + 0.5% 發卡行處理費）。實質淨回饋算式：`海外淨回饋率 = 宣稱回饋率 (如 3.0%) - 法定手續費 (1.5%) = 淨賺 1.5%`。同時警示 **DDC 雙幣轉換陷阱**：外國結算千萬別選新台幣，否則會被收取 4%~5% 惡劣手續費！",
         "金管會銀行局信用卡海外交易手續費揭露規範範本。",
         "提醒學生：『在日本或美國網站刷卡結算時，看到畫面問你要用 TWD 還是當地貨幣，一律選擇當地貨幣！』"),
        (14, 2, "點數經濟 (Points Economy)：LINE Points、街口幣與過期陷阱",
         "剖析發卡行最愛的「點數經濟 (Points Economy)」。許多信用卡不發放現金，改發放 LINE Points、街口幣、玉山 Pi 幣等自營點數（1 點 = 1 元）。點數好處是折抵零門檻，但隱藏大陷阱是**「效期限制」**（部分點數僅有 180 天效期）。若忘記使用，點數會在深夜自動歸零，形成銀行的「過期回沖收益 (Breakage)」。",
         "中華民國消費者文教基金會 (消基會) 點數效應調查報告。",
         "提問：『同學們有沒有過 LINE Points 或超商點數放到過期被扣掉的痛苦經驗？當時虧了多少錢？』"),
        (15, 2, "分期付款 (Installments) 的心理誘惑與實質利率算術",
         "揭露「單筆滿額 0 利率分期」的心理學與財務陷阱。心理學證明分期付款會模糊消費者的消費痛感，讓人誤以為自己買得起超越負擔能力的昂貴商品。隱藏代價：1. 許多信用卡規定分期消費不再享有現金回饋；2. 若忘了按時繳納，0% 利率會瞬間終止並啟動高達 **15% 的違約金與循環利息！**",
         "金融消費評議中心 (FOI) 信用卡分期糾紛裁決案例彙編。",
         "警示學生：『買 30,000 元手機分 12 期看似月付 2,500 很輕鬆，但如果你同時分期了手機、機車、衣服，每個月固定扣款就會壓垮你的生活費！』"),
        (16, 2, "循環信用利息 (Revolving Credit Interest) 的高利黑洞",
         "警示大一新鮮人絕對不能碰的卡奴黑洞——**循環信用利息**。當月帳單 20,000 元，若只繳納「最低應繳金額」（如 2,000 元），剩下的 18,000 元將啟動 **6%~15% 的年化循環利息**。可怕真相：計息起算日是往前溯及自「商家刷卡入帳日當天」開始按日計息！而且當月所有新刷卡消費都失去免息期。連續只繳最低應繳金額將導致債台高築。",
         "金管會銀行局信用卡循環信用利息計息標準規範。",
         "要求學生記住鐵律：『信用卡帳單來了，永遠只有一個選擇——【全額繳清】！』"),
        (17, 2, "開放銀行 (Open Banking) API 與三階段架構",
         "介紹金管會推動的「開放銀行 (Open Banking)」戰略。過去金融數據被鎖在各家銀行封閉資料庫，開放銀行透過開放 API，在客戶授權下讓第三方服務業者 (TSP) 整合數據：第一階段（公開商品查詢）、第二階段（客戶資訊提供，如在麻布記帳跨行查資產）、第三階段（交易資訊整合，直接跨行扣款）。讓金融數據所有權回歸消費者。",
         "財金資訊公司 (FISC) 開放 API 平台運作規範與技術標準。",
         "問學生：『如果你可以用一個 App 一次看清你在 5 家銀行的存款總額，這對你管理個人資產有何幫助？』"),
        (18, 2, "中央銀行數位貨幣 (CBDC) 的發展與實體現金未來",
         "對比中央銀行數位貨幣 (CBDC) 與去中心化加密貨幣（如比特幣、USDT）。CBDC 由國家央行發行並背書，具備 100% 法償效力與零信用風險，完全等同實體新台幣；而加密貨幣無國家背書且價格劇烈波動。介紹台灣央行正在進行的 CBDC 零售端試驗與未來離線感應支付發展。",
         "中央銀行 (CBC) 數位新台幣 (CBDC) 研究報告與試驗進度。",
         "引導思考：『如果未來全面使用央行數位新台幣，壓歲錢與紅包改成手機轉帳，實體紙鈔會不會徹底消失？』"),
        (19, 2, "嵌入式金融 (Embedded Finance)：先買後付 Buy Now Pay Later (BNPL)",
         "剖析針對無卡族與學生的「先買後付 (Buy Now Pay Later, BNPL，如 AFTEE、慢點付)」。BNPL 讓學生在電商結帳時只需手機號碼認證即可無卡分期。優點是免信用卡、免信用審查；缺點是缺乏聯徵監管，容易讓人不知不覺過度擴張信用，且逾期會被收取極高違約金與遲延利息！",
         "金融監督管理委員會 (FSC) BNPL 業務自律規範與消費警示。",
         "提醒：『BNPL 不是免費的禮物，本質上就是一種高風險的短期小額貸款！』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：海外刷卡扣除 1.5% 手續費才是淨回饋；注意 180 天點數過期；分期 0% 可能取消回饋且逾期按 15% 循環計息；開放銀行、CBDC 與 BNPL 的理性防範。預告第 2 小時 Modal 實務活動——「信用卡現金回饋與海外手續費計算器」，引導學生精算刷卡實質收益。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行信用卡刷卡回饋精算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "數位金融安全的防禦第一線：兩步驟驗證 (2FA) 與生物辨識",
         "講授數位金融資安的三大核心要素：1. **知識因素 (Knowledge)**：密碼、PIN碼（切勿用生日或123456）；2. **持有因素 (Possession)**：綁定手機、OTP 簡訊驗證碼；3. **生物特徵 (Inherence)**：Face ID 人臉辨識、指紋辨識。強調**兩步驟驗證 (2FA)** 的重要性，並立下鐵律：絕不透露 6 位數 OTP 驗證碼給任何人！",
         "內政部警政署 165 反詐騙諮詢專線資安指引。",
         "問學生：『如果自稱銀行客服打電話來，說系統故障要你念出簡訊裡的 6 位數驗證碼，你應該怎麼辦？（答案：立刻掛斷並撥打 165）』"),
        (22, 3, "全面防範簡訊與通訊軟體釣魚詐騙 (Phishing Attacks)",
         "剖析詐騙集團最常使用的「釣魚簡訊 (Phishing SMS)」手法。詐騙集團發送簡訊：「【國泰/富邦】您的帳戶資安異常，請點擊短網址更新密碼否則凍結！」點進去是複製得一模一樣的假銀行網站，一旦輸入帳密與 OTP，數十萬存款秒被轉走。防禦解法：認明奇怪域名（如 `.xyz`），永遠透過官方 App 登入，絕不點擊簡訊網址。",
         "國家通訊傳播委員會 (NCC) 暨 165 反詐騙資料庫。",
         "教導學生辨識網址：『真正的銀行官網必定是 `.com.tw` 結尾，任何帶有奇怪破折號或外國域名的網址都是詐騙！』"),
        (23, 3, "人頭帳戶與洗錢防制法：大一新鮮人絕對不能租借帳戶",
         "警示大一新鮮人絕對不能碰的法律紅線——**人頭帳戶**。網路常有徵才廣告求租/買金融帳戶（提款卡+密碼，宣稱一個月給 1 萬）。出借帳戶將直接觸犯《洗錢防制法》與《刑法》幫助詐欺罪！後果：全台所有銀行帳戶立刻被列為「警示帳戶」凍結管制，無法轉帳領薪水，面臨最重 3 年以上有期徒刑與民事賠償，信用破產！",
         "法務部與最高檢察署洗錢防制法懲處案例宣導。",
         "警告學生：『天下沒有白吃的午餐！把提款卡借給別人，你就是詐騙集團的替死鬼與洗錢共犯！』"),
        (24, 3, "實證數據調取：金管會銀行局與中央銀行利率查詢",
         "手把手教導學生如何查詢權威官方金融數據：1. **金管會銀行局 (bank.gov.tw)**：查詢全台數位帳戶統計與銀行裁罰紀錄；2. **中央銀行 (cbc.gov.tw)**：查詢基準利率與全體銀行存放款平均利率；3. **中央存款保險公司 (cdic.gov.tw)**：查詢特定機構是否具備 300 萬合法存款保險。",
         "金融監督管理委員會 (FSC) 暨中央銀行 (CBC) 官方門戶。",
         "演示官方查詢流程，要求學生記錄權威機構網址。"),
        (25, 3, "大一新鮮人數位帳戶組合策略：1 薪轉 + 2 高利活存",
         "傳授大一新鮮人最佳數位帳戶組合策略——**「1 薪轉 + 2 高利活存」**：1. **1 戶薪轉/日常帳戶**：選擇國泰、中信等大型實體銀行數位戶，用於打工薪轉、繳學雜費與扣卡費；2. **2 戶高利活存戶**：選擇 LINE Bank、Richart 等純網銀/數位戶，專款專用儲蓄緊急預備金並賺取 2%~3% 高利活存。",
         "理財規劃師 (CFP) 新鮮人資產配置實務原則。",
         "請學生隨堂演練：『劃分你每月的打工收入與零用錢，說明你打算如何分配在 3 個帳戶中。』"),
        (26, 3, "建立個人信用聯徵紀錄 (JCIC)：信用卡全額繳清的重要性",
         "介紹金融界的「信用護照」——**財團法人金融聯合徵信中心 (JCIC)**。銀行決定未來要不要借錢給你買房買車，全看聯徵信用分數 (300~800 分)。加分項目：信用卡每月全額按時繳清、長期穩定往來；扣分黑名單：動用循環利息、預借現金、延遲繳款、短期內頻繁向多家銀行送件辦卡（聯徵多次查詢）。",
         "財團法人金融聯合徵信中心 (JCIC) 信用評分說明冊。",
         "強調：『從大一辦第一張卡開始，每一次按時全額繳清，都是在為你未來買房子的低利房貸積攢信用資產！』"),
        (27, 3, "自動扣繳與免費跨轉/跨提次數極致利用技巧",
         "傳授新鮮人懶人理財與省錢技巧。數位帳戶通常提供每月 10~99 次免費跨轉跨提優惠。技巧：1. **發薪日自動劃轉**：設定發薪日隔天自動將 3,000 元轉入高利活存戶（先儲蓄後消費）；2. **善用免費轉帳**：利用 LINE Bank 或 Richart 免費轉帳分帳給朋友，一年省下數百元跨轉手續費（一次 15 元）。",
         "各數位存款帳戶免費跨轉跨提次數統計。",
         "算一算：『如果一個月跨行轉帳 10 次，一年下來光手續費就能省下 `10 * 15 * 12 = 1,800 元`，等於免費吃兩大頓大餐！』"),
        (28, 3, "大一新鮮人數位金融理性使用 4 大金律",
         "總結新鮮人終身受用的數位金融四大金律：1. **信用卡全額按時繳清**（絕不按最低應繳，遠離 15% 循環利息）；2. **開啟 2FA 絕不透密碼**（保護 OTP 驗證碼，不點擊簡訊網址）；3. **精算高利活存額度**（看清優惠上限，不過度消費）；4. **絕不出租出借金融帳戶**（拒絕高價求租，保護清白信用）。",
         "金管會銀行局 (FSC) 消費者金融保護原則。",
         "請學生齊聲朗讀四金律，深化數位金融安全信念。"),
        (29, 3, "第六週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第六週 30 頁純教學卡片進行整體串聯：Hour 1 (存放款利差 NIM ➔ 信用創造 ➔ 數位帳戶 ➔ 高利活存上限 ➔ CDIC 300萬存保) ➔ Hour 2 (支付演進 ➔ 四角架構 ➔ 海外 1.5% 手續費 ➔ 點數過期 ➔ 15% 循環利息 ➔ 開放銀行 & CBDC) ➔ Hour 3 (2FA & OTP ➔ 防釣魚 ➔ 人頭帳戶警告 ➔ 1薪轉+2數位戶 ➔ JCIC 分數 ➔ 4 大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第七週預告 (個人外匯、匯率變動與國際貿易)",
         "恭喜學生完成第六週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「數位帳戶權益比較與防詐防護報告」。預告第七週課程主題：「個人外匯、匯率變動與國際貿易」，下週將帶大家探索美金與日圓換匯最佳時機、即期 vs 現金匯率算術與護國神山出口升貶值影響。",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並預習第七週個人外匯主題。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：數位帳戶高利活存與利息精算盤",
         "1. 活動目標：讓學生親自精算高利活存額度上限 (Cap) 對實質年化利率的稀釋影響，避免盲目被行銷大字誤導。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入預計存款金額（例如 50,000 元）、宣稱優惠利率（例如 3.5%）、上限金額（例如 20,000 元）與基本利率（例如 0.8%），點擊試算。\n"
         "3. 診斷反思：引導學生觀察超過上限部分稀釋後的「實質綜合有效利率 (Effective Rate)」，教育學生如何依據資金規模合理分配在不同數位帳戶中。"),

        ("🎯 第 2 小時實務活動：信用卡現金回饋與海外手續費計算器",
         "1. 活動目標：幫助學生掌握海外刷卡 1.5% 手續費與回饋率的算術抵扣，學會精算實質淨收益與拒絕 DDC 雙幣轉換。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入刷卡金額、宣稱回饋率（例如 3.0%），並勾選是否為海外交易，系統自動計算扣除 1.5% 手續費後之淨獲利。\n"
         "3. 決策學習：當淨回饋率 > 0% 時判定划算，反之提示手續費吃掉回饋，培養海外消費精打細算的理性習慣。"),

        ("🎮 第 3 小時小遊戲：數位金融達人大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：存款保險保障機制；關卡 2：海外刷卡手續費算術；關卡 3：信用卡循環利息起算點；關卡 4：數位資安與人頭帳戶防禦）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 數位金融達人徽章 (Fintech Expert)」，未滿分獲頒「🥉 數位金融初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：數位帳戶權益比較與防詐防護報告",
         "1. 作業題目：請學生自行挑選 2 家數位帳戶/純網銀（如 LINE Bank vs Richart），線上填寫寫比較報告。\n"
         "2. 分析要項：(1) 比較兩者之活存利率、額度上限與免費跨轉次數；(2) 設定個人大一新鮮人「1薪轉 + 2數位戶」資產劃分規劃；(3) 擬定個人 2FA 資安防護與 165 防詐 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第六週_課程教學指引_數位金融高利活存與銀行業變革.docx'
    doc.save(doc_path)
    print("Created 第六週_課程教學指引_數位金融高利活存與銀行業變革.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
