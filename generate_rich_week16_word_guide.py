import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rich_word_guide():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft JhengHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Color Palette
    PRIMARY_NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("115 學年度 大一通識核心課程「管理探索二」")
    run_title.font.name = 'Microsoft JhengHei'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_BLUE

    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_title_p.add_run("第十六週 課程教學指引：加密貨幣、區塊鏈基礎、Web3 與數位資產風險")
    run_sub.font.name = 'Microsoft JhengHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY_NAVY

    # --- PART 1 ---
    h1 = doc.add_heading("一、課程基本資訊與學習架構總覽", level=1)
    h1.runs[0].font.name = 'Microsoft JhengHei'
    h1.runs[0].font.size = Pt(15)
    h1.runs[0].font.color.rgb = PRIMARY_NAVY
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    p1.paragraph_format.space_after = Pt(10)
    p1.add_run("本課程為 115 學年度大一通識核心課程「管理探索二」（財金蛙挖哇）第十六週完整教學指引。本單元帶領大一新鮮人解密數位新金融前沿「加密貨幣、區塊鏈與 Web3 資產風險」。課程深入剖析區塊鏈 (Blockchain) 去中心化分散式帳本技術；解密比特幣 (BTC) 創世白皮書與 2,100 萬枚總量上限抗通膨原理、4 年減半機制 (Halving) 供給衝擊；講授以太坊 (Ethereum) 智慧合約與 DApp 應用；剖析 Web1 唯讀 ➔ Web2 讀寫 ➔ Web3 擁有的文明演進；傳授冷錢包 (Cold Wallet) 私鑰 (Private Key) 保管鐵律「Not Your Keys, Not Your Coins」；解密 DeFi 流動性挖礦、NFT 數位憑證、USDT/USDC 穩定幣與 2022 FTX 倒閉歷史教訓；解讀台灣 VASP 洗錢防制監管與銀行信託；傳授大一小資族 5% 避險配置戰術，並提供 Telegram 幣圈防詐與合法交易所 KYC 出入金 SOP，輔以 CoinMarketCap 與金管會 VASP 洗錢防制聲明實證數據，引領學生建立理性的 Web3 數位金融資安素養。")

    # --- PART 2 ---
    h2 = doc.add_heading("二、3小時教學進度與活動配比表", level=1)
    h2.runs[0].font.name = 'Microsoft JhengHei'
    h2.runs[0].font.size = Pt(15)
    h2.runs[0].font.color.rgb = PRIMARY_NAVY

    table_schedule = doc.add_table(rows=4, cols=4)
    table_schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_schedule.rows[0].cells
    headers = ["授課時段", "核心單元主題與 30 頁純教學卡片", "講授重點與大一生活化生活實例", "專屬實務活動與互動 Modal"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    schedule_data = [
        ("第一小時\n(00:00-00:50)", "區塊鏈、比特幣減半與 Web3\n(Slide 01 - 10)", "區塊鏈定義、比特幣2100萬枚上限、BTC減半經濟學、以太坊智慧合約、Web3演進、冷熱錢包私鑰保管", "🎯 1小時活動：比特幣減半產量與冷熱錢包資安評估計算器"),
        ("第二小時\n(00:50-01:40)", "DeFi、穩定幣與 VASP 監管\n(Slide 11 - 20)", "DeFi去中心化金融、NFT數位資產、USDT/USDC穩定幣、FTX倒閉危機、台灣VASP監管與比特幣現貨ETF", "🎯 2小時活動：DeFi 年化收益風險與數位資產配置試算器"),
        ("第三小時\n(01:40-02:30)", "防詐 SOP、數位資產配置與 4 大金律\n(Slide 21 - 30)", "Telegram幣圈防詐、Web3三大素養、RWA資產代幣化、合法交易所KYC出入金、CoinGecko查閱與4大金律", "🎮 3小時小遊戲：Web3 數位資產防詐大挑戰 (4大關卡)\n📝 課堂實務作業：個人數位資產風險評估與 Web3 報告")
    ]

    for row_idx, data in enumerate(schedule_data, start=1):
        row_cells = table_schedule.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- PART 3 ---
    h3 = doc.add_heading("三、30頁純教學卡片詳細教案與講授指引", level=1)
    h3.runs[0].font.name = 'Microsoft JhengHei'
    h3.runs[0].font.size = Pt(15)
    h3.runs[0].font.color.rgb = PRIMARY_NAVY

    p_info = doc.add_paragraph()
    p_info.add_run("本部分提供 30 頁純教學卡片之逐頁詳細教案，每頁包含【講授核心與生活化說明】(150+字深入闡述)、【官方數據引述】與【課堂提問與引導技巧】，確保授課教師能精準掌握通識教學節奏。")
    p_info.paragraph_format.space_after = Pt(14)

    # 30 Detailed Slide Guides
    slides_breakdown = [
        # Hour 1 (Slide 01-10)
        (1, 1, "第十六週課程導論：加密貨幣、區塊鏈基礎、Web3 與數位資產風險", 
         "本頁為第十六週課程導論。講師可以引用中本聰名言開場：「密碼學貨幣的核心優勢，在於它不需要信任第三方仲介！」針對大一新鮮人身處數位原生時代的衝擊切入。說明加密貨幣已從非主流密碼學實驗，躍升為全球金融機構關注的數位資產類別。本單元將帶領大家拆解區塊鏈原理、比特幣 2,100 萬枚上限、冷熱錢包私鑰保管、FTX 倒閉教訓與 P2P 防詐 SOP。",
         "CoinMarketCap / 金管會 VASP 洗錢防制聲明名單 / 美國 SEC。",
         "問學生：『同學們平時在網路上或新聞中，有聽說過比特幣、以太幣或冷錢包嗎？』"),
        (2, 1, "第十六週 3 小時學習地圖與核心技能樹",
         "本頁為學生構建本週 3 小時的學習地圖。第一小時解密區塊鏈定義、比特幣 2,100 萬枚上限、BTC 減半經濟學、以太坊智慧合約、Web1-Web3 演進與冷熱錢包私鑰保管；第二小時聚焦於 DeFi 去中心化金融、NFT、USDT/USDC 穩定幣、FTX 倒閉教訓、台灣 VASP 監管、4 大風險與 5% 配置上限；第三小時傳授 Telegram 防詐 SOP、Web3 三大素養、RWA 資產代幣化、合法交易所 KYC 出入金與 4 大金律。",
         "115學年度「管理探索二」課程計畫與學習進度規劃網要。",
         "引導學生在講義上記下 Web3 與私鑰保管的關鍵鐵律，準備建立合理的數位資產配置觀念。"),
        (3, 1, "什麼是區塊鏈 (Blockchain)？去中心化分散式帳本技術",
         "講授區塊鏈 (Blockchain) 的核心技術架構：將交易數據包裝成「區塊 (Block)」，並透過密碼學哈希 (Hash) 串聯成「鏈 (Chain)」的分散式帳本技術。兩大特徵：1. **去中心化 (Decentralized)**：無須依賴中央銀行或仲介機構，全球數萬個節點共同記帳驗證；2. **不可篡改 (Immutability)**：歷史交易紀錄一旦寫入，無法被任何單一集體修改刪除。",
         "MIT Center for Future Blockchain Research。",
         "用生活比喻：『傳統記帳像全班只有班長一個人有帳本；區塊鏈像全班 50 個人人手一本同步更新的帳本，任何人想改答案都會被其他人抓包！』"),
        (4, 1, "比特幣 (Bitcoin BTC) 的誕生：2,100 萬枚總量上限抗通膨",
         "講授比特幣 (BTC) 的創世歷史：2008 年金融海嘯期間，化名中本聰 (Satoshi Nakamoto) 發布白皮書《點對點電子現金系統》。公式：`比特幣全球發行總量上限 = 21,000,000 枚 (永遠無法增發！)`。說明其硬性編碼的稀缺性，使其具備對抗中央銀行印鈔通膨的特質，被譽為「數位黃金」。",
         "Bitcoin Whitepaper (bitcoin.org)。",
         "提問：『各國央行在金融海嘯時無限量印鈔，與比特幣硬性限制 2,100 萬枚上限，在經濟學上形成了什麼強烈對比？（稀缺性 vs 通膨）』"),
        (5, 1, "比特幣減半機制 (BTC Halving) 與發行產量減半經濟學",
         "剖析比特幣減半 (BTC Halving) 的供給衝擊經濟學：每 21 萬個區塊 (約 4 年)，礦工挖礦的新幣獎勵減半一次。減半歷程：2009年 50 BTC ➔ 2012年 25 BTC ➔ 2016年 12.5 BTC ➔ 2020年 6.25 BTC ➔ **2024年 3.125 BTC！** 說明新幣發行速度減半帶來的供給卡位升力。",
         "Blockchain.com 比特幣網絡歷史數據。",
         "分析：『當新出廠的比特幣供給量每 4 年減半一次，而市場需求持續增加時，對價格會產生什麼影響？』"),
        (6, 1, "以太坊 (Ethereum ETH) 與智慧合約：DApp 去中心化應用",
         "講授以太坊 (ETH) 的升級革命：Vitalik Buterin 創立以太坊，引入**智慧合約 (Smart Contracts)**，即可程式化的自動執行合約程式碼。滿足條件即自動劃撥資產。以太坊不僅是貨幣，更是「全球去中心化世界電腦」，支撐起去中心化金融 (DeFi) 與 NFT 龐大生態系。",
         "Ethereum Foundation (ethereum.org)。",
         "比喻：『比特幣像功能單一的數位黃金；以太坊像智慧型手機 App Store，允許全球工程師在上面開發各種去中心化金融 App！』"),
        (7, 1, "Web1 -> Web2 -> Web3 演進：個人數位資產所有權革命",
         "解析網路文明演進的三大階段：1. **Web1 (1990s 唯讀 Read-only)**：靜態網頁與新聞，數據歸網站所有；2. **Web2 (2000s 讀寫 Read-Write)**：社群平台互動，數據歸 Meta/Google 巨頭所有；3. **Web3 (現今 讀寫與擁有 Read-Write-Own)**：個人擁有區塊鏈私鑰主權，資產完全由個人掌管！",
         "a16z Crypto《Web3 產業發展報告》。",
         "強調：『Web3 的核心革命，在於將數據與資產的所有權，從科技巨頭手中重新交還給每一個網路使用者！』"),
        (8, 1, "實證數據：CoinMarketCap 與金管會 VASP 登記專區",
         "引述 2026 最新實證數據：全球加密貨幣總市值突破 **2.5 兆美元 (約 80 兆台幣)**，比特幣市占率約 55%。介紹台灣金管會 VASP 洗錢防制專區，說明已完成洗錢防制法遵聲明之合法虛擬資產平台（如 MAX、ACE），客戶台幣資產享有本土銀行信託保管。",
         "CoinMarketCap / 金管會證期局 VASP 專區。",
         "指導學生：『投資加密貨幣，第一步必須查驗交易所是否在金管會合規名單上！』"),
        (9, 1, "冷錢包 (Cold Wallet) vs. 熱錢包 (Hot Wallet)：私鑰鐵律",
         "講授幣圈資安第一鐵律：**「Not Your Keys, Not Your Coins!」**。區分兩大錢包：1. **熱錢包 (Hot Wallet 聯網)**：手機 App / 網頁外掛 (Metamask)，方便但易遭木馬駭客攻擊；2. **冷錢包 (Cold Wallet 離線)**：硬體 USB 設備 (Ledger, Trezor)，私鑰完全離線儲存，防駭客遠端竊取。",
         "國家資通安全研究院 (NICS) 區塊鏈資安指引。",
         "警惕：『把加密貨幣放在交易所，私鑰是由交易所保管；唯有將幣提出至個人冷錢包，資產才真正屬於你！』"),
        (10, 1, "第一小時小結與專屬活動預告",
         "總結第一小時核心觀念：區塊鏈去中心化特徵；比特幣 2,100 萬枚上限與 4 年減半；以太坊智慧合約；Web3 資產主權；「Not Your Keys, Not Your Coins」冷錢包鐵律。預告第 1 小時 Modal 實務活動——「比特幣減半產量與冷熱錢包資安評估計算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行比特幣估值與冷熱錢包資安評級試算。"),

        # Hour 2 (Slide 11-20)
        (11, 2, "去中心化金融 (DeFi) 運作原理：流動性挖礦與 AMM 協議",
         "講授無須銀行的去中心化金融 (DeFi) 原理：1. **自動做市商 (AMM)**：如 Uniswap，透過數學公式 `x * y = k` 替代傳統券商搓合，實現無中介幣幣交易；2. **流動性挖礦 (Yield Farming)**：將加密貨幣注入資金池提供流動性，賺取手續費分潤與治理代幣。",
         "Uniswap Whitepaper & DeFi Pulse 數據。",
         "剖析：『DeFi 拿掉了傳統銀行的巨大中介利潤，將利息直接回饋給資金提供者，但同時需承擔合約漏洞風險。』"),
        (12, 2, "非同質化代幣 (NFT) 與數位藝術：智慧財產權與 Web3 憑證",
         "講授非同質化代幣 (NFT) 原理：具備獨一無二、不可替代特性的區塊鏈代幣。用途：數位藝術品、演唱會門票、遊戲虛寶與品牌會員憑證。智慧財產權：買方擁有區塊鏈上記錄的數位所有權證明與轉售分潤智慧合約。",
         "OpenSea 數位資產交易白皮書。",
         "提問：『為什麼一張 Jpeg 圖片做成 NFT 後能賣錢？（答案：賣的不是圖片本身，而是區塊鏈上獨一無二的數位所有權憑證）』"),
        (13, 2, "穩定幣 (Stablecoins) 經濟學：USDT/USDC 錨定與演算法暴雷",
         "解析連結法幣與加密世界的穩定幣：1. **法幣儲備型 (Fiat-backed)**：1:1 存入美元現金與短天期美債，如 **USDT (Tether), USDC (Circle)**，相對穩健；2. **演算法型 (Algorithmic)**：靠數學雙代幣套利維持價格，如 **UST (Luna 體系 2022 年歸零暴雷！)**。",
         "BIS 國際清算銀行穩定幣報告。",
         "警惕：『2022 年 UST 演算法穩定幣暴雷歸零，幾百億美元市值瞬間蒸發，給全球投資人上了慘痛的一課！』"),
        (14, 2, "重大歷史教訓：FTX 交易所倒閉事件 (2022) 剖析與資產隔離",
         "剖析 2022 年全球第二大加密交易所 FTX 72 小時倒閉事件：創辦人 SBF 違法挪用數十億美元客戶存款至旗下對沖基金 (Alameda) 炒股高槓桿爆倉，導致擠兌倒閉！核心教訓：絕對不要把 100% 資產放在中心化交易所，長期資產必須提領至冷錢包！",
         "美國德拉瓦州破產法院 FTX 案件訴訟卷宗。",
         "強調：『FTX 事件證明，未經監管的交易所再大都可能倒閉！必須嚴格執行資產分散與冷錢包保管。』"),
        (15, 2, "台灣 VASP 虛擬資產監管框架：洗錢防制 (AML) 與 KYC",
         "講授金管會納管台灣加密貨幣交易所 (VASP) 四大鐵律：1. **洗錢防制 (AML) 聲明**：必須向金管會完成聲明；2. **實名認證 (KYC)**：開戶須驗證雙證件；3. **客戶資產隔離信託**：台幣資產必須交由台灣本土銀行（如遠東、凱基）進行信託保管，嚴防挪用。",
         "金管會《虛擬資產平台及交易業務事業 (VASP) 指導原則》。",
         "提醒：『選台灣交易所一定要確認台幣資產是否有本土銀行信託保管！』"),
        (16, 2, "加密貨幣投資風險解析：波動性、Rug Pull 與私鑰遺失",
         "剖析幣圈 4 大致命風險：1. **價格劇烈波動**：單日暴漲暴跌 30% 是常態；2. **Rug Pull (拉地毯詐騙)**：匿名團隊發行新項目，募集資金後抽乾資金池逃跑；3. **黑客合約漏洞**：智慧合約遭攻擊；4. **私鑰遺失**：忘記助記詞，資產永遠無法找回。",
         "Chainalysis 全球加密貨幣犯罪報告。",
         "警告：『幣圈沒有「忘記密碼」按鈕！忘記助記詞，資產就等於永久消失。』"),
        (17, 2, "合理配置數位資產：小資族上限 (5%-10%) 避險戰術",
         "講授小資大一新鮮人數位資產避險配置公式：`加密貨幣占個人總投資資產比率 <= 5% ~ 10% (絕對不要借錢槓桿！)`。戰術：95% 資金配置於 0050/00878 等傳統優質 ETF；僅用 5% 閒錢參與加密貨幣，歸零不影響生活，大漲享爆發力！",
         "理財規劃師 (CFP) 數位資產風險控管模型。",
         "結論：『用 5% 的閒錢去博取潛在收益，永遠保留 95% 的本金安全，這是最聰明的黑天鵝避險戰術！』"),
        (18, 2, "全球現貨 ETF 批准里程碑：美國 SEC 批准 BTC/ETH 現貨 ETF",
         "講授加密貨幣進入華爾街主流金融體系的歷史性突破：2024 年美國 SEC 正式批准貝萊德 (BlackRock)、富達等巨頭推出的 **BTC 現貨 ETF 與 ETH 現貨 ETF** 上市！機構資金透過傳統美股帳戶即可合規買賣比特幣。",
         "美國證券交易委員會 (SEC) 批准公告。",
         "分析：『現貨 ETF 的通過，代表傳統法幣資本市場正式為比特幣打開了萬億美元的資金大門！』"),
        (19, 2, "央行數位貨幣 (CBDC) 趨勢：數位新台幣與去中心化對決",
         "解析國家級央行數位貨幣 (CBDC) 趨勢：由國家中央銀行直接發行背書的數位法定貨幣（如中央銀行實驗中的「數位新台幣 e-NTD」）。區分：CBDC 為**高度中心化監管**；比特幣為**去中心化抗審查**。兩者平行發展。",
         "中華民國中央銀行 CBDC 研究報告。",
         "提問：『數位新台幣 (CBDC) 與比特幣最大的本質區別是什麼？（答案：CBDC 仍由央行中心化發行控管，比特幣去中心化）』"),
        (20, 2, "第二小時小結與專屬活動預告",
         "總結第二小時核心觀念：DeFi AMM 協議；USDT/USDC 穩定幣；FTX 倒閉教訓；台灣 VASP 銀行信託；4 大致命風險；5% 配置上限；美股現貨 ETF 上市。預告第 2 小時 Modal 實務活動——「DeFi 年化收益風險與數位資產配置試算器」。",
         "115管理探索二團隊建置。",
         "指導學生開啟網頁版 Modal，準備進行 5% 數位資產避險配置金額試算演練。"),

        # Hour 3 (Slide 21-30)
        (21, 3, "大一新鮮人資安防禦：防範假交易所與 Telegram 幣圈詐騙",
         "講授幣圈防詐實務：大學生最易陷入的 3 大陷阱：1. Telegram 私訊「美女客服搬磚套利」；2. 假冒 MAX 交易所釣魚網址；3. Line 群組「保證獲利 100%」。破解 2 鐵律：**絕對不點陌生網址；交易僅限金管會 VASP 核准之正規交易所！**",
         "內政部警政署 165 反詐騙專線幣圈防詐指引。",
         "警告：『Telegram 上主動私訊你的正妹客服或交易大師，1000% 是詐騙集團！』"),
        (22, 3, "大一新鮮人面對 Web3 時代 3 大素養：密碼、資安與風險",
         "講授 Web3 時代三大硬核素養：1. **密碼學與私鑰保管素養**：落實「Not Your Keys, Not Your Coins」；2. **資安自治與網址查驗能力**：每筆交易檢查 Smart Contract 授權與 URL 域名；3. **風險控制與紀律執行力**：嚴格限制 5% 資金上限，拒絕合約槓桿賭博。",
         "Stanford Center for Blockchain Research 指南。",
         "強調：『Web3 帶來了絕對的資產主權，但也意味著你必須對自己的資安負 100% 的責任！』"),
        (23, 3, "區塊鏈在傳統金融的應用：跨國清算與資產代幣化 (RWA)",
         "講授傳統金融巨頭擁抱區塊鏈趨勢：**真實世界資產代幣化 (RWA)**，將美債、房地產與黃金轉化為區塊鏈代幣（如 BlackRock BUIDL 美債基金）。優勢：實現 24/7 全天候即時清算、微型化分割投資與低手續費。",
         "BlackRock 貝萊德 RWA 白皮書。",
         "展現：『連全球最大的資產管理公司貝萊德都在發行 RWA 代幣化美債，證明區塊鏈技術的實用價值！』"),
        (24, 3, "幣圈出金/入金法幣 SOP：台灣合法監管交易所 KYC 驗證",
         "手把手教學法幣安全出入金 SOP：1. 註冊台灣合法 VASP 交易所 (如 MAX)；2. 完成實名認證 (KYC) 與銀行帳戶綁定；3. 使用綁定銀行帳戶台幣入金（享銀行信託保管）；4. 出金回綁定台幣銀行帳戶，依法留存紀錄防止洗錢。",
         "台灣虛擬資產同業公會 (VASP Guild) 安全指南。",
         "提醒：『絕對不要進行未經 KYC 的 P2P 私下場外現金交易，否則容易淪為洗錢人頭帳戶遭凍結！』"),
        (25, 3, "實證數據調取：CoinGecko (coingecko.com) 實務查閱",
         "手把手教導學生登入 CoinGecko 官網：1. 登入 `coingecko.com` 查閱 BTC / ETH 流通量與歷史走勢；2. 在 CoinGecko 複製正版 Token 合約地址防範假代幣；3. 對照金管會 VASP 洗錢防制專區交叉驗證交易所合法性。",
         "CoinGecko (coingecko.com) / 金管會證期局。",
         "演示進入 CoinGecko 網站複製官方正確合約地址的步驟。"),
        (26, 3, "大一新鮮人 Web3 與數位資產 4 大金律",
         "總結 Web3 與數位資產 4 大金律：1. **私鑰即資產** (Not Your Keys, Not Your Coins，長期資產提領至冷錢包)；2. **數位資產嚴格限制 5% 比率** (絕不借錢槓桿賭博)；3. **交易僅限金管會 VASP 合法平台** (認明銀行信託保管)；4. **嚴防 Telegram 與 P2P 幣圈詐騙** (不點擊陌生網址)。",
         "金管會與區塊鏈資安保護原則。",
         "請學生齊聲朗讀四金律，建立強固的幣圈資安防線。"),
        (27, 3, "數位資產配置試算實例：5% 避險配置兼顧爆發力與本金安全",
         "展示大一新鮮人小明 10 萬元資產配置 SOP：傳統資產 95% (95,000 元，定期定額買 0050+00878 奠定根基)；數位資產 5% (5,000 元，定期定額買 BTC+ETH 提領至冷錢包)。避險效果：幣圈暴跌最大損失僅 5,000 元，完全不傷生活，大漲共享爆發力！",
         "115管理探索二教案數位資產組。",
         "強調：『95% 穩健 + 5% 爆發力，這是大一新鮮人參與 Web3 數位資產的最完美配置！』"),
        (28, 3, "Web3 數位資產與未來金融技術總整合",
         "總結 Web3 與區塊鏈的未來價值：區塊鏈正在重塑全球金融基礎設施。大一新鮮人身為數位原住民，應保持理性態度：技術為我所用，理解去中心化優勢；同時堅守資安防線與 5% 資金上限，成為有智慧的數位金融公民！",
         "115學年度「管理探索二」核心價值。",
         "激勵學生：『駕馭 Web3 新技術，開創安全的數位金融未來！』"),
        (29, 3, "第十六週全景知識體系圖與觀念整合",
         "以全景邏輯結構圖將第十六週 30 頁純教學卡片進行整體串聯：Hour 1 (區塊鏈定義 ➔ 比特幣2100萬枚上限 ➔ BTC減半經濟學 ➔ 以太坊智慧合約 ➔ Web1-Web3演進 ➔ CoinMarketCap數據 ➔ 冷熱錢包私鑰保管) ➔ Hour 2 (DeFi AMM協議 ➔ NFT數位資產 ➔ USDT/USDC穩定幣 ➔ FTX倒閉教訓 ➔ 台灣VASP四大監管 ➔ 4大風險 ➔ 5%配置上限 ➔ 美國現貨ETF ➔ CBDC) ➔ Hour 3 (Telegram防詐 ➔ Web3三大素養 ➔ RWA資產代幣化 ➔ 法幣合法KYC出入金 ➔ CoinGecko Portal查閱 ➔ 5%配置SOP ➔ 4大金律)。",
         "115學年度「管理探索二」課程組全景圖。",
         "引導學生回顧整週學習歷程，確認每一個節點均已融會貫通。"),
        (30, 3, "學習反思與第十七週預告 (期末專案發表、實戰財經簡報與總評量)",
         "恭喜學生完成第十六週學習！指導學生點擊 navbar 上方的「📝 本週課堂作業」按鈕，線上填寫並提交「個人數位資產風險評估與 Web3 報告」。預告第十七週課程主題：「期末專案發表、實戰財經簡報與總評量」，下週學生團隊將同台發表大一理財專案簡報！",
         "115學年度「管理探索二」課程計畫。",
         "提醒學生課後完成作業提交，並與小組成員準備第十七週期末專案簡報發表。")
    ]

    for slide in slides_breakdown:
        s_id, s_hour, s_title, s_content, s_data, s_q = slide
        
        # Slide Header Title
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(12)
        p_s.paragraph_format.space_after = Pt(4)
        run_s = p_s.add_run(f"Slide {s_id:02d} (第 {s_hour} 小時) | {s_title}")
        run_s.font.name = 'Microsoft JhengHei'
        run_s.font.size = Pt(12)
        run_s.font.bold = True
        run_s.font.color.rgb = SECONDARY_BLUE

        # Slide Table Box
        t_slide = doc.add_table(rows=3, cols=2)
        t_slide.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Col Widths
        t_slide.columns[0].width = Inches(1.8)
        t_slide.columns[1].width = Inches(4.8)

        labels = ["【講授核心與生活化說明】", "【官方數據引述】", "【課堂提問與引導技巧】"]
        contents = [s_content, s_data, s_q]

        for idx in range(3):
            cell_lbl = t_slide.rows[idx].cells[0]
            cell_val = t_slide.rows[idx].cells[1]

            cell_lbl.text = labels[idx]
            cell_val.text = contents[idx]

            set_cell_background(cell_lbl, "F2F4F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

            p_l = cell_lbl.paragraphs[0]
            p_l.runs[0].font.name = 'Microsoft JhengHei'
            p_l.runs[0].font.bold = True
            p_l.runs[0].font.size = Pt(9.5)
            p_l.runs[0].font.color.rgb = PRIMARY_NAVY

            p_v = cell_val.paragraphs[0]
            p_v.runs[0].font.name = 'Microsoft JhengHei'
            p_v.runs[0].font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PART 4 ---
    h4 = doc.add_heading("四、3大小時活動與課堂作業詳細執行手冊", level=1)
    h4.runs[0].font.name = 'Microsoft JhengHei'
    h4.runs[0].font.size = Pt(15)
    h4.runs[0].font.color.rgb = PRIMARY_NAVY

    act_text = [
        ("🎯 第 1 小時實務活動：比特幣減半產量與冷熱錢包資安評估計算器",
         "1. 活動目標：幫助學生計算比特幣持有資產之法幣估值，並評估冷錢包與交易所熱錢包之資安等級。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入比特幣數量與保管錢包類型，點擊精算。\n"
         "3. 診斷反思：系統算出台幣估值與資安評級，驗證「Not Your Keys, Not Your Coins，大額資產提領至冷錢包」之結論。"),

        ("🎯 第 2 小時實務活動：DeFi 年化收益風險與數位資產配置試算器",
         "1. 活動目標：讓學生親自精算個人總資產中配置 5% 加密貨幣之避險金額與風險曝險控制。\n"
         "2. 操作步驟：學生開啟網頁版 Modal，輸入總投資資產與預計加密貨幣比率，點擊精算。\n"
         "3. 決策學習：系統自動產出 95% 傳統資產與 5% 數位資產金額，深化「即便加密貨幣歸零也不傷害生活根基」的控制觀念。"),

        ("🎮 第 3 小時小遊戲：Web3 數位資產防詐大挑戰 (4 大關卡)",
         "1. 遊戲機制：包含 4 大主題關卡（關卡 1：比特幣 2100 萬枚上限；關卡 2：Not Your Keys, Not Your Coins 鐵律；關卡 3：台灣 VASP 洗錢防制金管會監管；關卡 4：5% 數位資產避險上限）。\n"
         "2. 評分與徽章：每答對一題得 100 分，總分 400 分。滿分獲頒「🏆 Web3 數位資產達人徽章 (Web3 Master)」，未滿分獲頒「🥉 區塊鏈初學者徽章」，附帶即時詳細解析。"),

        ("📝 第 3 小時課堂實務作業：個人數位資產風險評估與 Web3 報告",
         "1. 作業題目：請學生評估加密貨幣風險，設計專屬 5% 數位資產配置與資安防禦 SOP。\n"
         "2. 分析要項：(1) 說明區塊鏈去中心化與比特幣 2,100 萬枚上限抗通膨原理；(2) 剖析 FTX 倒閉教訓與冷錢包保管私鑰鐵律；(3) 規劃 95% 傳統資產 + 5% 數位資產配置方案；(4) 擬定防範 Telegram 詐騙與合法交易所 KYC 出入金 SOP。\n"
         "3. 繳交方式：於網頁版 Modal 表單填寫完成後點擊提交，教官與助教將於 3 日內完成線上評閱。")
    ]

    for title, desc in act_text:
        p_act_t = doc.add_paragraph()
        p_act_t.paragraph_format.space_before = Pt(10)
        p_act_t.paragraph_format.space_after = Pt(4)
        r_act_t = p_act_t.add_run(title)
        r_act_t.font.name = 'Microsoft JhengHei'
        r_act_t.font.size = Pt(12)
        r_act_t.font.bold = True
        r_act_t.font.color.rgb = SECONDARY_BLUE

        p_act_d = doc.add_paragraph()
        p_act_d.paragraph_format.line_spacing = 1.25
        p_act_d.paragraph_format.space_after = Pt(10)
        r_act_d = p_act_d.add_run(desc)
        r_act_d.font.name = 'Microsoft JhengHei'
        r_act_d.font.size = Pt(10)

    # Save Document
    doc_path = r'C:\Users\User\Desktop\115學年度\管理探索二\第十六週_課程教學指引_加密貨幣區塊鏈基礎Web3與數位資產風險.docx'
    doc.save(doc_path)
    print("Created 第十六週_課程教學指引_加密貨幣區塊鏈基礎Web3與數位資產風險.docx successfully!")

if __name__ == '__main__':
    create_rich_word_guide()
